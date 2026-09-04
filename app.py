import streamlit as st
import pandas as pd
import numpy as np
import joblib

# =========================================================
# 화학물질 관용명 → MSDS API 정식 국문명 변환표
# 형식: "사용자가 입력할 수 있는 이름": "API 검색용 정식 국문명"
# =========================================================

CHEMICAL_ALIASES = {
    # 알코올류
    "메탄올": "메틸알코올",
    "메틸알콜": "메틸알코올",
    "메틸 알콜": "메틸알코올",
    "메틸 알코올": "메틸알코올",

    "에탄올": "에틸알코올",
    "에틸알콜": "에틸알코올",
    "에틸 알콜": "에틸알코올",
    "에틸 알코올": "에틸알코올",

    "이소프로판올": "이소프로필알코올",
    "아이소프로판올": "이소프로필알코올",
    "IPA": "이소프로필알코올",
    "아이피에이": "이소프로필알코올",
    "이소프로필알콜": "이소프로필알코올",
    "자일렌": "크실렌",
    

    # 산류
    "염산": "염화수소",
    "빙초산": "아세트산",
    "초산": "아세트산",
    "식초산": "아세트산",
    "아세트산에틸": "초산에틸",
    "에틸렌아세테이트": "초산에틸",
    "아세트산메틸": "초산메틸",
    "메틸렌아세테이트": "초산메틸",
    "불산": "불화수소",
    "플루오르화수소": "불화수소",
    

    # 알칼리류
    "가성소다": "수산화나트륨",
    "가성소다액": "수산화나트륨",
    "소다회": "탄산나트륨",
    "가성칼리": "수산화칼륨",
    "가성가리": "수산화칼륨",
    "차염": "차아염소산나트륨",
    "차염소산나트륨": "차아염소산나트륨",

    # 용제류
    "MEK": "메틸에틸케톤",
    "엠이케이": "메틸에틸케톤",
    "메틸에틸케톤": "메틸에틸케톤",

    "MIBK": "메틸이소부틸케톤",
    "엠아이비케이": "메틸이소부틸케톤",
    "메틸아이소부틸케톤": "메틸이소부틸케톤",

    "MC": "염화메틸렌",
    "메틸렌클로라이드": "염화메틸렌",
    "디클로로메탄": "염화메틸렌",
    "ECH": "에피클로로히드린",
    "1,2-디클로로벤젠": "o-디클로로벤젠",

    # 기타 현장 사용명
    "과산화수소수": "과산화수소",
    "과수": "과산화수소",
    "암모니아수": "암모니아 용액",
    "암수": "암모니아 용액",
    "수산화 암모늄": "암모니아 용액",
    "포르말린": "포름알데히드",
    "포름알린": "포름알데히드",
    "포름산": "개미산",
    "MDI": "메틸렌 디페닐 디이소시아네이트",
    "디이소시안산디페닐메탄": "메틸렌 디페닐 디이소시아네이트",
    "NMP": "1-메틸-2-피롤리디논",
    "TDI": "톨루엔-2,4/2,6-디이소시아네이트",
}
import requests
import xml.etree.ElementTree as ET
from datetime import datetime, timedelta, time as dt_time
import time
import os
from pathlib import Path
import textwrap
import streamlit.components.v1 as components
import re
import base64
import copy
from PIL import Image
from difflib import SequenceMatcher
from openai import OpenAI
from supabase import create_client
from io import BytesIO
from docxtpl import DocxTemplate
from docx.shared import Emu
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from streamlit_js_eval import streamlit_js_eval
from streamlit_searchbox import st_searchbox
from streamlit_drawable_canvas import st_canvas

# =========================
# 1) 파일 불러오기
# =========================
model = joblib.load("model_histgb_v2.pkl")
model_columns = joblib.load("model_columns_histgb_v2.pkl")
mapping_df = pd.read_excel("hazard_mapping.xlsx")

# TBM 문구 DB 불러오기
risk_message_df = pd.read_excel("tbm_message_db.xlsx", sheet_name="risk_message_db")
accident_case_df = pd.read_excel("tbm_message_db.xlsx", sheet_name="accident_case_db")
final_result = pd.read_excel(
    "final_result.xlsx",
    dtype={"CHEMID": str}
)

final_result.columns = final_result.columns.str.strip()
final_result.columns = final_result.columns.str.strip()

ACCIDENT_ARCHIVE_PATH = "한국산업안전보건공단_산업재해 고위험요인 아카이브_20260401.xlsx"
ACCIDENT_ARCHIVE_COLUMNS = ["재해개요", "기인물", "고위험작업·상황", "재해유발요인", "위험성 감소대책(예시)"]


@st.cache_data
def load_accident_archive():
    """산업재해 고위험요인 아카이브(제조업등 시트)를 불러온다.
    3번째 행이 실제 헤더라 header=2로 읽는다. AI 프롬프트 참고자료로만 쓰이며
    앱 실행 중 반복 로드를 피하기 위해 st.cache_data로 캐싱한다."""
    df = pd.read_excel(
        ACCIDENT_ARCHIVE_PATH,
        sheet_name="아카이브(제조업등)",
        header=2
    )

    df = df.dropna(subset=["재해개요"]).reset_index(drop=True)

    for col in ACCIDENT_ARCHIVE_COLUMNS:
        df[col] = df[col].astype(str).str.strip()

    return df


accident_archive_df = load_accident_archive()

# 엑셀 컬럼명 앞뒤 공백 제거
risk_message_df.columns = risk_message_df.columns.str.strip()
accident_case_df.columns = accident_case_df.columns.str.strip()


st.set_page_config(
    page_title="화학안전 TBM",
    page_icon="🛡️",
    layout="centered",
    initial_sidebar_state="collapsed"
)

st.markdown(
    '<link rel="stylesheet" href="https://fonts.googleapis.com/css2?family=Material+Symbols+Rounded" />',
    unsafe_allow_html=True
)

st.markdown("""
<style>
/* 전체 배경 */
.stApp {
    background: #f1f5f9;
}

/* Streamlit 기본 여백
   상단 여백은 상단바(고정 헤더)가 자체 spacer로 확보하므로 여기서는
   Streamlit 기본 헤더(높이 60px)를 가릴 정도만 최소로 잡는다. */
.block-container {
    max-width: 480px;
    padding-top: 0.25rem;
    padding-bottom: 7rem;
}

/* 상단 로고 영역 */
.login-header {
    text-align: center;
    margin-bottom: 18px;
}

.login-logo-row {
    display: flex;
    justify-content: center;
    align-items: center;
    gap: 8px;
    margin-bottom: 6px;
}

.login-logo-icon {
    width: 34px;
    height: 34px;
    border-radius: 10px;
    background: #2170e4;
    color: white;
    display: flex;
    align-items: center;
    justify-content: center;
    font-size: 19px;
    font-weight: 800;
}

.login-title {
    font-size: 25px;
    font-weight: 800;
    color: #091426;
}

.login-subtitle {
    font-size: 14px;
    color: #45474c;
}

/* 메인 카드 */
.login-card {
    background: #ffffff;
    border: 1px solid #c5c6cd;
    border-radius: 18px;
    padding: 28px;
    box-shadow: 0 8px 30px rgba(0,0,0,0.08);
    margin-bottom: 20px;
}

/* 모드 카드 */
.mode-card-wrap {
    display: flex;
    gap: 10px;
    margin-bottom: 18px;
}

.mode-card {
    flex: 1;
    border-radius: 14px;
    padding: 16px 10px;
    text-align: center;
    border: 2px solid transparent;
    background: #f0edef;
    color: #45474c;
    font-weight: 700;
}

.mode-card-active {
    background: #2170e4;
    color: #ffffff;
    border: 2px solid #2170e4;
}

.mode-icon {
    font-size: 25px;
    margin-bottom: 5px;
}

.mode-label {
    font-size: 14px;
}

/* 안내 박스 */
.info-box {
    background: #f5f3f4;
    border: 1px solid rgba(197,198,205,0.5);
    border-radius: 14px;
    padding: 14px;
    display: flex;
    gap: 10px;
    margin-top: 18px;
    margin-bottom: 22px;
}

.info-icon {
    color: #ba1a1a;
    font-size: 20px;
    line-height: 1.4;
}

.info-text {
    color: #45474c;
    font-size: 13px;
    line-height: 1.55;
}

/* 입력창 */
div[data-baseweb="input"] {
    border-radius: 12px;
}

/* 기본 버튼: "TBM 회의록 출력하기"(st.download_button 기본 스타일)와 톤을 맞춰
   흰 배경 + 얇은 테두리로 통일한다. */
div.stButton > button {
    border-radius: 12px;
    font-size: 16px;
    font-weight: 800;
    background: #f7f7f7;
    color: #1b1b1d;
    border: 1px solid #d7d9e0;
    min-height: 46px;
}

div.stButton > button:hover {
    background: #f7f8fa;
    border-color: #b9bcc6;
    color: #1b1b1d;
}

/* 하단 푸터 */
.login-footer {
    text-align: center;
    margin-top: 28px;
    color: #75777d;
    font-size: 12px;
}

/* 상단 그라데이션 */
.top-gradient-line {
    position: fixed;
    top: 0;
    left: 0;
    width: 100%;
    height: 4px;
    background: linear-gradient(90deg, #2170e4, #091426, #2170e4);
    z-index: 9999;
}

/* 하단 네비게이션 spacer는 show_bottom_nav()에서 실제 바 높이에 맞춰 정의됨 */

/* =========================
   위험도 결과 화면 전체 스타일
========================= */


.result-page-topbar {
    position: sticky;
    top: 2px;
    z-index: 100;
    background: #fbf8fa;
    border-bottom: 1px solid #c5c6cd;
    padding: 14px 4px 12px 4px;
    margin-bottom: 20px;
}

.result-topbar-row {
    display: flex;
    align-items: center;
    justify-content: space-between;
}

.result-topbar-left {
    display: flex;
    align-items: center;
    gap: 10px;
}

.result-back-icon {
    width: 36px;
    height: 36px;
    border-radius: 999px;
    background: #f0edef;
    display: flex;
    align-items: center;
    justify-content: center;
    color: #091426;
    font-weight: 800;
}

.result-app-title {
    font-size: 22px;
    font-weight: 800;
    color: #091426;
}

.result-report-label {
    font-size: 11px;
    letter-spacing: 0.08em;
    font-weight: 800;
    color: #75777d;
    margin-bottom: 4px;
}

.result-title {
    font-size: 30px;
    font-weight: 900;
    color: #091426;
    line-height: 1.25;
    margin-bottom: 6px;
}

.result-subtitle {
    font-size: 14px;
    color: #45474c;
    line-height: 1.45;
    margin-bottom: 18px;
}

.result-card {
    background: #ffffff;
    border: 1px solid #c5c6cd;
    border-radius: 18px;
    padding: 18px;
    margin-bottom: 16px;
    box-shadow: 0 4px 14px rgba(15, 23, 42, 0.07);
}

.result-card-header {
    display: flex;
    align-items: center;
    gap: 8px;
    padding-bottom: 12px;
    border-bottom: 1px solid #e5e7eb;
    margin-bottom: 4px;
}

.result-card-icon {
    font-size: 20px;
}

.result-card-title {
    font-size: 18px;
    font-weight: 900;
    color: #091426;
}

.message-item {
    display: flex;
    gap: 10px;
    padding: 12px 0;
    border-bottom: 1px solid #e5e7eb;
}

.message-item:last-child {
    border-bottom: none;
}

.message-num {
    font-size: 13px;
    font-weight: 900;
    color: #ba1a1a;
    min-width: 26px;
}

.message-text {
    font-size: 14px;
    line-height: 1.5;
    color: #1b1b1d;
}

.measure-item {
    display: flex;
    gap: 10px;
    padding: 11px 0;
    border-bottom: 1px solid #e5e7eb;
}

.measure-item:last-child {
    border-bottom: none;
}

.measure-check {
    color: #2170e4;
    font-weight: 900;
    min-width: 20px;
}

.measure-text {
    font-size: 14px;
    line-height: 1.5;
    color: #1b1b1d;
}

/* 중점위험요인/대책 카드(버튼) 선택 상태: 배경은 그대로 흰색 유지, 테두리만 강조 */
div[data-testid="stButton"] button[kind="primary"] {
    background-color: #ffffff !important;
    color: #1b1b1d !important;
    border: 3px solid #2170e4 !important;
    box-shadow: 0 0 0 1px rgba(33, 112, 228, 0.15) !important;
}

div[data-testid="stButton"] button[kind="primary"]:hover {
    background-color: #f5f9ff !important;
    color: #1b1b1d !important;
    border-color: #2170e4 !important;
}

div[data-testid="stButton"] button[kind="primary"] p {
    color: #1b1b1d !important;
}

/* 중점위험요인/대책 카드(버튼) 선택 안 된 상태: 앱 기본 버튼 색(파란 배경)을 덮어써서
   흰 배경 + 검은 글씨로 표시한다. Streamlit이 각 위젯을 st-key-{key} 클래스가 붙은
   stElementContainer로 감싸주므로, 버튼의 key 접두사로 카드 버튼만 정확히 골라낸다. */
div[class*="st-key-tc_selected_hazard_idx_btn_"] button[kind="secondary"],
div[class*="st-key-tc_selected_measure_idx_btn_"] button[kind="secondary"] {
    background-color: #ffffff !important;
    color: #1b1b1d !important;
    border: 1px solid #d7d9e0 !important;
}

div[class*="st-key-tc_selected_hazard_idx_btn_"] button[kind="secondary"] p,
div[class*="st-key-tc_selected_measure_idx_btn_"] button[kind="secondary"] p {
    color: #1b1b1d !important;
}

div[class*="st-key-tc_selected_hazard_idx_btn_"] button[kind="secondary"]:hover,
div[class*="st-key-tc_selected_measure_idx_btn_"] button[kind="secondary"]:hover {
    background-color: #f7f8fa !important;
    border-color: #b9bcc6 !important;
}

.incident-placeholder {
    background: linear-gradient(135deg, #091426, #1e293b);
    color: white;
    border-radius: 14px;
    padding: 18px;
    margin-top: 12px;
}

.incident-placeholder-title {
    font-size: 15px;
    font-weight: 900;
    margin-bottom: 8px;
}

.incident-placeholder-desc {
    font-size: 13px;
    line-height: 1.45;
    opacity: 0.9;
}

.result-info-caption {
    font-size: 12px;
    color: #75777d;
    margin-top: 8px;
    margin-bottom: 14px;
}

/* 작업정보 입력 화면 */
.app-topbar {
    position: fixed;
    top: 4px;
    left: 50%;
    transform: translateX(-50%);

    width: min(480px, calc(100% - 28px));

    z-index: 9999;

    background: #fbf8fa;
    border-bottom: 1px solid #c5c6cd;

    padding: 14px 4px 12px 4px;

    box-shadow: 0 4px 14px rgba(15,23,42,0.08);
}

.topbar-row {
    display: flex;
    align-items: center;
    justify-content: space-between;
}

.topbar-left {
    display: flex;
    align-items: center;
    gap: 10px;
}

.back-btn {
    width: 36px;
    height: 36px;
    border-radius: 999px;
    background: #f0edef;
    display: flex;
    align-items: center;
    justify-content: center;
    color: #091426;
    font-weight: 800;
}

.app-title {
    font-size: 22px;
    font-weight: 800;
    color: #091426;
}

.hero-title {
    font-size: 27px;
    line-height: 1.25;
    font-weight: 800;
    color: #091426;
    margin-bottom: 8px;
}

.hero-subtitle {
    font-size: 14px;
    color: #45474c;
    margin-bottom: 22px;
}


.field-label {
    font-size: 12px;
    letter-spacing: 0.05em;
    font-weight: 800;
    color: #45474c;
    margin-bottom: 8px;
}

/* 작업 시작 시간(시/분) 숫자 입력 두 칸을 나란히 배치한다. 이 앱은
   max-width: 480px짜리 모바일 폭 레이아웃이라, st.columns를 그대로 두면
   Streamlit이 좁은 화면에서 컬럼을 자동으로 세로로 쌓아버리므로 막아준다. */
div[class*="st-key-tc_time_input_row"] div[data-testid="stHorizontalBlock"] {
    flex-wrap: nowrap !important;
}

.selected-chip {
    display: inline-flex;
    align-items: center;
    gap: 5px;
    padding: 6px 12px;
    border-radius: 999px;
    background: rgba(33,112,228,0.10);
    color: #0058be;
    border: 1px solid rgba(0,88,190,0.20);
    font-size: 13px;
    font-weight: 700;
    margin-top: 8px;
}

.plant-image-card {
    margin-top: 24px;
    min-height: 180px;
    border-radius: 18px;
    overflow: hidden;
    background: linear-gradient(135deg, #091426, #2170e4);
    color: white;
    padding: 24px;
    display: flex;
    align-items: flex-end;
    font-size: 14px;
    line-height: 1.5;
    box-shadow: 0 6px 18px rgba(15, 23, 42, 0.14);
}

</style>
""", unsafe_allow_html=True)


st.markdown("""
<style>

/* =========================
   TBM 체크리스트 화면
========================= */

.checklist-topbar {
    position: sticky;
    top: 4px;
    z-index: 100;
    background: #fbf8fa;
    border-bottom: 1px solid #c5c6cd;
    padding: 14px 4px 12px 4px;
    margin-bottom: 20px;
}

.checklist-topbar-row {
    display: flex;
    align-items: center;
    justify-content: space-between;
}

.checklist-topbar-left {
    display: flex;
    align-items: center;
    gap: 10px;
}

.checklist-back-icon {
    width: 36px;
    height: 36px;
    border-radius: 999px;
    background: #f0edef;
    display: flex;
    align-items: center;
    justify-content: center;
    color: #091426;
    font-weight: 800;
}

.checklist-app-title {
    font-size: 22px;
    font-weight: 900;
    color: #091426;
}

.checklist-info-card {
    background: #ffffff;
    border: 1px solid #c5c6cd;
    border-radius: 18px;
    padding: 20px;
    margin-bottom: 20px;
    box-shadow: 0 4px 14px rgba(15, 23, 42, 0.07);
}

.checklist-badge {
    display: inline-block;
    background: #d8e2ff;
    color: #0058be;
    padding: 6px 12px;
    border-radius: 999px;
    font-size: 11px;
    font-weight: 900;
    letter-spacing: 0.06em;
    margin-bottom: 10px;
}

.checklist-title {
    font-size: 25px;
    font-weight: 900;
    color: #091426;
    line-height: 1.25;
    margin-bottom: 6px;
}

.checklist-subtitle {
    font-size: 14px;
    color: #45474c;
    line-height: 1.45;
}

.checklist-section-title {
    font-size: 19px;
    font-weight: 900;
    color: #091426;
    margin: 22px 0 12px 0;
    display: flex;
    align-items: center;
    gap: 8px;
}

.checklist-item-card {
    background: #ffffff;
    border: 1px solid #c5c6cd;
    border-radius: 16px;
    padding: 14px 14px;
    margin-bottom: 10px;
    box-shadow: 0 3px 10px rgba(15, 23, 42, 0.05);
}

.checklist-remark-card {
    background: #ffffff;
    border: 1px solid #c5c6cd;
    border-radius: 16px;
    padding: 16px;
    margin-bottom: 18px;
    box-shadow: 0 3px 10px rgba(15, 23, 42, 0.05);
}

.checklist-small-label {
    font-size: 11px;
    color: #75777d;
    font-weight: 900;
    letter-spacing: 0.08em;
    margin-bottom: 8px;
}

.checklist-save-note {
    font-size: 12px;
    color: #75777d;
    line-height: 1.45;
    margin-top: 8px;
    margin-bottom: 12px;
}

</style>
""", unsafe_allow_html=True)


st.markdown("""
<style>

/* =========================
   나의 작업일지 화면
========================= */


.journal-topbar {
    position: sticky;
    top: 4px;
    z-index: 100;
    background: #fbf8fa;
    border-bottom: 1px solid #c5c6cd;
    padding: 14px 4px 12px 4px;
    margin-bottom: 20px;
}

.journal-topbar-row {
    display: flex;
    align-items: center;
    justify-content: space-between;
}

.journal-topbar-left {
    display: flex;
    align-items: center;
    gap: 10px;
}

.journal-back-icon {
    width: 36px;
    height: 36px;
    border-radius: 999px;
    background: #f0edef;
    display: flex;
    align-items: center;
    justify-content: center;
    color: #091426;
    font-weight: 800;
}

.journal-app-title {
    font-size: 22px;
    font-weight: 900;
    color: #091426;
}

.journal-title {
    font-size: 30px;
    font-weight: 900;
    color: #091426;
    line-height: 1.25;
    margin-bottom: 6px;
}

.journal-subtitle {
    font-size: 14px;
    color: #45474c;
    line-height: 1.45;
    margin-bottom: 20px;
}

.journal-card {
    background: #ffffff;
    border: 1px solid #c5c6cd;
    border-radius: 18px;
    padding: 18px;
    margin-bottom: 16px;
    box-shadow: 0 4px 14px rgba(15, 23, 42, 0.07);
}

.journal-card-header {
    display: flex;
    align-items: center;
    justify-content: space-between;
    margin-bottom: 14px;
}

.journal-card-title {
    font-size: 18px;
    font-weight: 900;
    color: #091426;
}

.journal-summary-row {
    display: flex;
    justify-content: space-between;
    align-items: center;
    border-bottom: 1px solid #e5e7eb;
    padding: 10px 0;
    gap: 12px;
}

.journal-summary-label {
    font-size: 12px;
    color: #75777d;
    font-weight: 900;
    letter-spacing: 0.05em;
}

.journal-summary-value {
    font-size: 14px;
    color: #091426;
    font-weight: 800;
    text-align: right;
}

.journal-risk-badge {
    display: inline-block;
    padding: 5px 10px;
    border-radius: 8px;
    background: #ffdad6;
    color: #93000a;
    font-size: 12px;
    font-weight: 900;
}

.journal-message-box {
    background: #f5f3f4;
    border-radius: 14px;
    padding: 14px;
    font-size: 14px;
    line-height: 1.5;
    color: #1b1b1d;
    margin-top: 12px;
}

.journal-small-label {
    font-size: 11px;
    color: #75777d;
    font-weight: 900;
    letter-spacing: 0.08em;
    margin-bottom: 8px;
}


.journal-submit-card {
    background: rgba(33, 112, 228, 0.06);
    border: 1px solid rgba(33, 112, 228, 0.25);
    border-radius: 18px;
    padding: 18px;
    margin-top: 20px;
    margin-bottom: 18px;
}

.journal-submit-title {
    font-size: 18px;
    font-weight: 900;
    color: #0058be;
    margin-bottom: 6px;
}

.journal-submit-desc {
    font-size: 14px;
    color: #45474c;
    line-height: 1.45;
}

</style>
""", unsafe_allow_html=True)


st.markdown("""
<style>


/* =========================
   작업팀 접속 첫 화면
========================= */


.team-create-wrap {
    display: flex;
    justify-content: flex-end;
    margin-bottom: 90px;
}

.team-create-btn {
    border: 1px solid rgba(255,255,255,0.5);
    border-radius: 999px;
    padding: 10px 18px;
    color: white;
    font-size: 14px;
    font-weight: 800;
    background: rgba(255,255,255,0.06);
}

.team-icon {
    width: 104px;
    height: 104px;
    border-radius: 999px;
    background: rgba(255,255,255,0.12);
    display: flex;
    align-items: center;
    justify-content: center;
    margin: 0 auto 24px auto;
    font-size: 54px;
    box-shadow: 0 0 30px rgba(33,112,228,0.35);
}

.team-title {
    text-align: center;
    font-size: 34px;
    font-weight: 900;
    margin-bottom: 42px;
}

.team-helper {
    display: flex;
    justify-content: center;
    align-items: center;
    gap: 8px;
    margin: 16px 0 18px 0;
    color: rgba(255,255,255,0.72);
    font-size: 13px;
}

.team-footer {
    text-align: center;
    margin-top: 58px;
    color: rgba(255,255,255,0.4);
    font-size: 11px;
    letter-spacing: 0.16em;
    font-weight: 800;
}

.team-plant-bg {
    position: fixed;
    left: 0;
    bottom: 0;
    width: 100%;
    height: 180px;
    opacity: 0.16;
    background: linear-gradient(to top, rgba(0,0,0,0.7), rgba(0,0,0,0));
    pointer-events: none;
}


/* 팀 접속 화면 입력창 */


.team-access-input div[data-baseweb="input"] {
    background: #ffffff !important;
    border-radius: 999px !important;
    min-height: 58px;
}

.team-access-input input {
    color: #091426 !important;
    font-size: 16px !important;
    font-weight: 700 !important;
}

.team-access-input input::placeholder {
    color: #75777d !important;
}

</style>
""", unsafe_allow_html=True)

st.markdown("""
<style>

/* =========================
   팀 생성 화면
========================= */


.create-team-topbar {
    position: sticky;
    top: 4px;
    z-index: 100;
    background: #fbf8fa;
    border-bottom: 1px solid #c5c6cd;
    padding: 14px 4px 12px 4px;
    margin-bottom: 20px;
}

.create-team-topbar-row {
    display: flex;
    align-items: center;
    justify-content: space-between;
}

.create-team-left {
    display: flex;
    align-items: center;
    gap: 10px;
}

.create-team-back {
    width: 36px;
    height: 36px;
    border-radius: 999px;
    background: #f0edef;
    display: flex;
    align-items: center;
    justify-content: center;
    color: #091426;
    font-weight: 900;
}

.create-team-app-title {
    font-size: 22px;
    font-weight: 900;
    color: #091426;
}

.create-team-title {
    font-size: 30px;
    font-weight: 900;
    color: #091426;
    line-height: 1.25;
    margin-bottom: 6px;
}

.create-team-subtitle {
    font-size: 14px;
    color: #45474c;
    line-height: 1.45;
    margin-bottom: 22px;
}

.create-field-label {
    font-size: 12px;
    color: #45474c;
    font-weight: 900;
    letter-spacing: 0.05em;
    margin: 14px 0 8px 0;
}

.worker-chip {
    display: inline-flex;
    align-items: center;
    width: 100%;
    background: rgba(33,112,228,0.10);
    border: 1px solid rgba(33,112,228,0.25);
    color: #0058be;
    padding: 10px 13px;
    border-radius: 999px;
    font-size: 14px;
    font-weight: 900;
    margin-bottom: 6px;
}

.worker-list-title {
    font-size: 18px;
    font-weight: 900;
    color: #091426;
    margin-bottom: 12px;
}

.worker-chip-wrap {
    display: flex;
    flex-wrap: wrap;
    gap: 8px;
    margin-top: 12px;
}

.worker-chip {
    display: inline-flex;
    align-items: center;
    background: rgba(33,112,228,0.10);
    border: 1px solid rgba(33,112,228,0.25);
    color: #0058be;
    padding: 7px 11px;
    border-radius: 999px;
    font-size: 13px;
    font-weight: 900;
}

.create-team-info-box {
    display: flex;
    gap: 10px;
    background: rgba(216,226,255,0.55);
    border: 1px solid #adc6ff;
    color: #004395;
    padding: 14px;
    border-radius: 14px;
    font-size: 13px;
    line-height: 1.45;
    margin-bottom: 16px;
}

.input-info-box {
    background: #f5f7fb;
    border: 1px solid #d6d9e0;
    border-radius: 14px;
    padding: 14px 16px;
    margin-bottom: 16px;
    font-size: 15px;
    font-weight: 600;
    color: #31343a;
}

</style>
""", unsafe_allow_html=True)
st.markdown("""
<style>
/* 앱 상단바 (render_topbar()가 만드는 st.container(key="topbar_*")를 타겟)
   position: sticky는 Streamlit의 실제 스크롤 컨테이너(section.stMain)가
   내부적으로 중첩되어 있고 컴포넌트(검색창 등)가 마운트될 때 스스로 스크롤을
   이동시키는 경우가 있어 고정(stuck) 동작이 깨지는 것을 확인했다. 대신
   position: fixed를 쓰고, 아래 문서 흐름에는 실제 렌더링 높이보다 넉넉한
   spacer를 둬서 콘텐츠가 가려지지 않도록 한다. */
div[class*="st-key-topbar_"] {
    position: fixed !important;
    top: 60px !important;
    left: 50% !important;
    transform: translateX(-50%) !important;

    width: min(480px, calc(100% - 28px)) !important;

    z-index: 9998 !important;

    margin: 0 !important;
    box-sizing: border-box !important;

    background: #fbf8fa;
    border-bottom: 1px solid #c5c6cd;
    border-radius: 0 0 14px 14px;
    padding: 10px 10px !important;

    box-shadow: 0 4px 14px rgba(15,23,42,0.08);

    /* 내부 요소가 어떤 이유로든 자기 폭을 못 줄이더라도 상단바 박스
       밖으로 시각적으로 삐져나가지 않도록 하는 안전장치 */
    overflow: hidden !important;
}

/* 상단바가 fixed라서 본문이 가려지지 않도록 여백 추가.
   Streamlit이 topbar 컨테이너 자리에 자체적으로 예약해두는 높이가 있어
   (실측 약 128px) 이 값과 아래 padding-top(0.25rem=4px)을 함께 고려해서
   상단바 바로 아래에 본문이 붙도록 역산한 값이다. */
.fixed-topbar-spacer {
    height: 8px;
}

/* 관리자 대시보드 상단바는 기존 어두운 디자인 유지 */
div[class*="st-key-topbar_manager"] {
    background: #091426 !important;
    border-bottom: none !important;
    border-radius: 0 0 18px 18px !important;
}

/* 뒤로가기(←) / 안내(ℹ️) 버튼을 동그란 아이콘처럼 보이도록 스타일링
   (44px = 모바일 터치 권장 최소 크기) */
div[class*="st-key-tbback_"] button,
div[class*="st-key-tbhelp_"] button {
    width: 44px !important;
    height: 44px !important;
    min-width: 44px !important;
    min-height: 44px !important;
    flex-shrink: 0 !important;
    padding: 0 !important;
    border-radius: 999px !important;
    background: #f0edef !important;
    color: #091426 !important;
    font-weight: 800 !important;
    font-size: 16px !important;
    border: none !important;
}

div[class*="st-key-tbback_manager"] button,
div[class*="st-key-tbhelp_manager"] button {
    background: rgba(255,255,255,0.14) !important;
    color: white !important;
}

/* Streamlit의 st.columns는 기본적으로 좁은 화면(약 640px 이하)에서
   세로로 쌓이도록 되어 있다. 상단바(←/제목/ℹ️)는 항상 한 줄로 유지되어야
   하므로 이 자동 stack 동작을 명시적으로 막는다. */
div[class*="st-key-topbar_"] div[data-testid="stHorizontalBlock"] {
    flex-wrap: nowrap !important;
}

/* 컬럼 비율(예: [1, 6, 1])은 Streamlit이 inline flex-basis로 강제하기 때문에
   단순 min-width:0/flex-shrink 재정의만으로는 좁은 화면에서 제목이 안내 버튼을
   밀어내는 문제가 해결되지 않는다. 버튼이 들어있는 컬럼(뒤로가기/안내)은
   내용 크기만큼만 차지하고, 버튼이 없는 컬럼(제목)만 남은 공간을 갖도록
   구조적으로 구분해서 강제한다. */
div[class*="st-key-topbar_"] [data-testid="stColumn"]:has(button) {
    flex: 0 0 auto !important;
    width: auto !important;
    min-width: 0 !important;
}

div[class*="st-key-topbar_"] [data-testid="stColumn"]:not(:has(button)) {
    flex: 1 1 0% !important;
    width: auto !important;
    min-width: 0 !important;
}

div[class*="st-key-topbar_"] [data-testid="stColumn"]:not(:has(button)) * {
    min-width: 0 !important;
}

/* overflow:hidden을 제목 텍스트 바로 위 요소에 걸면 한글 글자(받침)가
   세로로 잘려 다른 글자처럼 보이는 렌더링 문제가 있어, 여기서는 줄바꿈만
   막고(nowrap) 실제 클리핑은 상단바 컨테이너 바깥쪽 overflow:hidden에 맡긴다. */
div[class*="st-key-topbar_"] [data-testid="stMarkdownContainer"] > div {
    white-space: nowrap;
}

/* 모바일에서 살짝 조정 */
@media (max-width: 640px) {
    div[class*="st-key-topbar_"] {
        width: 100% !important;
        border-radius: 0 !important;
    }
}

/* 입력창/텍스트 옆에 작은 버튼이 나란히 오는 여러 화면(작업자 명단 등록의
   "추가"/"X" 삭제, 저장된 템플릿의 "불러오기"/"삭제", TBM 이력의 다운로드
   아이콘 등)에서, 좁은 화면에서 Streamlit이 컬럼을 자동으로 세로로 쌓아버리는
   기본 동작을 막는다. 이런 행은 모두 key를 "inline_row_"로 시작하게 지어
   이 셀렉터 하나로 한꺼번에 잡는다(TBM 이력 목록은 기존 tbm_history_list
   컨테이너를 그대로 재사용). 버튼이 든 마지막 컬럼은 내용 크기만큼만 차지하게
   고정하고, 앞쪽 컬럼(입력창/칩/텍스트)이 남는 공간에 맞춰 줄어들도록 한다. */
div[class*="st-key-inline_row_"] div[data-testid="stHorizontalBlock"],
div[class*="st-key-tbm_history_list"] div[data-testid="stHorizontalBlock"] {
    flex-wrap: nowrap !important;
}

/* 저장된 템플릿의 "불러오기"/"삭제"처럼 버튼 두 개가 동일한 비율로 나란히
   와야 하는 행은, 마지막 컬럼을 고정폭으로 만들면 두 버튼 크기가 서로
   달라져 버리므로 줄바꿈만 막고 컬럼 비율은 그대로 둔다. */
div[class*="st-key-inline_pair_"] div[data-testid="stHorizontalBlock"] {
    flex-wrap: nowrap !important;
}

div[class*="st-key-inline_row_"] [data-testid="stColumn"]:last-child,
div[class*="st-key-tbm_history_list"] [data-testid="stColumn"]:last-child {
    flex: 0 0 auto !important;
    width: auto !important;
    min-width: 0 !important;
}
</style>
""", unsafe_allow_html=True)
# ========= 화면 구현 CCS =====

if "page" not in st.session_state:
    st.session_state.page = "team_access"

query_params = st.query_params

if query_params.get("page"):
    st.session_state.page = query_params.get("page")
    # 이 값을 URL에 남겨두면 다음 rerun마다(버튼 클릭 등 세션 내부 이동에도) 이 값으로
    # session_state.page가 계속 덮어써지므로, 한 번 반영한 뒤에는 즉시 지운다.
    del st.query_params["page"]

if "mode" not in st.session_state:
    st.session_state.mode = "작업자"

if "work_data" not in st.session_state:
    st.session_state.work_data = {}

if "team_name" not in st.session_state:
    st.session_state.team_name = ""

if "team_password" not in st.session_state:
    st.session_state.team_password = ""

if "team_id" not in st.session_state:
    st.session_state.team_id = ""

if "created_teams" not in st.session_state:
    st.session_state.created_teams = {}

if "temp_workers" not in st.session_state:
    st.session_state.temp_workers = []

# =========================
# 안내 이미지 팝업 설정
# =========================
HELP_INFO_DIR = Path(__file__).parent / "help_info"

HELP_IMAGES = {
    "login": ["mode_choice"],
    "create_team": ["create_team"],
    "input": ["input_work"],
    "result": ["risk_1", "risk_2"],
    "checklist": ["checklist"],
    "journal": ["journal"],
    "manager": ["manager_1", "manager_2", "manager_3"],
}

def get_help_image_path(image_name):
    """
    확장자를 직접 쓰지 않아도 png, jpg, jpeg, webp 순서로 찾아줌
    """
    for ext in ["png", "jpg", "jpeg", "webp"]:
        path = HELP_INFO_DIR / f"{image_name}.{ext}"
        if path.exists():
            return path
    return None


def _close_help_popup():
    st.session_state.active_help_page = None


def _image_to_data_uri(path):
    ext = path.suffix.lower().lstrip(".")
    mime = {
        "png": "image/png",
        "jpg": "image/jpeg",
        "jpeg": "image/jpeg",
        "webp": "image/webp",
    }.get(ext, "image/png")
    data = base64.b64encode(path.read_bytes()).decode("ascii")
    return f"data:{mime};base64,{data}"


def render_help_carousel(image_list):
    """
    이전/다음 버튼 대신 좌우 스와이프(모바일 터치)/드래그(데스크톱 마우스)로
    넘기는 캐러셀. 순수 Streamlit 위젯으로는 터치 제스처를 잡을 수 없어
    st.iframe으로 HTML/CSS/JS를 직접 삽입한다
    (components.html은 이 Streamlit 버전에서 사실상 동작하지 않아 st.iframe 사용).
    """
    paths = []
    for name in image_list:
        path = get_help_image_path(name)
        if path is None:
            st.error(f"안내 이미지를 찾을 수 없습니다: {name}")
            st.caption(f"확인 경로: {HELP_INFO_DIR}")
            continue
        paths.append(path)

    if not paths:
        return

    # 이미지들 중 가장 세로로 긴 비율에 맞춰 슬라이드 높이를 잡는다.
    # (object-fit: contain으로 넣으므로 비율이 다른 이미지도 잘리지 않는다)
    max_ratio = 0.0
    for path in paths:
        with Image.open(path) as im:
            w, h = im.size
            if w > 0:
                max_ratio = max(max_ratio, h / w)

    slide_width = 620  # st.dialog(width="large") 내부 실사용 폭 근사값
    slide_height = int(slide_width * max_ratio) if max_ratio > 0 else 420
    slide_height = max(240, min(slide_height, 620))

    slides_html = "".join(
        f'<div class="hc-slide"><img src="{_image_to_data_uri(path)}" /></div>'
        for path in paths
    )
    dots_html = "".join(
        f'<span class="hc-dot{" active" if i == 0 else ""}"></span>'
        for i in range(len(paths))
    )

    html = f"""
<style>
  html, body {{ margin:0; padding:0; }}
  .hc-wrap {{ font-family: -apple-system, "Malgun Gothic", sans-serif; }}
  .hc-track {{
      display: flex;
      overflow-x: auto;
      scroll-snap-type: x mandatory;
      -webkit-overflow-scrolling: touch;
      scrollbar-width: none;
      cursor: grab;
      user-select: none;
      border-radius: 12px;
      background: #f1f5f9;
  }}
  .hc-track::-webkit-scrollbar {{ display: none; }}
  .hc-track.dragging {{ cursor: grabbing; scroll-snap-type: none; }}
  .hc-slide {{
      flex: 0 0 100%;
      scroll-snap-align: center;
      display: flex;
      align-items: center;
      justify-content: center;
      height: {slide_height}px;
  }}
  .hc-slide img {{
      max-width: 100%;
      max-height: 100%;
      object-fit: contain;
      pointer-events: none;
      border-radius: 12px;
  }}
  .hc-dots {{
      display: flex;
      justify-content: center;
      gap: 8px;
      margin-top: 12px;
  }}
  .hc-dot {{
      width: 8px;
      height: 8px;
      border-radius: 999px;
      background: #cbd2dc;
      transition: background 0.15s, transform 0.15s;
  }}
  .hc-dot.active {{
      background: #2170e4;
      transform: scale(1.25);
  }}
</style>
<div class="hc-wrap">
  <div class="hc-track" id="hcTrack">
    {slides_html}
  </div>
  <div class="hc-dots" id="hcDots">
    {dots_html}
  </div>
</div>
<script>
(function() {{
  const track = document.getElementById("hcTrack");
  const dots = document.querySelectorAll("#hcDots .hc-dot");
  const slideCount = {len(paths)};

  function setActive(idx) {{
      dots.forEach((d, i) => d.classList.toggle("active", i === idx));
  }}

  function currentIndex() {{
      const w = track.clientWidth || 1;
      return Math.round(track.scrollLeft / w);
  }}

  let scrollTimer = null;
  track.addEventListener("scroll", function() {{
      if (scrollTimer) clearTimeout(scrollTimer);
      scrollTimer = setTimeout(function() {{
          setActive(Math.max(0, Math.min(slideCount - 1, currentIndex())));
      }}, 60);
  }});

  dots.forEach((dot, i) => {{
      dot.addEventListener("click", function() {{
          track.scrollTo({{ left: i * track.clientWidth, behavior: "smooth" }});
      }});
  }});

  // 데스크톱: 마우스로 드래그해도 넘어가도록.
  // 컨테이너 폭의 50%를 기준으로 삼으면(currentIndex()) 데스크톱처럼 넓은
  // 화면에서는 수백 px를 끌어야 다음 장으로 넘어가 버려서, 폭과 무관하게
  // 일정 픽셀(60px)만 끌면 한 장 넘어가도록 고정 임계값을 쓴다.
  let isDown = false;
  let startX = 0;
  let lastDx = 0;
  let baseIndex = 0;
  const DRAG_THRESHOLD = 60;

  track.addEventListener("mousedown", function(e) {{
      isDown = true;
      track.classList.add("dragging");
      startX = e.pageX;
      lastDx = 0;
      baseIndex = Math.max(0, Math.min(slideCount - 1, currentIndex()));
  }});

  window.addEventListener("mouseup", function() {{
      if (!isDown) return;
      isDown = false;
      track.classList.remove("dragging");
      let idx = baseIndex;
      if (lastDx <= -DRAG_THRESHOLD) idx = baseIndex + 1;
      else if (lastDx >= DRAG_THRESHOLD) idx = baseIndex - 1;
      idx = Math.max(0, Math.min(slideCount - 1, idx));
      track.scrollTo({{ left: idx * track.clientWidth, behavior: "smooth" }});
  }});

  window.addEventListener("mousemove", function(e) {{
      if (!isDown) return;
      e.preventDefault();
      lastDx = e.pageX - startX;
      track.scrollLeft = baseIndex * track.clientWidth - lastDx;
  }});
}})();
</script>
"""
    st.iframe(html, height=slide_height + 40)


@st.dialog("화면 안내", width="large", on_dismiss=_close_help_popup)
def show_help_popup(page_key):
    image_list = HELP_IMAGES.get(page_key, [])

    if not image_list:
        st.warning("이 화면에 등록된 안내 이미지가 없습니다.")
        return

    render_help_carousel(image_list)


def show_active_help_popup():
    active_page = st.session_state.get("active_help_page")

    if active_page in HELP_IMAGES:
        show_help_popup(active_page)

def render_topbar_spacer():
    st.markdown('<div class="fixed-topbar-spacer"></div>', unsafe_allow_html=True)

def render_topbar(topbar_key, title, title_class, back_page=None, help_key=None):
    """
    뒤로가기/안내 아이콘이 있는 상단바.
    <a href="?page=..."> 링크는 브라우저 풀 리로드를 유발해 st.session_state(team_id 등)가
    통째로 초기화되므로, st.button + st.rerun()으로 같은 세션 안에서 화면만 전환한다.
    """
    with st.container(key=f"topbar_{topbar_key}"):
        if back_page:
            col_back, col_title, col_info = st.columns([1, 6, 1], vertical_alignment="center")
        else:
            col_back = None
            col_title, col_info = st.columns([6, 1], vertical_alignment="center")

        if col_back is not None:
            with col_back:
                if st.button("←", key=f"tbback_{topbar_key}"):
                    st.session_state.page = back_page
                    st.rerun()

        with col_title:
            st.markdown(f'<div class="{title_class}">{title}</div>', unsafe_allow_html=True)

        with col_info:
            if help_key:
                if st.button("ℹ️", key=f"tbhelp_{topbar_key}"):
                    st.session_state.active_help_page = help_key
                    st.rerun()

    render_topbar_spacer()

def _redirect_with_message(target_page, message):
    """앞 단계를 건너뛰고 화면에 진입했을 때, 원래 있어야 할 화면으로 돌려보내고
    이유를 안내한다. 하단 네비게이션 버튼 클릭뿐 아니라 URL의 ?page= 쿼리파라미터로
    직접 진입하는 경우(브라우저 뒤로가기가 사실상 히스토리 없이 동작하는 대신 남는
    유일한 우회 경로)도 이 가드로 함께 막는다."""
    st.session_state.pending_redirect_message = message
    st.session_state.page = target_page
    st.rerun()


def _flash_pending_message():
    """직전 rerun에서 _redirect_with_message가 남긴 안내 메시지를 한 번만 보여준다."""
    message = st.session_state.pop("pending_redirect_message", None)
    if message:
        st.warning(message)


def show_team_access():

    # Material Icons
    st.markdown("""
<link rel="stylesheet" href="https://fonts.googleapis.com/css2?family=Material+Symbols+Rounded" />
""", unsafe_allow_html=True)

    # CSS
    st.markdown("""
<style>

.stApp {
    background: linear-gradient(180deg, #091426 0%, #0b2f7a 100%);
}

.block-container {
    max-width: 420px;
    padding-top: 1.5rem;
    padding-bottom: 2rem;
}

/* 상단 팀생성 버튼 */
.team-create-wrap {
    display: flex;
    justify-content: flex-end;
    margin-bottom: 50px;
}

.team-create-btn {
    background: #f7f7f7;
    color: #1b1b1d;
    border: 1px solid #d7d9e0;
    border-radius: 999px;
    padding: 12px 20px;
    font-weight: 900;
    font-size: 14px;
    box-shadow: 0 6px 16px rgba(0,0,0,0.22);
    display: inline-block;
}

/* 아이콘 */
.team-icon {
    width: 108px;
    height: 108px;
    border-radius: 999px;
    background: rgba(255,255,255,0.14);
    display: flex;
    align-items: center;
    justify-content: center;
    margin: 0 auto 26px auto;
    box-shadow: 0 0 34px rgba(33,112,228,0.35);
}

.material-symbols-rounded {
    font-variation-settings:
    'FILL' 1,
    'wght' 500,
    'GRAD' 0,
    'opsz' 48;
}

.team-helmet-icon {
    font-size: 56px;
    color: white;
}

/* 제목 */
.team-title {
    text-align: center;
    color: white;
    font-size: 34px;
    font-weight: 900;
    margin-bottom: 42px;
}

/* 입력창 */
div[data-baseweb="input"] {
    background: white !important;
    border-radius: 999px !important;
    min-height: 58px !important;
}

div[data-baseweb="input"] input {
    height: 58px !important;
    color: #091426 !important;
    font-weight: 700 !important;
    font-size: 16px !important;
}

/* 버튼: 앱 전체 버튼 스타일(흰 배경 + 테두리)과 통일 */
div.stButton > button {
    border-radius: 999px;
    font-size: 17px;
    font-weight: 900;
    background: #f7f7f7;
    color: #1b1b1d;
    border: 1px solid #d7d9e0;
    box-shadow: 0 6px 16px rgba(0,0,0,0.22);
    margin-top: 8px;
}

div.stButton > button:hover {
    background: #f7f8fa;
    border-color: #b9bcc6;
    color: #1b1b1d;
}

/* 안내문구 */
.team-helper {
    text-align: center;
    color: rgba(255,255,255,0.72);
    font-size: 13px;
    margin: 16px 0 20px 0;
    line-height: 1.5;
}

/* 하단 */
.team-footer {
    text-align: center;
    color: rgba(255,255,255,0.4);
    font-size: 11px;
    letter-spacing: 0.16em;
    font-weight: 800;
    margin-top: 48px;
}

</style>
""", unsafe_allow_html=True)

    _flash_pending_message()

    # 상단 팀 생성 버튼
    st.markdown("""
<div class="team-create-wrap">
    <a href="?page=create_team" target="_self" style="text-decoration:none;">
        <div class="team-create-btn">+ 팀 생성</div>
    </a>
</div>
""", unsafe_allow_html=True)

    # 아이콘 + 제목
    st.markdown("""
<div class="team-icon">
    <span class="material-symbols-rounded team-helmet-icon">
        engineering
    </span>
</div>

<div class="team-title">작업팀 접속</div>
""", unsafe_allow_html=True)

    # 입력창
    team_name = st.text_input(
        "팀명",
        placeholder="팀명을 입력하세요",
        label_visibility="collapsed",
        key="team_name_input"
    )

    team_password = st.text_input(
        "팀 비밀번호",
        placeholder="비밀번호를 입력하세요",
        type="password",
        label_visibility="collapsed",
        key="team_password_input"
    )

    # 안내문구
    st.markdown("""
<div class="team-helper">
🛡️ 안전관리자가 생성한 팀정보로 접속합니다.
</div>
""", unsafe_allow_html=True)

    # 접속 버튼
    if st.button(
        "TBM 접속하기",
        key="team_access_submit",
        use_container_width=True
    ):
        if not team_name.strip():
            st.warning("팀명을 입력해 주세요.")

        elif not team_password.strip():
            st.warning("비밀번호를 입력해 주세요.")

        else:
            team = login_team(
                team_name.strip(),
                team_password.strip()
            )
            if not team:
                st.error("팀명 또는 비밀번호가 일치하지 않습니다.")
                return

            st.session_state.team_name = team["team_name"]
            st.session_state.team_password = team_password.strip()
            st.session_state.team_id = team["id"]            

            st.query_params.clear()

            st.session_state.page = "login"
            st.rerun()

    # 하단 문구
    st.markdown("""
<div class="team-footer">
SAFETY TBM SYSTEM
</div>
""", unsafe_allow_html=True)

def _populate_worker_session_from_task(worker_name, selected_task, is_tbm_leader, leader_department, leader_position):
    """선택된 작업(task) 정보를 work_data/result 세션에 채운다.
    "TBM 시작"과 "작업로그 작성하기" 양쪽에서 같은 작업/작업자 정보를 동일하게 채우기 위해 공유한다."""
    st.session_state.work_data["작업자명"] = worker_name
    st.session_state.work_data["작업명"] = selected_task.get("work_name", "")
    st.session_state.work_data["작업내용"] = selected_task.get("work_content", "")
    st.session_state.work_data["작업장소"] = selected_task.get("work_location", "")
    st.session_state.work_data["TBM장소"] = selected_task.get("tbm_place", "")
    st.session_state.work_data["예정시간"] = selected_task.get("scheduled_time", "")
    st.session_state.work_data["task_id"] = selected_task.get("id", "")

    st.session_state.work_data["TBM리더여부"] = is_tbm_leader
    st.session_state.work_data["리더소속"] = leader_department.strip()
    st.session_state.work_data["리더직책"] = leader_position.strip()
    st.session_state.work_data["리더성명"] = worker_name if is_tbm_leader else ""

    # 안전관리자가 "오늘 작업 입력" 화면에서 미리 분석해 둔 위험도 결과를
    # 그대로 가져와 보여준다 (작업자가 다시 분석하지 않음).
    st.session_state.result = {
        "score": selected_task.get("risk_score"),
        "level": selected_task.get("risk_level"),
        "chem_name": selected_task.get("chem_name"),
        "chem_id": selected_task.get("chem_id"),
        "작업유형": selected_task.get("work_type_display"),
        "작업유형_모델값": selected_task.get("work_type_code"),
        "취급물질": selected_task.get("material_name"),
        "작업시간": selected_task.get("start_time"),
        "작업시간대": selected_task.get("time_slot"),
        "similar_accident_text": selected_task.get("similar_accident_text", ""),
        "task_id": selected_task.get("id", ""),
    }
    st.session_state.work_data["작업유형"] = selected_task.get("work_type_display", "")
    st.session_state.work_data["취급물질"] = selected_task.get("material_name", "")
    st.session_state.work_data["main_hazard_1"] = selected_task.get("main_hazard_1", "")
    st.session_state.work_data["main_hazard_2"] = selected_task.get("main_hazard_2", "")
    st.session_state.work_data["main_hazard_3"] = selected_task.get("main_hazard_3", "")
    st.session_state.work_data["safety_measure_1"] = selected_task.get("safety_measure_1", "")
    st.session_state.work_data["safety_measure_2"] = selected_task.get("safety_measure_2", "")
    st.session_state.work_data["safety_measure_3"] = selected_task.get("safety_measure_3", "")


def show_login():

    if not st.session_state.get("team_id"):
        st.warning("팀 접속 정보가 없습니다. 작업팀 접속 화면에서 다시 접속해 주세요.")

        if st.button("작업팀 접속 화면으로 이동", use_container_width=True):
            st.session_state.page = "team_access"
            st.rerun()

        return

    # =========================
    # 상단바
    # =========================
    render_topbar("login", "작업 모드 선택", "create-team-app-title", back_page="team_access", help_key="login")
    _flash_pending_message()

    # 접속 기기의 로컬 날짜. "오늘 작업" 목록을 서버(UTC) 날짜로 조회하면
    # 한국 시간 00시~09시 사이에는 하루 어긋난 목록이 나오므로 이걸 기준으로 삼는다.
    login_client_dt = get_client_datetime()

    # =========================
    # 로그인 헤더
    # =========================
    login_header_html = '<div class="login-header"><div class="login-logo-row"><div class="login-logo-icon">✓</div><div class="login-title">Safety TBM</div></div><div class="login-subtitle">안전한 작업의 시작, 스마트 안전관리 플랫폼</div></div>'
    st.markdown(login_header_html, unsafe_allow_html=True)


    # =========================
    # 모드 선택
    # =========================
    mode = st.radio(
        "접속 모드",
        ["작업자", "안전관리자"],
        horizontal=True,
        label_visibility="collapsed",
        key="login_mode_radio"
    )

    # =========================
    # 모드 카드 UI
    # =========================
    if mode == "작업자":
        mode_card_html = '<div class="mode-card-wrap"><div class="mode-card mode-card-active"><div class="mode-icon">👷</div><div class="mode-label">작업자 모드</div></div><div class="mode-card"><div class="mode-icon">🛡️</div><div class="mode-label">안전관리자 모드</div></div></div>'
    else:
        mode_card_html = '<div class="mode-card-wrap"><div class="mode-card"><div class="mode-icon">👷</div><div class="mode-label">작업자 모드</div></div><div class="mode-card mode-card-active"><div class="mode-icon">🛡️</div><div class="mode-label">안전관리자 모드</div></div></div>'

    st.markdown(mode_card_html, unsafe_allow_html=True)

    # =========================
    # 입력 영역
    # =========================
    if mode == "작업자":
        worker_name = st.text_input(
            "작업자명",
            placeholder="성명을 입력하세요",
            key="login_worker_name"
        )
        
        today_tasks = get_today_work_tasks(
            st.session_state.get("team_id", ""),
            client_date=(login_client_dt or datetime.now()).strftime("%Y-%m-%d")
        )
        selected_task = None
        if today_tasks:
            task_options = {
                f'{task.get("work_name", "-")} / {task.get("scheduled_time", "-")}': task
                for task in today_tasks
        }
        
            selected_task_label = st.selectbox(
                "작업명 선택",
                list(task_options.keys()),
                key="login_selected_task"
            )
            selected_task = task_options[selected_task_label]
        
        else:
            st.warning("안전관리자가 등록한 오늘 작업이 없습니다.")
        
        

        is_tbm_leader = st.checkbox(
            "TBM 리더입니다",
            key="is_tbm_leader"
        )

        leader_department = ""
        leader_position = ""
        

        if is_tbm_leader:

            leader_department = st.text_input(
                "소속",
                placeholder="예: 생산1팀",
                key="leader_department"
            )

            leader_position = st.text_input(
                "직책",
                placeholder="예: 반장",
                key="leader_position"
            )


        manager_password_input = ""

    else:
        manager_password_input = st.text_input(
            "안전관리자 모드 진입 비밀번호",
            placeholder="팀 생성 시 등록한 안전관리자 비밀번호 입력",
            type="password",
            key="login_manager_password"
        )

        worker_name = ""
        work_name = "안전관리자 대시보드"

    # =========================
    # 안내 문구
    # =========================
    info_box_html = '<div class="info-box"><div class="info-icon">ℹ</div><div class="info-text">본 시스템은 산업안전보건법에 따른 <b>TBM(Tool Box Meeting)</b>의 디지털 기록을 위해 사용됩니다. 정확한 정보를 입력해 주시기 바랍니다.</div></div>'
    st.markdown(info_box_html, unsafe_allow_html=True)

    # =========================
    # 시작 버튼
    # =========================
    if st.button("TBM 시작  ▶", key="login_start_btn", use_container_width=True):
        st.session_state.work_data["접속시간"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")


        if mode == "작업자":

            if not worker_name.strip():
                st.warning("작업자명을 입력해 주세요.")
                return

            if selected_task is None:
                st.warning("등록된 작업을 선택해 주세요.")
                return

            st.session_state.mode = mode
            _populate_worker_session_from_task(
                worker_name.strip(), selected_task, is_tbm_leader, leader_department, leader_position
            )
            st.session_state.work_data["접속모드"] = mode
            st.session_state.work_data["접속시간"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

            if st.session_state.work_data.get("TBM리더여부", False):
                st.session_state.work_data["TBM시작시간"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

            st.query_params.clear()

            st.session_state.page = "task_info"
            st.rerun()

        else:

            if not manager_password_input.strip():
                st.warning("안전관리자 비밀번호를 입력해 주세요.")
                return

            result = (
                supabase.table("teams")
                .select("*")
                .eq("id", st.session_state.team_id)
                .execute()
            )

            if not result.data:
                st.error("팀 정보를 찾을 수 없습니다.")
                return

            team_info = result.data[0]
            saved_manager_password = team_info.get("manager_password", "")

            if manager_password_input.strip() != saved_manager_password:
                st.error("안전관리자 비밀번호가 일치하지 않습니다.")
                return

            st.session_state.mode = mode
            st.session_state.manager_name = team_info.get("manager_name", "안전관리자")

            st.session_state.work_data["작업자명"] = st.session_state.get("manager_name", "안전관리자")
            st.session_state.work_data["작업명"] = "안전관리자 대시보드"
            st.session_state.work_data["접속모드"] = mode
            st.session_state.work_data["접속시간"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

            st.query_params.clear()

            st.session_state.page = "manager"
            st.rerun()

    # =========================
    # 작업로그 작성하기 (TBM 체크리스트를 이미 완료한 작업/작업자에 대해
    # 작업일지만 이어서 작성)
    # =========================
    if mode == "작업자":
        if st.button("📝 작업로그 작성하기", key="login_journal_btn", use_container_width=True):

            if not worker_name.strip():
                st.warning("작업자명을 입력해 주세요.")
                return

            if selected_task is None:
                st.warning("등록된 작업을 선택해 주세요.")
                return

            journal_task_id = selected_task.get("id", "")

            try:
                existing_check = (
                    supabase.table("work_logs")
                    .select("id")
                    .eq("task_id", journal_task_id)
                    .eq("worker_name", worker_name.strip())
                    .limit(1)
                    .execute()
                )
            except Exception as e:
                st.error("작업 기록 조회 중 오류가 발생했습니다.")
                st.write(str(e))
                return

            if not existing_check.data:
                st.warning("먼저 TBM 체크리스트를 완료해야 작업로그를 작성할 수 있습니다.")
                return

            st.session_state.mode = mode
            _populate_worker_session_from_task(
                worker_name.strip(), selected_task, is_tbm_leader, leader_department, leader_position
            )
            st.session_state.work_data["접속모드"] = mode

            st.query_params.clear()

            st.session_state.page = "journal"
            st.rerun()

    show_bottom_nav()
    # =========================
    # 푸터
    # =========================
    footer_html = '<div class="login-footer">이용약관 · 개인정보처리방침<br>© 2026 Korea Environment Corporation. All Rights Reserved.</div>'
    st.markdown(footer_html, unsafe_allow_html=True)

def show_create_team():
    render_topbar("create_team", "팀 생성", "create-team-app-title", back_page="team_access", help_key="create_team")

    st.markdown("""
<div class="create-team-title">새 TBM 작업방 만들기</div>
<div class="create-team-subtitle">
    안전관리자가 팀명, 비밀번호, 작업자 명단을 등록합니다.
</div>
""", unsafe_allow_html=True)

    st.markdown('<div class="create-field-label">팀명</div>', unsafe_allow_html=True)
    team_name = st.text_input(
        "팀명",
        placeholder="예: OO공장 정비1팀",
        label_visibility="collapsed",
        key="create_team_name"
    )

    st.markdown('<div class="create-field-label">팀 비밀번호</div>', unsafe_allow_html=True)
    team_password = st.text_input(
        "팀 비밀번호",
        placeholder="작업자에게 공유할 비밀번호 입력",
        type="password",
        label_visibility="collapsed",
        key="create_team_password"
    )

    st.markdown('<div class="create-field-label">안전관리자 비밀번호</div>', unsafe_allow_html=True)
    manager_password = st.text_input(
        "안전관리자 비밀번호",
        placeholder="안전관리자 모드 진입 비밀번호",
        type="password",
        label_visibility="collapsed",
        key="create_manager_password"
    )

    st.markdown('<div class="create-field-label">안전관리자 비밀번호 확인</div>', unsafe_allow_html=True)
    manager_password_confirm = st.text_input(
        "안전관리자 비밀번호 확인",
        placeholder="비밀번호를 한 번 더 입력하세요",
        type="password",
        label_visibility="collapsed",
        key="create_manager_password_confirm"
    )

    st.markdown('<div class="create-field-label">안전관리자명</div>', unsafe_allow_html=True)
    manager_name = st.text_input(
        "안전관리자명",
        placeholder="예: 홍길동",
        label_visibility="collapsed",
        key="create_manager_name"
    )

    st.markdown("""
<div class="worker-list-title" style="margin-top:24px;">작업자 명단 등록</div>
""", unsafe_allow_html=True)

    with st.container(key="inline_row_create_team_add"):
        col_worker, col_add = st.columns([4, 1])

        with col_worker:
            worker_name = st.text_input(
                "작업자 이름 입력",
                placeholder="작업자 이름 입력",
                label_visibility="collapsed",
                key="worker_name_input"
            )

        with col_add:
            if st.button("추가", key="add_worker_btn", use_container_width=True):
                if not worker_name.strip():
                    st.warning("작업자 이름을 입력해 주세요.")
                elif worker_name.strip() in st.session_state.temp_workers:
                    st.warning("이미 등록된 작업자입니다.")
                else:
                    st.session_state.temp_workers.append(worker_name.strip())
                    st.rerun()

    if st.session_state.temp_workers:
        for idx, worker in enumerate(st.session_state.temp_workers):
            with st.container(key=f"inline_row_create_team_worker_{idx}"):
                col_name, col_del = st.columns([5, 1])

                with col_name:
                    st.markdown(
                        f'<div class="worker-chip">{worker}</div>',
                        unsafe_allow_html=True
                    )

                with col_del:
                    if st.button("X", key=f"delete_worker_{idx}"):
                        st.session_state.temp_workers.remove(worker)
                        st.rerun()
    else:
        st.caption("아직 등록된 작업자가 없습니다.")

    st.markdown("""
<div class="create-team-info-box">
ℹ️ 등록된 작업자만 해당 팀의 작업자로 선택할 수 있습니다.
</div>
""", unsafe_allow_html=True)

    if st.button("작업팀 생성하기", key="create_team_submit", use_container_width=True):
        if not team_name.strip():
            st.warning("팀명을 입력해 주세요.")

        elif not team_password.strip():
            st.warning("팀 비밀번호를 입력해 주세요.")

        elif not manager_password.strip():
            st.warning("안전관리자 비밀번호를 입력해 주세요.")

        elif not manager_password_confirm.strip():
            st.warning("안전관리자 비밀번호 확인을 입력해 주세요.")

        elif manager_password.strip() != manager_password_confirm.strip():
            st.error("안전관리자 비밀번호와 비밀번호 확인이 일치하지 않습니다.")

        elif not manager_name.strip():
            st.warning("안전관리자명을 입력해 주세요.")

        elif not st.session_state.temp_workers:
            st.warning("작업자 명단을 1명 이상 등록해 주세요.")

        else:
            try:
                saved_team = create_team(
                    team_name=team_name.strip(),
                    team_password=team_password.strip(),
                    manager_name=manager_name.strip(),
                    manager_password=manager_password.strip(),
                    workers=st.session_state.temp_workers.copy()
                )

                if not saved_team:
                    st.error("DB 저장 결과가 비어 있습니다.")
                    return

                st.session_state.team_name = team_name.strip()
                st.session_state.team_password = team_password.strip()
                st.session_state.team_id = saved_team[0]["id"]
                st.session_state.manager_name = manager_name.strip()
                st.session_state.manager_password = manager_password.strip()
                st.session_state.temp_workers = []

                st.success("작업팀이 생성되었습니다.")
                st.session_state.page = "team_access"
                st.rerun()

            except Exception as e:
                if "duplicate key value" in str(e):
                    st.error("이미 존재하는 팀명입니다. 다른 팀명을 입력해 주세요.")
                else:
                    st.error("Supabase 저장 중 오류가 발생했습니다.")
                    st.write(str(e))
                return

def run_risk_scoring(chemical, work_type, time_slot):
    """
    화학물질명/작업유형코드/시간대를 받아 위험도 점수를 계산한다.
    작업자 모드(show_work_input)와 안전관리자 모드(show_task_create)가 공유해서 쓴다.
    실패 시 (None, 에러메시지)를 반환한다.
    """
    chem_id, chem_name, err = get_chemid_by_name(
        chemical,
        SERVICE_KEY
    )

    if err:
        return None, err

    status_code, detail_text = get_hazard_by_chemid(
        chem_id,
        SERVICE_KEY
    )

    if status_code != 200:
        return None, f"상세 위해성 API 호출 실패: {status_code}"

    classification_text = extract_classification_text(detail_text)

    hazard_scores = map_hazard_scores_by_excel(
        classification_text,
        mapping_df
    )

    input_df = make_input_data(work_type, time_slot)
    pred_prob = model.predict_proba(input_df)[0][1]

    score, level, detail = calculate_final_score(
        input_df=input_df,
        pred_prob=pred_prob,
        work_type=work_type,
        time_slot=time_slot,
        chem_info_missing=0
    )

    result = {
        "score": score,
        "level": level,
        "detail": detail,
        "input_df": input_df,
        "chem_name": chem_name,
        "chem_id": chem_id,
        "hazard_scores": hazard_scores,
        "classification_text": classification_text,
    }

    return result, None


def show_work_input():

    if not st.session_state.get("work_data", {}).get("task_id"):
        _redirect_with_message("login", "먼저 작업자명과 오늘 작업을 선택해 주세요.")
        return

    # =========================
    # 상단바
    # =========================
    render_topbar("input", "Safety TBM", "app-title", back_page="login", help_key="input")
    _flash_pending_message()

    st.markdown("""
    <div>
        <div class="hero-title">오늘의 작업 정보를 확인하세요.</div>
        <div class="hero-subtitle">안전관리자가 등록한 작업 정보를 기준으로 TBM을 진행합니다.</div>
    </div>
    """, unsafe_allow_html=True)

    # =========================
    # 선택 작업 정보 자동 표시
    # =========================
    work_name = st.session_state.work_data.get("작업명", "-")
    work_content = st.session_state.work_data.get("작업내용", "-")
    work_location = st.session_state.work_data.get("작업장소", "-")
    scheduled_time = st.session_state.work_data.get("예정시간", "-")
    tbm_place = st.session_state.work_data.get("TBM장소", "-")

    st.markdown(f"""
<div class="result-card">
    <div class="result-card-header">
        <div class="result-card-icon">📋</div>
        <div class="result-card-title">선택한 작업 정보</div>
    </div>
    <div class="message-text">
        <b>작업명</b> : {work_name}<br>
        <b>작업내용</b> : {work_content}<br>
        <b>작업장소</b> : {work_location}<br>
        <b>예정시간</b> : {scheduled_time}<br>
        <b>TBM 장소</b> : {tbm_place}
    </div>
</div>
""", unsafe_allow_html=True)

    # =========================
    # 작업 유형
    # =========================
    st.markdown('<div class="input-card">', unsafe_allow_html=True)
    st.markdown('<div class="field-label">작업 유형</div>', unsafe_allow_html=True)

    work_type_display = st.selectbox(
        "작업 유형",
        [
            "정기작업",
            "비작업(순찰·경비)",
            "유지보수",
            "화기작업",
            "시운전·정지",
            "세척작업"
        ],
        label_visibility="collapsed"
    )

    # =========================
    # 작업 시간
    # =========================
    time_slot, current_dt = get_current_time_slot(get_client_datetime())
    current_time = current_dt.strftime("%Y-%m-%d %H:%M")

    st.markdown('<div class="input-card input-card-muted">', unsafe_allow_html=True)
    st.markdown('<div class="field-label">작업 시간 자동 입력</div>', unsafe_allow_html=True)

    st.text_input(
        "작업 시간",
        value=f"{current_time} · {time_slot}",
        disabled=True,
        label_visibility="collapsed"
    )

    # =========================
    # 취급물질 입력
    # =========================
    st.markdown('<div class="input-card">', unsafe_allow_html=True)
    st.markdown('<div class="field-label">취급물질 입력</div>', unsafe_allow_html=True)

    chemical = st_searchbox(
        search_chemical_candidates,
        key="chemical_searchbox",
        placeholder="화학물질명을 입력하세요. 예: 황산",
        clear_on_submit=False,
    )

    # =========================
    # 위험도 분석 버튼
    # =========================
    if st.button("⚠️ 위험도 분석하기", use_container_width=True):

        if not work_location or work_location == "-":
            st.warning("선택한 작업의 작업장소 정보가 없습니다. 안전관리자에게 작업정보를 확인해 주세요.")
            return

        if not chemical or not str(chemical).strip():
            st.warning("취급물질을 입력하고 검색 결과에서 물질을 선택해 주세요.")
            return

        chemical = resolve_chemical_alias(str(chemical).strip())

        work_type_map = {
            "정기작업": "ROUTINE",
            "비작업(순찰·경비)": "IDLE",
            "유지보수": "MAINTENANCE",
            "화기작업": "HOT_WORK",
            "시운전·정지": "STARTUP_SHUTDOWN",
            "세척작업": "CLEANING"
        }

        work_type = work_type_map[work_type_display]

        st.session_state.work_data.update({
            "작업명": work_name,
            "작업내용": work_content,
            "작업장소": work_location,
            "예정시간": scheduled_time,
            "TBM장소": tbm_place,
            "작업유형": work_type_display,
            "작업유형_모델값": work_type,
            "취급물질": chemical,
            "작업시간": current_time,
            "작업시간대": time_slot
        })

        try:
            result, err = run_risk_scoring(chemical, work_type, time_slot)

            if err:
                st.error(err)
                return

            st.session_state.hazard_scores = result["hazard_scores"]
            st.session_state.chem_id = result["chem_id"]
            st.session_state.chem_name = result["chem_name"]
            st.session_state.classification_text = result["classification_text"]

            st.session_state.result = {
                **result,
                "작업명": work_name,
                "작업내용": work_content,
                "작업장소": work_location,
                "예정시간": scheduled_time,
                "TBM장소": tbm_place,
                "작업유형": work_type_display,
                "작업유형_모델값": work_type,
                "취급물질": chemical,
                "작업시간": current_time,
                "작업시간대": time_slot,
                "task_id": st.session_state.work_data.get("task_id", "")
            }

            st.session_state.risk_score = result["score"]
            st.session_state.page = "result"
            st.rerun()

        except Exception as e:
            st.error("위험도 분석 중 오류가 발생했습니다.")
            st.exception(e)

    st.markdown("""
    <div class="plant-image-card">
        실시간 현장 데이터 기반 위험 분석 알고리즘이 가동 중입니다.
    </div>
    """, unsafe_allow_html=True)

    show_bottom_nav()
def show_bottom_nav():
    # <a href="?page=..."> 링크는 브라우저 풀 리로드를 유발해 st.session_state(team_id 등)가
    # 통째로 초기화되므로, st.button + st.rerun()으로 같은 세션 안에서 화면만 전환한다.
    #
    # 네비게이션 대상은 실제로 지금 진행 중인 흐름(작업자: task_info→checklist→journal,
    # 안전관리자: manager→task_create)의 화면으로만 연결한다. 예전에 쓰던 수동 위험도
    # 분석 흐름(input/result)은 현재 작업 등록 기반 흐름과 맞지 않아(위험도 재계산,
    # AI 재호출 등) 더 이상 네비게이션 대상으로 노출하지 않는다.

    current_page = st.session_state.get("page", "task_info")
    current_mode = st.session_state.get("mode", "작업자")

    if current_mode == "안전관리자":
        nav_items = [
            ("manager", "▦"),
            ("task_create", "📝"),
        ]
    else:
        nav_items = [
            ("task_info", "⌂"),
            ("checklist", "☑"),
            ("journal", "✎"),
        ]

    st.markdown("""
<style>
.bottom-nav-spacer {
    height: 88px;
}

div[class*="st-key-bottomnavbar"] {
    position: fixed !important;
    left: 50% !important;
    bottom: 14px !important;
    transform: translateX(-50%) !important;
    width: min(440px, calc(100% - 28px)) !important;
    background: #fbf8fa;
    border: 1px solid #c5c6cd;
    border-radius: 22px;
    box-shadow: 0 10px 30px rgba(15,23,42,0.18);
    z-index: 9999 !important;
    padding: 6px !important;
}

/* 하단 네비게이션 아이콘이 항상 한 줄로 유지되어야 한다 */
div[class*="st-key-bottomnavbar"] div[data-testid="stHorizontalBlock"] {
    flex-wrap: nowrap !important;
}

div[class*="st-key-bottomnavbar"] [data-testid="stColumn"] {
    min-width: 0 !important;
    width: auto !important;
    flex-shrink: 1 !important;
}

div[class*="st-key-navbtn_"] button {
    border: none !important;
    background: transparent !important;
    color: #8a8f98 !important;
    font-size: 26px !important;
    font-weight: 900 !important;
    min-height: 52px !important;
    width: 100% !important;
    border-radius: 16px !important;
}

div[class*="st-key-navbtnactive_"] button {
    color: #0b3fa5 !important;
    background: rgba(11,63,165,0.08) !important;
}

@media (max-width: 640px) {
    div[class*="st-key-bottomnavbar"] {
        width: calc(100% - 22px) !important;
        bottom: 10px !important;
    }
}
</style>
<div class="bottom-nav-spacer"></div>
""", unsafe_allow_html=True)

    with st.container(key="bottomnavbar"):
        cols = st.columns(len(nav_items))
        for col, (page_name, icon) in zip(cols, nav_items):
            with col:
                key = f"navbtnactive_{page_name}" if current_page == page_name else f"navbtn_{page_name}"
                if st.button(icon, key=key):
                    st.session_state.page = page_name
                    st.rerun()

def split_db_text(value):
    """
    엑셀 셀 안에 여러 문장이 들어간 경우 보기 좋게 나누기 위한 함수.
    줄바꿈, |, ; 기준으로 분리한다.
    """
    if pd.isna(value):
        return []

    text = str(value).strip()

    if not text or text.lower() == "nan":
        return []

    for sep in ["|", ";", "\n", "\r"]:
        text = text.replace(sep, "\n")

    items = [
        x.strip()
        for x in text.split("\n")
        if x.strip()
    ]

    return items if items else [str(value).strip()]


def get_score_style(score):
    if score < 40:
        return {
            "level_text": "안전유의",
            "score_color": "#22c55e",
            "badge_bg": "#22c55e",
            "badge_color": "#111827",
            "green_class": "traffic-light green-on",
            "yellow_class": "traffic-light",
            "red_class": "traffic-light"
        }
    elif score < 70:
        return {
            "level_text": "작업주의",
            "score_color": "#facc15",
            "badge_bg": "#facc15",
            "badge_color": "#111827",
            "green_class": "traffic-light",
            "yellow_class": "traffic-light yellow-on",
            "red_class": "traffic-light"
        }
    else:
        return {
            "level_text": "위험경고",
            "score_color": "#ef4444",
            "badge_bg": "#ef4444",
            "badge_color": "#ffffff",
            "green_class": "traffic-light",
            "yellow_class": "traffic-light",
            "red_class": "traffic-light red-on"
        }


def get_risk_and_measure_messages(result):
    """
    AI 분석 주요 위험요인과 안전 및 사고 예방대책 문구를 DB에서 추출.
    """

    hazard_scores = result.get("hazard_scores", {})
    work_type_display = result.get("작업유형", "")

    material_risk_items = []
    work_risk_items = []

    material_measure_items = []
    work_measure_items = []

    # =========================
    # 1) risk_message_db: 물질군 기준 매칭
    # =========================
    active_hazards = []

    for hazard_name, hazard_score in hazard_scores.items():
        try:
            score_value = float(hazard_score)
        except Exception:
            score_value = 0

        if score_value > 0:
            active_hazards.append(str(hazard_name).strip())

    if active_hazards:
        temp_risk_df = risk_message_df.copy()
        temp_risk_df["물질군"] = temp_risk_df["물질군"].astype(str).str.strip()

        matched_risk = temp_risk_df[
            temp_risk_df["물질군"].isin(active_hazards)
        ]

        for _, row in matched_risk.iterrows():
            material_risk_items.extend(
    split_db_text(row.get("유해위험요인", ""))
)
            material_measure_items.extend(
    split_db_text(row.get("취급시 주의사항 및 예방조치", ""))
)

    # =========================
    # 2) accident_case_db: 작업명 기준 매칭
    # =========================
    temp_accident_df = accident_case_df.copy()

    if "작업명" in temp_accident_df.columns:
        work_col = "작업명"
    elif "작업유형" in temp_accident_df.columns:
        work_col = "작업유형"
    else:
        work_col = None

    if work_col:
        temp_accident_df[work_col] = temp_accident_df[work_col].astype(str).str.strip()

        matched_work = temp_accident_df[
            temp_accident_df[work_col] == str(work_type_display).strip()
        ]

        for _, row in matched_work.iterrows():
            work_risk_items.extend(
    split_db_text(row.get("주요 위험요인", ""))
)
            work_measure_items.extend(
    split_db_text(row.get("안전대책", ""))
)

    # =========================
    # 3) 90% 이상 유사 문구 중복 제거
    # =========================
    def remove_similar_text(items, similarity_threshold=0.90):
        cleaned = []
        seen_keys = []

        def normalize_for_similarity(value):
            text = str(value)

            text = (
                text
                .replace("\n", " ")
                .replace("\r", " ")
                .replace("\t", " ")
                .replace("\u00a0", " ")
                .replace("\u3000", " ")
            )

            text = " ".join(text.split())
            key = re.sub(r"\d+", "", text)

            key = (
                key
                .replace(" ", "")
                .replace(",", "")
                .replace("，", "")
                .replace(".", "")
                .replace("。", "")
                .replace("·", "")
                .replace("ㆍ", "")
                .replace("-", "")
                .replace("–", "")
                .replace("~", "")
                .replace("(", "")
                .replace(")", "")
                .replace("[", "")
                .replace("]", "")
                .replace(":", "")
                .replace(";", "")
                .strip()
            )

            return text, key

        for item in items:
            if item is None:
                continue

            try:
                if pd.isna(item):
                    continue
            except Exception:
                pass

            text, key = normalize_for_similarity(item)

            if not text or text.lower() == "nan":
                continue

            is_duplicate = False

            for old_key in seen_keys:
                similarity = SequenceMatcher(None, key, old_key).ratio()

                if similarity >= similarity_threshold:
                    is_duplicate = True
                    break

            if not is_duplicate:
                seen_keys.append(key)
                cleaned.append(text)

        return cleaned

    material_risk_items = remove_similar_text(material_risk_items, similarity_threshold=0.90)
    work_risk_items = remove_similar_text(work_risk_items, similarity_threshold=0.90)
    material_measure_items = remove_similar_text(material_measure_items, similarity_threshold=0.90)
    work_measure_items = remove_similar_text(work_measure_items, similarity_threshold=0.90)

    # =========================
    # 4) 기본 문구
    # =========================
    if not material_risk_items:
        material_risk_items = [
            "입력된 물질정보를 기준으로 누출, 접촉, 흡입 가능성을 확인해야 합니다.."
        ]

    if not work_risk_items:
        work_risk_items = [
            "입력된 작업유형을 기준으로 오조작, 점화원, 작업환경 위험요인을 확인해야 합니다."
        ]

    if not material_measure_items:
        material_measure_items = [
            "작업 전 MSDS 확인, 보호구 착용, 환기 상태 확인이 필요합니다."
        ]

    if not work_measure_items:
        work_measure_items = [
            "작업 전 안전절차 확인, 작업구역 통제, 비상대응 절차 숙지가 필요합니다."
        ]

    return (
    material_risk_items[:5],
    work_risk_items[:5],
    material_measure_items[:5],
    work_measure_items[:5]
)

# =========================
# 산업재해 아카이브 참고자료 검색
# (hazard_mapping.xlsx 매칭 로직과 별개로, AI 프롬프트에 추가로 참고시킬
#  실제 산업재해 사례 3건을 키워드 매칭으로 추려낸다.)
# =========================

_ARCHIVE_STOPWORDS = {"작업", "관련", "등", "및", "위한", "대한", "작업중", "중"}

# hazard_mapping.xlsx의 물질군(hazard_scores 키)을 아카이브 "기인물" 표현으로 매핑.
# 기인물 컬럼은 "화학물질(황산)", "인화성물질(톨루엔)"처럼 분류+물질명 조합이지만,
# 실제 값을 전수 확인한 결과 "인화성물질"/"유해가스"/"화학물질" 세 표현만 쓰이고
# "부식성"·"독성"·"산화성"·"고압가스" 같은 표현은 존재하지 않는다. 따라서 실제
# 아카이브에 있는 표현으로만 매핑해야 검색이 의미가 있다.
_HAZARD_GROUP_TO_ARCHIVE_KEYWORDS = {
    "금속부식성 물질": ["화학물질"],
    "급성독성(흡입)": ["유해가스", "화학물질"],
    "인체독성": ["화학물질", "유해가스"],
    "인화성": ["인화성물질"],
    "산화성": ["화학물질"],
    "고압가스": ["유해가스", "화학물질"],
    "급성 수생환경 유해성": ["화학물질"],
    "자연발화 및 과산화물": ["인화성물질", "화학물질"],
}


def _extract_archive_keywords(text, min_len=2):
    tokens = re.split(r"[\s/,()\[\]·\-]+", str(text or ""))
    return [
        t.strip() for t in tokens
        if len(t.strip()) >= min_len and t.strip() not in _ARCHIVE_STOPWORDS
    ]


def get_archive_reference_cases(work_name, chem_name, hazard_scores, top_n=3):
    """작업명·취급물질을 기준으로 산업재해 아카이브에서 관련성 높은 사례를 찾는다.
    프롬프트 비용 관리를 위해 최대 top_n건만 반환하고, 관련 사례가 없으면 빈 리스트를 준다."""
    df = accident_archive_df

    if df is None or df.empty:
        return []

    work_keywords = _extract_archive_keywords(work_name)

    material_keywords = set()
    chem_name_clean = str(chem_name or "").strip()

    if chem_name_clean and chem_name_clean.lower() != "nan":
        material_keywords.add(chem_name_clean)
        # "메틸 알코올"처럼 공백이 섞인 정식 화학물질명은 기인물 표기와
        # 그대로 일치하지 않는 경우가 많아, 공백 기준 조각도 함께 추가한다.
        for token in chem_name_clean.split():
            if len(token) >= 2:
                material_keywords.add(token)

    for hazard_name, score in (hazard_scores or {}).items():
        try:
            score_value = float(score)
        except Exception:
            score_value = 0

        if score_value > 0:
            for kw in _HAZARD_GROUP_TO_ARCHIVE_KEYWORDS.get(str(hazard_name).strip(), []):
                material_keywords.add(kw)

    if not work_keywords and not material_keywords:
        return []

    work_text = df["재해개요"] + " " + df["고위험작업·상황"]
    origin_text = df["기인물"]

    work_score = pd.Series(0, index=df.index)
    for kw in work_keywords:
        work_score = work_score + work_text.str.contains(re.escape(kw), na=False).astype(int)

    material_score = pd.Series(0, index=df.index)
    for kw in material_keywords:
        material_score = material_score + origin_text.str.contains(re.escape(kw), na=False).astype(int)

    # 물질(기인물) 일치가 작업명 키워드 일치보다 더 신뢰도 높은 신호라 가중치를 더 준다.
    total_score = work_score.astype(float) + (material_score.astype(float) * 1.5)

    scored_df = df.assign(_score=total_score)
    scored_df = scored_df[scored_df["_score"] > 0].sort_values("_score", ascending=False)

    cases = []
    for _, row in scored_df.head(top_n).iterrows():
        cases.append({
            "재해개요": row["재해개요"][:100],
            "고위험작업·상황": row["고위험작업·상황"][:40],
            "재해유발요인": row["재해유발요인"][:120],
            "위험성감소대책": row["위험성 감소대책(예시)"][:120],
        })

    return cases


def format_archive_reference_text(cases):
    if not cases:
        return "관련 산업재해 사례 없음"

    lines = []
    for i, case in enumerate(cases, 1):
        lines.append(
            f"{i}) 재해개요: {case['재해개요']}\n"
            f"   재해유발요인: {case['재해유발요인']}\n"
            f"   위험성 감소대책: {case['위험성감소대책']}"
        )

    return "\n".join(lines)


def generate_ai_text(
    material_risk_items,
    work_risk_items,
    material_measure_items,
    work_measure_items,
    similar_text,
    archive_reference_text="관련 산업재해 사례 없음",
    work_name="",
    work_content=""
):
    material_risk_source = "\n".join(material_risk_items)
    work_risk_source = "\n".join(work_risk_items)
    material_measure_source = "\n".join(material_measure_items)
    work_measure_source = "\n".join(work_measure_items)

    work_name_text = work_name.strip() if work_name else "정보 없음"
    work_content_text = work_content.strip() if work_content else "정보 없음"

    prompt = f"""
너는 화학안전 TBM 전문가다.

아래 원문을 바탕으로 작업 전 TBM 안내문을 작성하라.

[이번 작업 정보]
작업명: {work_name_text}
작업내용: {work_content_text}

[이번 작업 반영 원칙]
- 이번 작업의 핵심은 위 [이번 작업 정보]다. 아래 참고자료(물질 위험성 DB, 산업재해 사례)는 이 작업을 뒷받침하는 보조 자료로만 활용할 것
- 참고자료 중에서도 [이번 작업 정보]의 작업명·작업내용에 등장하는 설비·행위(예: 펌프·회전체, 밸브, 배관, 탱크, 용접 등)와 직접 관련된 항목을 최우선으로 골라 반영할 것
- 참고자료가 일반적인 배관·밸브 위주 내용이더라도, 작업명이 가리키는 실제 설비 특성(예: 회전체·베어링·커플링 등 펌프 관련 위험)이 있다면 그 특성을 위험요인/안전대책에 반드시 반영할 것

[공통 작성 원칙]
- 제공된 원문 내용만 사용할 것
- 새로운 위험요인, 사고사례, 법령, 수치를 만들지 말 것
- 같은 의미의 문장을 반복하지 말 것
- 작업반장이 작업 전 1분 브리핑하듯 작성할 것
- 어려운 전문용어 사용을 최소화할 것
- 작업자가 바로 이해할 수 있도록 짧고 명확하게 작성할 것

[참고자료 활용 원칙]
- [참고자료 1] 물질/작업 위험성 DB와 [참고자료 2] 실제 산업재해 사례를 모두 고려하여 이 작업의 위험요인과 안전대책을 구체적으로 작성할 것
- [참고자료 2]는 실제 사고 사례이므로, 유사한 사고 흐름(끼임·누출·낙하 등 사고 유형)이 있다면 위험요인 작성에 반영할 것
- [참고자료 2]에 관련 사례가 없다고 나오면 [참고자료 1]만으로 작성할 것

[AI 주요 위험요인 작성 원칙]
- 물질군 위험요인과 작업유형 위험요인을 따로 나열하지 말 것
- 반드시 물질의 위험성과 작업유형의 위험상황을 한 문장 안에서 연결해서 설명할 것
- 특정 화학물질명, 제품명, CAS 번호를 사용하지 말 것
- [이번 작업 정보]의 작업명·작업내용 문장을 그대로 베껴 쓰지 말 것 (단, 그 안에 담긴 설비·행위 특성은 위험상황 설명에 반영할 것)
- 물질명 대신 부식성, 인화성, 독성, 산화성 등 유해·위험 특성으로 표현할 것
- 작업명을 문장 그대로 반복하는 대신, 설비 개방, 연결부 분리, 잔압 제거, 누출, 비산, 회전체 접촉 등 이번 작업에 실제로 해당하는 위험상황으로 구체적으로 표현할 것
- “질산 취급 시”, “메탄올 작업 중”, “배관 교체 시”와 같은 문장 형식을 사용하지 말 것
- 일반적인 화학안전 설명보다 현재 작업상황 중심으로 작성할 것
- 개별 위험요인을 단순 나열하지 말고 작업 흐름처럼 연결해서 설명할 것
- 각 항목은 최대 3개까지만 작성할 것
- 한 문장에는 최대 2개의 위험요인만 포함할 것

[안전 및 예방조치 작성 원칙]
- 작업자가 바로 실행할 수 있는 행동 중심으로 작성할 것
- 물질 특성과 작업유형 상황을 함께 고려한 예방조치로 작성할 것
- 특정 화학물질명을 사용하지 말 것
- [이번 작업 정보]의 작업명·작업내용 문장을 그대로 베껴 쓰지 말 것 (단, 그 안에 담긴 설비·행위 특성은 조치에 반영할 것)
- 물질명 대신 해당 물질의 유해·위험 특성을 사용하여 작성할 것
- 작업명을 문장 그대로 반복하는 대신, 이번 작업의 설비 상태 및 작업자의 행동을 중심으로 구체적으로 작성할 것
- 각 항목은 최대 3개까지만 작성할 것
- 한 문장에는 최대 2개의 조치만 포함할 것

[유사사고 작성 원칙]
- 사고일시, 지역, 작업유형을 반드시 포함할 것
- 유사사고 문장은 “2021년 5월 충청북도 제천시 소재 사업장에서 정기작업 중 ...” 형태를 기본으로 작성할 것
- 사고내용은 삭제하거나 과도하게 축약하지 말 것
- 사고내용의 핵심 단어를 유지하면서 자연스럽게 문장 형태로 재구성할 것
- 작업 전 주의사항 형태로 마무리할 것
- 최대 2문장으로 작성할 것
- 새로운 사고내용 생성 금지


[참고자료 1] 물질/작업 위험성 DB

[물질군 유해위험요인]
{material_risk_source}

[작업유형 주요 위험요인]
{work_risk_source}

[물질군 취급시 주의사항 및 예방조치]
{material_measure_source}

[작업유형 안전대책]
{work_measure_source}

[참고자료 2] 실제 산업재해 사례 (한국산업안전보건공단 아카이브)
{archive_reference_text}

[유사사고 원문]
{similar_text}

[출력 형식]
[AI 주요 위험요인]
- 물질 특성과 작업유형을 연결한 위험요인 3개

[안전 및 예방조치]
- 물질 특성과 작업유형을 연결한 예방조치 3개

[유사사고]
- 유사사고를 작업 전 주의사항 형태로 1~2문장
"""

    response = client.chat.completions.create(
        model="gpt-4.1-mini",
        messages=[
            {"role": "system", "content": "너는 화학안전 전문가다."},
            {"role": "user", "content": prompt}
        ],
        temperature=0.3
    )

    return response.choices[0].message.content

def parse_ai_result(ai_result):
    sections = {
        "risk": "",
        "measure": "",
        "accident": ""
    }

    text = str(ai_result)

    try:
        risk_part = text.split("[AI 주요 위험요인]")[1].split("[안전 및 예방조치]")[0].strip()
        measure_part = text.split("[안전 및 예방조치]")[1].split("[유사사고]")[0].strip()
        accident_part = text.split("[유사사고]")[1].strip()

        sections["risk"] = risk_part
        sections["measure"] = measure_part
        sections["accident"] = accident_part

    except Exception:
        sections["risk"] = text

    return sections

def create_team(
    team_name,
    team_password,
    manager_name,
    manager_password,
    workers
):

    data = {
        "team_name": team_name,
        "password": team_password,
        "manager_name": manager_name,
        "manager_password": manager_password,
        "workers": workers
    }

    result = (
        supabase.table("teams")
        .insert(data)
        .execute()
    )

    return result.data

def login_team(team_name, password):

    result = (
        supabase.table("teams")
        .select("*")
        .eq("team_name", team_name)
        .eq("password", password)
        .execute()
    )

    if result.data:
        return result.data[0]

    return None

def login_manager(team_name, manager_password):

    result = (
        supabase.table("teams")
        .select("*")
        .eq("team_name", st.session_state.get("team_name", ""))
        .execute()
    )

    if result.data:
        return result.data[0]

    return None

def create_work_task(
    team_id,
    team_name,
    work_name,
    work_content,
    work_location,
    work_date,
    scheduled_time,
    tbm_place,
    assigned_workers,
    risk_assessment_done=True,
    work_type_code="",
    work_type_display="",
    start_time="",
    time_slot="",
    material_name="",
    chem_id="",
    chem_name="",
    risk_score=None,
    risk_level="",
    main_hazard_1="",
    main_hazard_2="",
    main_hazard_3="",
    safety_measure_1="",
    safety_measure_2="",
    safety_measure_3="",
    selected_hazard="",
    selected_measure="",
    similar_accident_text=""
):
    data = {
        "team_id": team_id,
        "team_name": team_name,
        "work_name": work_name,
        "work_content": work_content,
        "work_location": work_location,
        "work_date": work_date,
        "scheduled_time": scheduled_time,
        "tbm_place": tbm_place,
        "assigned_workers": assigned_workers,
        "risk_assessment_done": risk_assessment_done,
        "work_type_code": work_type_code,
        "work_type_display": work_type_display,
        "start_time": start_time,
        "time_slot": time_slot,
        "material_name": material_name,
        "chem_id": chem_id,
        "chem_name": chem_name,
        "risk_score": risk_score,
        "risk_level": risk_level,
        "main_hazard_1": main_hazard_1,
        "main_hazard_2": main_hazard_2,
        "main_hazard_3": main_hazard_3,
        "safety_measure_1": safety_measure_1,
        "safety_measure_2": safety_measure_2,
        "safety_measure_3": safety_measure_3,
        "selected_hazard": selected_hazard,
        "selected_measure": selected_measure,
        "similar_accident_text": similar_accident_text
    }

    result = (
        supabase.table("work_tasks")
        .insert(data)
        .execute()
    )

    return result.data


def get_work_templates(team_id):
    if not team_id:
        return []

    result = (
        supabase.table("work_templates")
        .select("*")
        .eq("team_id", team_id)
        .order("created_at", desc=True)
        .execute()
    )

    return result.data if result.data else []


def create_work_template(
    team_id,
    team_name,
    template_name,
    work_name,
    work_content,
    work_location,
    tbm_place,
    work_type_code,
    work_type_display,
    material_name
):
    data = {
        "team_id": team_id,
        "team_name": team_name,
        "template_name": template_name,
        "work_name": work_name,
        "work_content": work_content,
        "work_location": work_location,
        "tbm_place": tbm_place,
        "work_type_code": work_type_code,
        "work_type_display": work_type_display,
        "material_name": material_name,
    }

    result = (
        supabase.table("work_templates")
        .insert(data)
        .execute()
    )

    return result.data


def delete_work_template(template_id):
    supabase.table("work_templates").delete().eq("id", template_id).execute()


def get_team_workers(team_id):
    if not team_id:
        return []

    result = (
        supabase.table("teams")
        .select("workers")
        .eq("id", team_id)
        .execute()
    )

    if result.data:
        return result.data[0].get("workers") or []

    return []


def get_work_tasks_by_date(team_id, work_date):
    if not team_id:
        return []

    result = (
        supabase.table("work_tasks")
        .select("*")
        .eq("team_id", team_id)
        .eq("work_date", work_date)
        .order("created_at", desc=True)
        .execute()
    )

    return result.data if result.data else []


def get_today_work_tasks(team_id, client_date=None):
    today = client_date or datetime.now().strftime("%Y-%m-%d")
    return get_work_tasks_by_date(team_id, today)

def get_all_work_tasks(team_id):
    if not team_id:
        return []

    result = (
        supabase.table("work_tasks")
        .select("*")
        .eq("team_id", team_id)
        .order("created_at", desc=True)
        .execute()
    )

    return result.data if result.data else []

def get_recent_past_work_tasks(team_id, today_str, days=7):
    """오늘을 제외한, 최근 days일 이내(work_date가 [오늘-days, 오늘) 범위) 작업을
    work_date 내림차순(최근 지난 작업이 먼저)으로 반환한다."""
    if not team_id:
        return []

    cutoff_dt = datetime.strptime(today_str, "%Y-%m-%d") - timedelta(days=days)
    cutoff_str = cutoff_dt.strftime("%Y-%m-%d")

    all_tasks = get_all_work_tasks(team_id)

    past_tasks = [
        t for t in all_tasks
        if t.get("work_date") and cutoff_str <= t.get("work_date") < today_str
    ]

    past_tasks.sort(key=lambda t: t.get("work_date", ""), reverse=True)

    return past_tasks

def _set_cell_lines(cell, lines):
    """빈 병합 셀에 여러 줄을 문단 단위로 채워 넣는다."""
    if not lines:
        lines = ["-"]

    cell.paragraphs[0].text = lines[0]
    for line in lines[1:]:
        cell.add_paragraph(line)


def _mark_checkbox(cell, yes):
    """'예 □ 아니오 □' 형태의 셀에서 해당하는 □ 런을 ■로 치환한다."""
    boxes = []
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            if "□" in run.text:
                boxes.append(run)

    if len(boxes) < 2:
        return

    target = boxes[0] if yes else boxes[1]
    target.text = target.text.replace("□", "■")


SIGNATURE_BUCKET = "tbm-signatures"

ATTENDEE_NAME_COLS = (0, 5, 9)
ATTENDEE_NAME_TO_SIG_COL = {0: 2, 5: 8, 9: 11}
ATTENDEE_FIRST_ROW = 22
ATTENDEE_ORIGINAL_LAST_ROW = 25

SIGNATURE_CELL_WIDTH_EMU = 950000   # 표의 서명 칸 실측 폭(약 1.1인치)보다 살짝 작게
SIGNATURE_CELL_HEIGHT_EMU = 320000  # 행 높이가 과도하게 늘어나지 않도록 하는 상한


def upload_signature(task_id, team_id, worker_name, png_bytes):
    """서명 PNG를 Supabase Storage에 올리고 tbm_signatures 테이블에 기록한다."""
    if not png_bytes:
        return None

    # Supabase Storage 객체 키는 한글 등 비-ASCII 문자를 허용하지 않아(InvalidKey),
    # 작성자명은 경로에 넣지 않고 tbm_signatures 테이블 컬럼으로만 관리한다.
    saved_at = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    file_stamp = datetime.now().strftime("%Y%m%d%H%M%S%f")
    storage_path = f"{task_id}/sig_{file_stamp}.png"

    try:
        supabase.storage.from_(SIGNATURE_BUCKET).upload(
            storage_path,
            png_bytes,
            {"content-type": "image/png"}
        )

        supabase.table("tbm_signatures").insert({
            "task_id": task_id,
            "team_id": team_id,
            "worker_name": worker_name,
            "storage_path": storage_path,
            "created_at": saved_at,
        }).execute()

        return storage_path
    except Exception as e:
        print("서명 이미지 업로드 오류:", e)
        return None


def upsert_work_log(fields, task_id, worker_name):
    """TBM 체크리스트 완료 → 작업일지 제출, 두 단계에 걸쳐 채워지는 work_logs 레코드를
    하나로 유지한다. (task_id, worker_name) 조합으로 기존 행을 찾아 있으면 갱신하고,
    없으면 새로 만든다 (체크리스트 단계에서 먼저 만들고, 작업일지 제출 단계에서 갱신)."""
    existing = (
        supabase.table("work_logs")
        .select("id")
        .eq("task_id", task_id)
        .eq("worker_name", worker_name)
        .limit(1)
        .execute()
    )

    if existing.data:
        row_id = existing.data[0]["id"]
        supabase.table("work_logs").update(fields).eq("id", row_id).execute()
        return row_id

    inserted = supabase.table("work_logs").insert(fields).execute()
    return inserted.data[0]["id"] if inserted.data else None


def fetch_signatures_for_task(task_id):
    """작업(task_id)에 대해 제출된 서명 이미지를 작성자명 기준으로 내려받는다."""
    signatures_by_worker = {}

    try:
        result = (
            supabase.table("tbm_signatures")
            .select("*")
            .eq("task_id", task_id)
            .order("created_at", desc=False)
            .execute()
        )

        for row in (result.data or []):
            worker_name = row.get("worker_name")
            storage_path = row.get("storage_path")

            if not worker_name or not storage_path:
                continue

            try:
                png_bytes = supabase.storage.from_(SIGNATURE_BUCKET).download(storage_path)
                signatures_by_worker[worker_name] = png_bytes
            except Exception as e:
                print(f"서명 이미지 다운로드 오류 ({worker_name}):", e)

    except Exception as e:
        print("서명 이미지 조회 오류:", e)

    return signatures_by_worker


def cleanup_expired_signatures(team_id):
    """TBM 워드 파일이 생성된 작업에 한해, 저장된 지 48시간 지난 서명 이미지를 정리한다."""
    if not team_id:
        return

    try:
        tasks_result = (
            supabase.table("work_tasks")
            .select("id, tbm_docx_generated_at")
            .eq("team_id", team_id)
            .execute()
        )

        eligible_task_ids = [
            t["id"] for t in (tasks_result.data or [])
            if t.get("tbm_docx_generated_at")
        ]

        if not eligible_task_ids:
            return

        cutoff = (datetime.now() - timedelta(hours=48)).strftime("%Y-%m-%d %H:%M:%S")

        sig_result = (
            supabase.table("tbm_signatures")
            .select("id, storage_path, task_id, created_at")
            .in_("task_id", eligible_task_ids)
            .lt("created_at", cutoff)
            .execute()
        )

        expired = sig_result.data or []
        if not expired:
            return

        paths = [row["storage_path"] for row in expired if row.get("storage_path")]
        if paths:
            supabase.storage.from_(SIGNATURE_BUCKET).remove(paths)

        ids = [row["id"] for row in expired]
        supabase.table("tbm_signatures").delete().in_("id", ids).execute()

    except Exception as e:
        print("서명 이미지 자동 삭제 오류:", e)


def _add_attendee_row(table):
    """참석자 표의 마지막 행과 동일한 구조(병합 포함)로 새 행을 하나 복제해 추가한다."""
    last_tr = table.rows[-1]._tr
    new_tr = copy.deepcopy(last_tr)
    table._tbl.append(new_tr)
    return table.rows[-1]


def _insert_signature_image(cell, png_bytes):
    """서명 이미지를 비율을 유지한 채 칸 크기에 맞춰 삽입한다."""
    if not png_bytes:
        return

    try:
        with Image.open(BytesIO(png_bytes)) as im:
            width_px, height_px = im.size
    except Exception as e:
        print("서명 이미지 열기 오류:", e)
        return

    if not width_px or not height_px:
        return

    ratio = width_px / height_px
    target_w = SIGNATURE_CELL_WIDTH_EMU
    target_h = int(target_w / ratio)

    if target_h > SIGNATURE_CELL_HEIGHT_EMU:
        target_h = SIGNATURE_CELL_HEIGHT_EMU
        target_w = int(target_h * ratio)

    cell.paragraphs[0].add_run().add_picture(
        BytesIO(png_bytes), width=Emu(target_w), height=Emu(target_h)
    )


TBM_DOCX_FONT = "맑은 고딕"


def _set_run_font(run, font_name=TBM_DOCX_FONT):
    """
    run의 서체를 지정한다. python-docx의 font.name은 영문(ascii/hAnsi)만 바꾸고
    한글은 eastAsia 서체를 따로 봐서, 이걸 안 맞춰주면 셀마다 서체가 다르게 보인다.
    """
    run.font.name = font_name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)


def _normalize_docx_fonts(doc, font_name=TBM_DOCX_FONT):
    """
    문서 전체(본문 문단 + 표, 중첩된 표 포함)의 모든 run 서체를 하나로 통일한다.
    cell.text=... / add_paragraph(...)로 새로 채운 텍스트는 템플릿 서체를 안 물려받고
    기본 서체(Calibri 등)로 들어가서, 렌더링 후 마지막에 한 번 덮어써야 한다.
    """
    def process_paragraphs(paragraphs):
        for paragraph in paragraphs:
            for run in paragraph.runs:
                _set_run_font(run, font_name)

    def process_table(table):
        for row in table.rows:
            for cell in row.cells:
                process_paragraphs(cell.paragraphs)
                for nested_table in cell.tables:
                    process_table(nested_table)

    process_paragraphs(doc.paragraphs)
    for table in doc.tables:
        process_table(table)


def generate_tbm_docx(task, logs, signatures_by_worker=None):
    template_path = "tbm_template.docx"

    doc = DocxTemplate(template_path)

    leader_log = None
    for log in logs:
        if log.get("is_tbm_leader") is True:
            leader_log = log
            break

    if leader_log is None and logs:
        leader_log = logs[0]

    # TBM 일시 = 그날 작업자 중 TBM을 가장 먼저 "완료 및 저장"한 시각
    # (work_tasks.first_tbm_submitted_at, show_checklist()에서 최초 1회만 기록됨).
    # 이 값이 없는 과거 작업(마이그레이션 이전 데이터)은 TBM 리더의 제출 시각으로 대체한다.
    tbm_end = task.get("first_tbm_submitted_at") or (
        leader_log.get("tbm_end_time", "") if leader_log else ""
    )

    tbm_datetime = ""
    try:
        end_dt = datetime.strptime(tbm_end, "%Y-%m-%d %H:%M:%S")
        tbm_datetime = end_dt.strftime("%Y년 %m월 %d일 %H:%M")
    except Exception:
        tbm_datetime = tbm_end

    worker_list = [
        log.get("worker_name", "")
        for log in logs
        if log.get("worker_name")
    ]

    context = {
        "tbm_datetime": tbm_datetime,
        "work_name": task.get("work_name", ""),
        "work_content": task.get("work_content", ""),
        "tbm_place": task.get("tbm_place", ""),
        "leader_department": leader_log.get("leader_department", "") if leader_log else "",
        "leader_potition": leader_log.get("leader_position", "") if leader_log else "",
        "leader_name": leader_log.get("leader_name", "") if leader_log else "",

        # 위험요인/안전대책은 이제 작업자별이 아니라 안전관리자가 "오늘 작업 입력"에서
        # 한 번 분석해 둔 값(work_tasks)을 그대로 쓴다.
        "main_hazard_1": task.get("main_hazard_1", ""),
        "main_hazard_2": task.get("main_hazard_2", ""),
        "main_hazard_3": task.get("main_hazard_3", ""),

        "safety_measure_1": task.get("safety_measure_1", ""),
        "safety_measure_2": task.get("safety_measure_2", ""),
        "safety_measure_3": task.get("safety_measure_3", ""),
    }

    doc.render(context)

    table = doc.tables[0]

    # 위험성평가 실시여부 체크 표시
    try:
        _mark_checkbox(table.cell(3, 10), bool(task.get("risk_assessment_done")))
    except Exception as e:
        print("위험성평가 실시여부 표시 오류:", e)

    # 중점위험요인 선정 / 대책 (안전관리자가 "오늘 작업 입력"에서 선택한 값)
    try:
        table.cell(8, 4).text = task.get("selected_hazard", "") or ""
        table.cell(9, 4).text = task.get("selected_measure", "") or ""
    except Exception as e:
        print("중점위험요인 선정/대책 입력 오류:", e)

    # 작업 전 일일 안전점검 시행 결과 (작성자별 목록)
    try:
        daily_check_lines = [
            f"- {log.get('worker_name', '-')}: {log.get('daily_safety_check_result')}"
            for log in logs
            if log.get("daily_safety_check_result")
        ]
        _set_cell_lines(table.cell(17, 0), daily_check_lines)
    except Exception as e:
        print("작업 전 일일 안전점검 시행 결과 입력 오류:", e)

    # 작업 후 종료 미팅 (작성자별 목록)
    try:
        closing_meeting_lines = [
            f"- {log.get('worker_name', '-')}: {log.get('closing_meeting_result')}"
            for log in logs
            if log.get("closing_meeting_result")
        ]
        _set_cell_lines(table.cell(19, 0), closing_meeting_lines)
    except Exception as e:
        print("작업 후 종료 미팅 입력 오류:", e)

    # 참석자 확인란 자동 입력 (행당 이름 칸: 열 0, 5, 9 / 서명 칸: 열 2, 8, 11)
    try:
        signatures_by_worker = signatures_by_worker or {}
        slots_per_row = len(ATTENDEE_NAME_COLS)
        base_row_count = ATTENDEE_ORIGINAL_LAST_ROW - ATTENDEE_FIRST_ROW + 1

        rows_needed = -(-len(worker_list) // slots_per_row) if worker_list else 0  # 올림 나눗셈
        rows_needed = max(rows_needed, base_row_count)

        for _ in range(rows_needed - base_row_count):
            _add_attendee_row(table)

        attendee_row_indices = list(range(ATTENDEE_FIRST_ROW, ATTENDEE_FIRST_ROW + rows_needed))
        attendee_cells = [
            (row, col)
            for row in attendee_row_indices
            for col in ATTENDEE_NAME_COLS
        ]

        for worker_name, cell_pos in zip(worker_list, attendee_cells):
            row_idx, name_col = cell_pos
            table.cell(row_idx, name_col).text = worker_name

            sig_col = ATTENDEE_NAME_TO_SIG_COL[name_col]
            sig_bytes = signatures_by_worker.get(worker_name)
            if sig_bytes:
                _insert_signature_image(table.cell(row_idx, sig_col), sig_bytes)

    except Exception as e:
        print("참석자 확인란 입력 오류:", e)

    try:
        _normalize_docx_fonts(doc)
    except Exception as e:
        print("서체 통일 오류:", e)

    output = BytesIO()
    doc.save(output)
    output.seek(0)

    return output

def add_worker_to_team(team_id, new_worker_name):
    result = (
        supabase.table("teams")
        .select("workers")
        .eq("id", team_id)
        .execute()
    )

    if not result.data:
        return False, "팀 정보를 찾을 수 없습니다."

    current_workers = result.data[0].get("workers") or []

    if new_worker_name in current_workers:
        return False, "이미 등록된 작업자입니다."

    current_workers.append(new_worker_name)

    supabase.table("teams").update({
        "workers": current_workers
    }).eq("id", team_id).execute()

    return True, "작업자가 추가되었습니다."

def split_ai_bullets(text):
    lines = []

    for line in str(text).splitlines():
        line = line.strip()

        if not line:
            continue

        line = line.lstrip("-").strip()
        lines.append(line)

    return lines

def make_casualty_text(row):
    death_direct = pd.to_numeric(row.get("사망(직접)", 0), errors="coerce")
    death_other = pd.to_numeric(row.get("사망(기타)", 0), errors="coerce")
    injury_direct = pd.to_numeric(row.get("부상(직접)", 0), errors="coerce")
    injury_other = pd.to_numeric(row.get("부상(기타)", 0), errors="coerce")

    death_count = int((0 if pd.isna(death_direct) else death_direct) + (0 if pd.isna(death_other) else death_other))
    injury_count = int((0 if pd.isna(injury_direct) else injury_direct) + (0 if pd.isna(injury_other) else injury_other))

    parts = []

    if death_count > 0:
        parts.append(f"사망자 {death_count}명")

    if injury_count > 0:
        parts.append(f"부상자 {injury_count}명")

    if not parts:
        return ""

    return ", ".join(parts) + "이 발생한 사고."


def find_similar_accident(result, current_month=None):
    """
    current_month를 명시적으로 넘기면 그 값을 그대로 쓰고(예: 접속 기기의 현재 월),
    안 넘기면 기존처럼 result["작업시간"]에서 월을 뽑아보고 실패하면 서버 시각으로 폴백한다.
    """
    df = final_result.copy()

    work_type = str(result.get("작업유형", "")).strip()
    chem_id = str(result.get("chem_id", "")).strip().zfill(6)
    df["CHEMID"] = df["CHEMID"].astype(str).str.strip().str.zfill(6)

    if current_month is None:
        work_time = result.get("작업시간", "")
        try:
            current_month = pd.to_datetime(work_time).month
        except Exception:
            current_month = datetime.now().month

    df["작업유형"] = df["작업유형"].astype(str).str.strip()
    df["CHEMID"] = df["CHEMID"].astype(str).str.strip()

    df["month"] = pd.to_numeric(df["month"], errors="coerce")
    df["date"] = pd.to_datetime(df["date"], errors="coerce")

    search_steps = [
        ("작업유형·물질·월 일치", df[
            (df["작업유형"] == work_type) &
            (df["CHEMID"] == chem_id) &
            (df["month"] == current_month)
        ]),
        ("작업유형·물질 일치", df[
            (df["작업유형"] == work_type) &
            (df["CHEMID"] == chem_id)
        ]),
        ("작업유형·월 일치", df[
            (df["작업유형"] == work_type) &
            (df["month"] == current_month)
        ]),
        ("작업유형 일치", df[
            (df["작업유형"] == work_type)
        ]),
    ]

    for match_level, matched_df in search_steps:
        matched_df = matched_df.dropna(subset=["date"])

        if not matched_df.empty:
            latest = matched_df.sort_values("date", ascending=False).iloc[0]
            return latest, match_level

    return None, None


def render_traffic_light(score):
    style = get_score_style(score)

    red_on = "on-red" if "red-on" in style["red_class"] else ""
    yellow_on = "on-yellow" if "yellow-on" in style["yellow_class"] else ""
    green_on = "on-green" if "green-on" in style["green_class"] else ""

    html = f"""
    <!DOCTYPE html>
    <html>
    <head>
    <style>
        body {{
            margin: 0;
            padding: 0;
            background: transparent;
            font-family: Arial, sans-serif;
        }}

        .result-card {{
            background: #ffffff;
            border: 1px solid #c5c6cd;
            border-radius: 18px;
            padding: 18px;
            box-shadow: 0 4px 14px rgba(15, 23, 42, 0.07);
            box-sizing: border-box;
        }}

        .result-card-title {{
            font-size: 18px;
            font-weight: 900;
            color: #091426;
            margin-bottom: 14px;
        }}

        .traffic-wrap {{
            display: flex;
            justify-content: center;
            padding: 12px 0 6px 0;
        }}

        .traffic-body-horizontal {{
            display: flex;
            align-items: center;
            justify-content: center;
            gap: 14px;
            background: #080b10;
            padding: 16px 18px;
            border-radius: 42px;
            box-shadow: inset 0 0 12px rgba(255,255,255,0.08), 0 10px 24px rgba(0,0,0,0.22);
        }}

        .traffic-light {{
            width: 66px;
            height: 66px;
            border-radius: 999px;
            border: 4px solid #222;
            background: #111;
            position: relative;
            overflow: hidden;
            opacity: 0.35;
            box-shadow: inset 0 4px 10px rgba(0,0,0,0.7);
        }}

        .traffic-light::before {{
            content: "";
            position: absolute;
            inset: 0;
            background:
                radial-gradient(circle at 35% 25%, rgba(255,255,255,0.55) 0%, transparent 28%),
                repeating-linear-gradient(45deg, transparent, transparent 3px, rgba(0,0,0,0.16) 3px, rgba(0,0,0,0.16) 6px);
        }}

        .on-red {{
            background: #ef4444;
            opacity: 1;
            box-shadow: 0 0 28px rgba(239,68,68,0.9), inset 0 4px 12px rgba(255,255,255,0.35);
        }}

        .on-yellow {{
            background: #facc15;
            opacity: 1;
            box-shadow: 0 0 28px rgba(250,204,21,0.9), inset 0 4px 12px rgba(255,255,255,0.35);
        }}

        .on-green {{
            background: #22c55e;
            opacity: 1;
            box-shadow: 0 0 28px rgba(34,197,94,0.9), inset 0 4px 12px rgba(255,255,255,0.35);
        }}
    </style>
    </head>
    <body>
        <div class="result-card">
            <div class="result-card-title">🚦 실시간 위험 수준</div>
            <div class="traffic-wrap">
                <div class="traffic-body-horizontal">
                    <div class="traffic-light {red_on}"></div>
                    <div class="traffic-light {yellow_on}"></div>
                    <div class="traffic-light {green_on}"></div>
                </div>
            </div>
        </div>
    </body>
    </html>
    """

    components.html(html, height=180, scrolling=False)


def render_score_card(score):
    style = get_score_style(score)

    dash_offset = 125 - (125 * score / 100)

    html = f"""
    <!DOCTYPE html>
    <html>
    <head>
    <style>
        body {{
            margin: 0;
            padding: 0;
            background: transparent;
            font-family: Arial, sans-serif;
        }}

        .result-card {{
            background: #ffffff;
            border: 1px solid #c5c6cd;
            border-radius: 18px;
            padding: 20px 16px;
            box-shadow: 0 4px 14px rgba(15, 23, 42, 0.07);
            box-sizing: border-box;
            text-align: center;
        }}

        .score-title {{
            font-size: 12px;
            letter-spacing: 0.22em;
            color: #8b8f98;
            font-weight: 800;
            margin-bottom: 8px;
        }}

        .gauge-box {{
            position: relative;
            width: 250px;
            height: 150px;
            margin: 0 auto;
        }}

        .gauge-box svg {{
            width: 250px;
            height: 150px;
        }}

        .score-center {{
            position: absolute;
            left: 0;
            right: 0;
            top: 58px;
            text-align: center;
        }}

        .score-small-label {{
            font-size: 13px;
            font-weight: 800;
            color: #45474c;
            margin-bottom: -4px;
        }}

        .score-number {{
            font-size: 64px;
            font-weight: 950;
            line-height: 1;
            -webkit-text-stroke: 1.8px #111827;
        }}

        .risk-badge {{
            display: inline-block;
            margin-top: 4px;
            padding: 8px 28px;
            border-radius: 999px;
            font-size: 16px;
            font-weight: 900;
        }}
    </style>
    </head>
    <body>
        <div class="result-card">
            <div class="score-title">SAFETY SCORE</div>

            <div class="gauge-box">
                <svg viewBox="0 0 100 60">
                    <defs>
                        <linearGradient id="safetyScoreGradient" x1="0%" y1="0%" x2="100%" y2="0%">
                            <stop offset="0%" stop-color="#22c55e"/>
                            <stop offset="50%" stop-color="#facc15"/>
                            <stop offset="100%" stop-color="#ef4444"/>
                        </linearGradient>
                    </defs>

                    <path d="M 10 50 A 40 40 0 0 1 90 50"
                        fill="none"
                        stroke="#e5e7eb"
                        stroke-linecap="round"
                        stroke-width="12"/>

                    <path d="M 10 50 A 40 40 0 0 1 90 50"
                        fill="none"
                        stroke="url(#safetyScoreGradient)"
                        stroke-dasharray="125"
                        stroke-dashoffset="{dash_offset}"
                        stroke-linecap="round"
                        stroke-width="12"/>
                </svg>

                <div class="score-center">
                    <div class="score-small-label">위험도</div>
                    <div class="score-number" style="color:{style['score_color']};">{score:.0f}</div>
                </div>
            </div>

            <div class="risk-badge" style="background:{style['badge_bg']}; color:{style['badge_color']};">
                {style['level_text']}
            </div>
        </div>
    </body>
    </html>
    """

    components.html(html, height=250, scrolling=False)

def render_risk_summary_card(score):
    style = get_score_style(score)

    dash_offset = 125 - (125 * score / 100)

    red_on = "on-red" if score >= 70 else ""
    yellow_on = "on-yellow" if 40 <= score < 70 else ""
    green_on = "on-green" if score < 40 else ""

    html = f"""
    <!DOCTYPE html>
    <html>
    <head>
    <style>
        body {{
            margin: 0;
            padding: 0;
            background: transparent;
            font-family: Arial, sans-serif;
        }}

        .summary-card {{
            background: #ffffff;
            border: 1px solid #c5c6cd;
            border-radius: 18px;
            padding: 20px 16px 18px 16px;
            box-shadow: 0 4px 14px rgba(15, 23, 42, 0.07);
            box-sizing: border-box;
            text-align: center;
        }}

        .score-title {{
            font-size: 11px;
            letter-spacing: 0.22em;
            color: #8b8f98;
            font-weight: 800;
            margin-bottom: 6px;
        }}

        .gauge-box {{
            position: relative;
            width: 230px;
            height: 132px;
            margin: 0 auto;
        }}

        .gauge-box svg {{
            width: 230px;
            height: 132px;
        }}

        .score-center {{
            position: absolute;
            left: 0;
            right: 0;
            top: 50px;
            text-align: center;
        }}

        .score-small-label {{
            font-size: 12px;
            font-weight: 800;
            color: #45474c;
            margin-bottom: -3px;
        }}

        .score-number {{
            font-size: 58px;
            font-weight: 950;
            line-height: 1;
            -webkit-text-stroke: 1.6px #111827;
        }}

        .risk-badge {{
            display: inline-block;
            margin-top: 2px;
            padding: 7px 24px;
            border-radius: 999px;
            font-size: 15px;
            font-weight: 900;
        }}

        .traffic-section {{
            margin-top: 18px;
            padding-top: 14px;
            border-top: 1px solid #e5e7eb;
        }}

        .traffic-label {{
            font-size: 13px;
            font-weight: 900;
            color: #091426;
            margin-bottom: 10px;
        }}

        .traffic-wrap {{
            display: flex;
            justify-content: center;
        }}

        .traffic-body {{
            display: flex;
            align-items: center;
            justify-content: center;
            gap: 9px;
            background: #080b10;
            padding: 10px 13px;
            border-radius: 999px;
            box-shadow: inset 0 0 10px rgba(255,255,255,0.08), 0 8px 18px rgba(0,0,0,0.18);
        }}

        .traffic-light {{
            width: 42px;
            height: 42px;
            border-radius: 999px;
            border: 3px solid #222;
            background: #111;
            position: relative;
            overflow: hidden;
            opacity: 0.35;
            box-shadow: inset 0 3px 8px rgba(0,0,0,0.7);
        }}

        .traffic-light::before {{
            content: "";
            position: absolute;
            inset: 0;
            background:
                radial-gradient(circle at 35% 25%, rgba(255,255,255,0.55) 0%, transparent 30%),
                repeating-linear-gradient(45deg, transparent, transparent 3px, rgba(0,0,0,0.16) 3px, rgba(0,0,0,0.16) 6px);
        }}

        .on-red {{
            background: #ef4444;
            opacity: 1;
            box-shadow: 0 0 20px rgba(239,68,68,0.9), inset 0 3px 9px rgba(255,255,255,0.35);
        }}

        .on-yellow {{
            background: #facc15;
            opacity: 1;
            box-shadow: 0 0 20px rgba(250,204,21,0.9), inset 0 3px 9px rgba(255,255,255,0.35);
        }}

        .on-green {{
            background: #22c55e;
            opacity: 1;
            box-shadow: 0 0 20px rgba(34,197,94,0.9), inset 0 3px 9px rgba(255,255,255,0.35);
        }}
    </style>
    </head>

    <body>
        <div class="summary-card">
            <div class="score-title">SAFETY SCORE</div>

            <div class="gauge-box">
                <svg viewBox="0 0 100 60">
                    <defs>
                        <linearGradient id="safetyScoreGradient" x1="0%" y1="0%" x2="100%" y2="0%">
                            <stop offset="0%" stop-color="#22c55e"/>
                            <stop offset="50%" stop-color="#facc15"/>
                            <stop offset="100%" stop-color="#ef4444"/>
                        </linearGradient>
                    </defs>

                    <path d="M 10 50 A 40 40 0 0 1 90 50"
                        fill="none"
                        stroke="#e5e7eb"
                        stroke-linecap="round"
                        stroke-width="12"/>

                    <path d="M 10 50 A 40 40 0 0 1 90 50"
                        fill="none"
                        stroke="url(#safetyScoreGradient)"
                        stroke-dasharray="125"
                        stroke-dashoffset="{dash_offset}"
                        stroke-linecap="round"
                        stroke-width="12"/>
                </svg>

                <div class="score-center">
                    <div class="score-small-label">위험도</div>
                    <div class="score-number" style="color:{style['score_color']};">{score:.0f}</div>
                </div>
            </div>

            <div class="risk-badge" style="background:{style['badge_bg']}; color:{style['badge_color']};">
                {style['level_text']}
            </div>

            <div class="traffic-section">
                <div class="traffic-label">실시간 위험 수준</div>
                <div class="traffic-wrap">
                    <div class="traffic-body">
                        <div class="traffic-light {red_on}"></div>
                        <div class="traffic-light {yellow_on}"></div>
                        <div class="traffic-light {green_on}"></div>
                    </div>
                </div>
            </div>
        </div>
    </body>
    </html>
    """

    st.iframe(html, height=365)


def render_ai_hazard_cards(risk_items):
    risk_html = ""
    for i, item in enumerate(risk_items, start=1):
        risk_html += f"""
<div class="message-item">
    <div class="message-num">{i:02d}</div>
    <div class="message-text">{item}</div>
</div>
"""

    html = f"""
<div class="result-card">
    <div class="result-card-header">
        <div class="result-card-icon">🧠</div>
        <div class="result-card-title">AI 분석 주요 위험요인</div>
    </div>
    {risk_html}
</div>
"""
    st.markdown(textwrap.dedent(html).strip(), unsafe_allow_html=True)


def render_safety_measure_cards(measure_items):
    measure_html = ""
    for item in measure_items:
        measure_html += f"""
<div class="measure-item">
    <div class="measure-check">✓</div>
    <div class="measure-text">{item}</div>
</div>
"""

    html = f"""
<div class="result-card">
    <div class="result-card-header">
        <div class="result-card-icon">🛡️</div>
        <div class="result-card-title">안전 및 사고 예방대책</div>
    </div>
    {measure_html}
</div>
"""
    st.markdown(textwrap.dedent(html).strip(), unsafe_allow_html=True)


def render_similar_accident_card(similar_accident_text):
    if similar_accident_text:
        html = f'''
<div class="result-card">
    <div class="result-card-header">
        <div class="result-card-icon">🕒</div>
        <div class="result-card-title">유사 사고사례</div>
    </div>
    <div class="incident-placeholder">
        <div class="incident-placeholder-desc">{similar_accident_text}</div>
    </div>
</div>
'''
    else:
        html = '''
<div class="result-card"><div class="result-card-header"><div class="result-card-icon">🕒</div><div class="result-card-title">유사 사고사례</div></div><div class="incident-placeholder"><div class="incident-placeholder-title">유사 사고사례 없음</div><div class="incident-placeholder-desc">입력한 작업유형과 물질정보를 기준으로 일치하는 사고사례를 찾지 못했습니다.</div></div></div>
'''
    st.markdown(html, unsafe_allow_html=True)


def render_selectable_hazard_cards(risk_items, session_key):
    """
    "AI 분석 주요 위험요인" 카드 자체를 버튼으로 만들어 렌더링한다 (안전관리자 전용).
    카드를 클릭하면 선택/해제(토글)되고, 선택된 항목의 인덱스는 st.session_state[session_key]에 저장된다.
    """
    st.markdown("""
<div class="result-card-header hazard-section-marker">
    <div class="result-card-icon">🧠</div>
    <div class="result-card-title">AI 분석 주요 위험요인 (탭하여 중점위험요인 선택)</div>
</div>
""", unsafe_allow_html=True)

    selected_idx = st.session_state.get(session_key)

    for i, item in enumerate(risk_items):
        is_selected = (selected_idx == i)

        check_prefix = "✅ " if is_selected else ""
        label = f"{check_prefix}:red[**{i + 1:02d}**]  {item}"

        if st.button(
            label,
            key=f"{session_key}_btn_{i}",
            use_container_width=True,
            type="primary" if is_selected else "secondary"
        ):
            st.session_state[session_key] = None if is_selected else i
            st.rerun()


def render_selectable_measure_cards(measure_items, session_key):
    """
    "안전 및 사고 예방대책" 카드 자체를 버튼으로 만들어 렌더링한다 (안전관리자 전용).
    카드를 클릭하면 선택/해제(토글)되고, 선택된 항목의 인덱스는 st.session_state[session_key]에 저장된다.
    """
    st.markdown("""
<div class="result-card-header measure-section-marker">
    <div class="result-card-icon">🛡️</div>
    <div class="result-card-title">안전 및 사고 예방대책 (탭하여 대책 선택)</div>
</div>
""", unsafe_allow_html=True)

    selected_idx = st.session_state.get(session_key)

    for i, item in enumerate(measure_items):
        is_selected = (selected_idx == i)

        check_prefix = "✅ " if is_selected else ""
        label = f"{check_prefix}:blue[**✓**]  {item}"

        if st.button(
            label,
            key=f"{session_key}_btn_{i}",
            use_container_width=True,
            type="primary" if is_selected else "secondary"
        ):
            st.session_state[session_key] = None if is_selected else i
            st.rerun()


def run_ai_hazard_analysis(result, current_month=None):
    """
    위험도 계산 결과(result)를 바탕으로 AI 위험요인/안전대책/유사사고 텍스트를 생성한다.
    작업자 모드(show_risk_result)와 안전관리자 모드(show_task_create)가 공유해서 쓴다.
    current_month를 넘기면 유사사고사례 매칭 시 그 월을 기준으로 쓴다.
    """
    material_risk_items, work_risk_items, material_measure_items, work_measure_items = get_risk_and_measure_messages(result)

    similar_accident, match_level = find_similar_accident(result, current_month=current_month)

    if similar_accident is not None:
        accident_date = similar_accident.get("date", "")
        accident_date_text = pd.to_datetime(accident_date).strftime("%Y-%m-%d")

        sido = str(similar_accident.get("도", "")).strip()
        sigungu = str(similar_accident.get("시군", "")).strip()
        location_text = f"{sido} {sigungu}"

        accident_work_type = str(similar_accident.get("작업유형", "-")).strip()
        accident_content = str(similar_accident.get("사고내용", "-")).strip()

        casualty_text = make_casualty_text(similar_accident)

        similar_text = f"""
    사고일시: {accident_date_text}
    사업장 소재지: {location_text}
    작업유형: {accident_work_type}
    사고내용: {accident_content}
    인명피해: {casualty_text}
    """
    else:
        similar_text = "유사사고 정보 없음"

    archive_cases = get_archive_reference_cases(
        work_name=result.get("작업명", ""),
        chem_name=result.get("chem_name", ""),
        hazard_scores=result.get("hazard_scores", {}),
        top_n=3
    )
    archive_reference_text = format_archive_reference_text(archive_cases)

    ai_result = generate_ai_text(
        material_risk_items,
        work_risk_items,
        material_measure_items,
        work_measure_items,
        similar_text,
        archive_reference_text=archive_reference_text,
        work_name=result.get("작업명", ""),
        work_content=result.get("작업내용", "")
    )

    ai_sections = parse_ai_result(ai_result)

    risk_items = split_ai_bullets(ai_sections["risk"])
    measure_items = split_ai_bullets(ai_sections["measure"])
    similar_accident_ai_text = ai_sections["accident"]

    return {
        "risk_items": risk_items,
        "measure_items": measure_items,
        "main_hazard_1": risk_items[0] if len(risk_items) > 0 else "",
        "main_hazard_2": risk_items[1] if len(risk_items) > 1 else "",
        "main_hazard_3": risk_items[2] if len(risk_items) > 2 else "",
        "safety_measure_1": measure_items[0] if len(measure_items) > 0 else "",
        "safety_measure_2": measure_items[1] if len(measure_items) > 1 else "",
        "safety_measure_3": measure_items[2] if len(measure_items) > 2 else "",
        "similar_accident": similar_accident,
        "similar_accident_ai_text": similar_accident_ai_text if similar_accident is not None else "",
        "match_level": match_level,
    }


def show_risk_result():
    if not st.session_state.get("result", {}).get("score"):
        _redirect_with_message("login", "먼저 작업자명과 오늘 작업을 선택해 주세요.")
        return

    if "ai_result" not in st.session_state:
        st.session_state.ai_result = None

    result = st.session_state.result
    score = float(result["score"])

    if st.session_state.ai_result is None:
        with st.spinner("AI가 TBM 내용을 생성 중입니다..."):
            st.session_state.ai_result = run_ai_hazard_analysis(result)

    ai_result = st.session_state.ai_result

    risk_items = ai_result["risk_items"]
    measure_items = ai_result["measure_items"]
    similar_accident = ai_result["similar_accident"]
    similar_accident_ai_text = ai_result["similar_accident_ai_text"]

    st.session_state.work_data["main_hazard_1"] = ai_result["main_hazard_1"]
    st.session_state.work_data["main_hazard_2"] = ai_result["main_hazard_2"]
    st.session_state.work_data["main_hazard_3"] = ai_result["main_hazard_3"]

    st.session_state.work_data["safety_measure_1"] = ai_result["safety_measure_1"]
    st.session_state.work_data["safety_measure_2"] = ai_result["safety_measure_2"]
    st.session_state.work_data["safety_measure_3"] = ai_result["safety_measure_3"]

    render_topbar("result", "Safety TBM", "result-app-title", back_page="input", help_key="result")


    # =========================
    # 메인 제목
    # =========================
    st.markdown("""
<div class="result-report-label">RISK ANALYSIS REPORT</div>
<div class="result-title">오늘의 작업 위험도는</div>
<div class="result-subtitle">
    입력한 작업유형과 취급물질 정보를 기준으로 오늘 작업의 위험수준을 분석했습니다.
</div>
""", unsafe_allow_html=True)

    # =========================
    # 위험도 점수 + 신호등 카드
    # =========================
    render_risk_summary_card(score)

    # =========================
    # AI 분석 주요 위험요인 / 안전대책 / 유사사고 카드
    # =========================
    render_ai_hazard_cards(risk_items)
    render_safety_measure_cards(measure_items)
    render_similar_accident_card(similar_accident_ai_text if similar_accident is not None else "")

    # =========================
    # 작업물질 정보
    # =========================
    if result.get("chem_name"):
        st.markdown(
            f"""
<div class="result-info-caption">
    작업물질: {result['chem_name']} / CHEMID: {result['chem_id']}
</div>
""",
            unsafe_allow_html=True
        )
    else:
        st.warning("작업물질 정보가 불확실하여 위험도를 보수적으로 산정했습니다.")

    # =========================
    # 체크리스트 이동 버튼
    # =========================
    if st.button("☑️ TBM 체크리스트 확인", use_container_width=True):
        st.session_state.page = "checklist"
        st.rerun()

    show_bottom_nav()


def show_task_info():
    """
    작업자가 오늘 작업을 선택하고 접속했을 때 보여주는 화면.
    위험도를 다시 계산하지 않고, 안전관리자가 "오늘 작업 입력"에서
    미리 분석해 둔 결과(work_tasks 테이블 값)를 그대로 보여준다.
    """
    work_data = st.session_state.get("work_data", {})

    # 작업모드 선택에서 작업자명/작업을 고르지 않고(예: 네비게이션이나 URL로) 곧바로
    # 들어온 경우, 빈 화면 대신 작업모드 선택 화면으로 돌려보낸다.
    if not (work_data.get("task_id") and work_data.get("작업자명")):
        _redirect_with_message("login", "먼저 작업자명과 오늘 작업을 선택해 주세요.")
        return

    render_topbar("task_info", "Safety TBM", "app-title", back_page="login", help_key="input")
    _flash_pending_message()

    result = st.session_state.get("result", {})

    work_name = work_data.get("작업명", "-")
    work_content = work_data.get("작업내용", "-")
    work_location = work_data.get("작업장소", "-")
    scheduled_time = work_data.get("예정시간", "-")
    tbm_place = work_data.get("TBM장소", "-")

    st.markdown(f"""
<div class="result-card">
    <div class="result-card-header">
        <div class="result-card-icon">📋</div>
        <div class="result-card-title">선택한 작업 정보</div>
    </div>
    <div class="message-text">
        <b>작업명</b> : {work_name}<br>
        <b>작업내용</b> : {work_content}<br>
        <b>작업장소</b> : {work_location}<br>
        <b>예정시간</b> : {scheduled_time}<br>
        <b>TBM 장소</b> : {tbm_place}
    </div>
</div>
""", unsafe_allow_html=True)

    score = result.get("score")

    if score is not None:
        st.markdown("""
<div class="result-report-label">RISK ANALYSIS REPORT</div>
<div class="result-title">오늘의 작업 위험도는</div>
<div class="result-subtitle">
    안전관리자가 사전에 분석한 오늘 작업의 위험수준입니다.
</div>
""", unsafe_allow_html=True)

        render_risk_summary_card(float(score))

        risk_items = [
            h for h in (
                work_data.get("main_hazard_1", ""),
                work_data.get("main_hazard_2", ""),
                work_data.get("main_hazard_3", "")
            ) if h
        ]
        measure_items = [
            m for m in (
                work_data.get("safety_measure_1", ""),
                work_data.get("safety_measure_2", ""),
                work_data.get("safety_measure_3", "")
            ) if m
        ]

        render_ai_hazard_cards(risk_items)
        render_safety_measure_cards(measure_items)
        render_similar_accident_card(result.get("similar_accident_text", ""))

        if result.get("chem_name"):
            st.markdown(
                f"""
<div class="result-info-caption">
    작업물질: {result['chem_name']} / CHEMID: {result.get('chem_id', '-')}
</div>
""",
                unsafe_allow_html=True
            )
    else:
        st.warning("안전관리자가 아직 이 작업의 위험도 분석을 등록하지 않았습니다. 안전관리자에게 확인해 주세요.")

    if st.button("☑️ TBM 체크리스트 확인", use_container_width=True):
        st.session_state.page = "checklist"
        st.rerun()

    show_bottom_nav()


def show_checklist():

    work_data = st.session_state.get("work_data", {})

    if not (work_data.get("task_id") and work_data.get("작업자명")):
        _redirect_with_message("login", "먼저 작업자명과 오늘 작업을 선택해 주세요.")
        return

    if st.session_state.get("show_checklist_success_modal"):
        show_checklist_success_popup()

    render_topbar("checklist", "Checklist", "checklist-app-title", back_page="task_info", help_key="checklist")
    _flash_pending_message()

    result = st.session_state.get("result", {})

    worker_name = work_data.get("작업자명", "미입력")
    work_name = work_data.get("작업명", "미입력")
    work_location = work_data.get("작업장소", "미입력")
    work_type = work_data.get("작업유형", "미입력")
    chemical = work_data.get("취급물질", "미입력")
    score = result.get("score", None)
    level = result.get("level", "미산정")

    client_dt = get_client_datetime()
    current_date = (client_dt or datetime.now()).strftime("%Y년 %m월 %d일")
    work_time = work_data.get("작업시간") or (client_dt or datetime.now()).strftime("%Y-%m-%d %H:%M")

    checklist_info_html = f'<div class="checklist-info-card"><div class="checklist-badge">DAILY INSPECTION</div><div class="checklist-title">작업 전 안전점검(TBM)</div><div class="checklist-subtitle">일시: {current_date}<br>점검자: {worker_name}<br>작업명: {work_name}<br>작업장소: {work_location}<br>작업유형: {work_type}</div></div>'
    st.markdown(checklist_info_html, unsafe_allow_html=True)

    checklist_section_html = '<div class="checklist-section-title">☑️ 점검 항목</div>'
    st.markdown(checklist_section_html, unsafe_allow_html=True)

    checklist_items = [
        "오늘 작업의 위험요인을 확인하였다.",
        "오늘 작업의 유사사고를 확인하였다.",
        "오늘 작업 중 사고발생 시 사용하는 보호구와 보호구 위치를 알고 있다.",
        "전날 과도한 음주를 하지 않았다.",
        "피로, 발열 없음 등 오늘 몸상태는 작업에 적절하다.",
        "위험요인, 불안전 발견 시 즉시 멈추고 생각한 후 작업하도록 주지했다.",
        "작업 후 정리정돈 방법을 알고 있다.",
        "사고발생 시 긴급연락처를 알고 있다.",
        "비상대피로 및 집결지를 확인하였다."
    ]

    checked_count = 0

    for idx, item in enumerate(checklist_items, start=1):
        checked = st.checkbox(
            item,
            key=f"tbm_check_{idx}"
        )

        if checked:
            checked_count += 1

    unchecked_items = [
        item for idx, item in enumerate(checklist_items, start=1)
        if not st.session_state.get(f"tbm_check_{idx}", False)
    ]

    all_checked = len(unchecked_items) == 0

    st.markdown(
        f"""
<div class="checklist-save-note">
    현재 {len(checklist_items)}개 항목 중 {checked_count}개 항목을 확인했습니다.
</div>
""",
        unsafe_allow_html=True
    )

    if not all_checked:
        st.warning("모든 점검 항목을 체크해야 TBM을 완료할 수 있습니다.")

    # =========================
    # 특이사항 입력
    # =========================
    st.markdown("""
<div class="checklist-section-title">📝 작업 전 일일 안전점검 시행 결과</div>
""", unsafe_allow_html=True)

    remark = st.text_area(
        "작업 전 일일 안전점검 시행 결과",
        placeholder="작업 전 일일 안전점검 시행 결과를 작성하세요(※ TBM리더 작성 필수).",
        height=140,
        key="checklist_remark",
        label_visibility="collapsed"
    )

    # =========================
    # 서명 입력
    # =========================
    st.markdown("""
<div class="checklist-section-title">✍️ 서명</div>
""", unsafe_allow_html=True)

    st.caption("터치 또는 마우스로 아래 칸에 직접 서명해 주세요.")

    signature_reset_count = st.session_state.get("signature_reset_count", 0)

    canvas_result = st_canvas(
        fill_color="rgba(0, 0, 0, 0)",
        stroke_width=3,
        stroke_color="#111111",
        background_color="",
        height=150,
        width=320,
        drawing_mode="freedraw",
        display_toolbar=False,
        key=f"signature_canvas_{signature_reset_count}",
    )

    if st.button("🧹 지우기 (다시 서명)"):
        st.session_state.signature_reset_count = signature_reset_count + 1
        st.session_state.pop("signature_confirmed_png", None)
        st.session_state.pop("signature_confirmed_reset_count", None)
        st.rerun()

    def _is_signature_drawn(result):
        try:
            image_data = getattr(result, "image_data", None)
            if image_data is None:
                return False
            return bool(np.any(np.asarray(image_data)[:, :, 3] > 0))
        except Exception:
            return False

    # streamlit-drawable-canvas 컴포넌트는 다른 위젯(체크박스, 텍스트 입력 등)이 먼저
    # 리런을 유발한 경우, 방금 그린 서명 결과가 세션에 아직 반영되기 전이라 image_data가
    # None이거나 canvas_result 자체가 CanvasResult 클래스로 반환될 수 있다(브라우저→세션
    # 반영 타이밍 문제). 이 경우 실제로는 서명을 그렸는데도 "없음"으로 잘못 판단하게 된다.
    # 이를 막기 위해, 이번 리런에서 서명이 정상적으로 감지되면 즉시 PNG로 인코딩해
    # session_state에 "확정" 저장해두고, 같은 캔버스(같은 reset_count)인 동안에는 이후
    # 리런에서 캔버스 값이 일시적으로 비어 보이더라도 이 확정 데이터를 신뢰한다.
    # "지우기(다시 서명)"를 누를 때만 확정 상태를 초기화한다.
    live_signature_drawn = _is_signature_drawn(canvas_result)

    if live_signature_drawn:
        try:
            _live_buf = BytesIO()
            Image.fromarray(
                np.asarray(canvas_result.image_data).astype("uint8"), mode="RGBA"
            ).save(_live_buf, format="PNG")
            st.session_state.signature_confirmed_png = _live_buf.getvalue()
            st.session_state.signature_confirmed_reset_count = signature_reset_count
        except Exception:
            pass

    signature_confirmed_png = (
        st.session_state.get("signature_confirmed_png")
        if st.session_state.get("signature_confirmed_reset_count") == signature_reset_count
        else None
    )
    signature_drawn = live_signature_drawn or signature_confirmed_png is not None

    if not signature_drawn:
        st.warning("서명을 입력해야 TBM을 완료할 수 있습니다.")

    # =========================
    # 저장 버튼
    # =========================
    if st.button("💾 TBM 완료 및 저장", use_container_width=True):
        unchecked_items = [
            item for idx, item in enumerate(checklist_items, start=1)
            if not st.session_state.get(f"tbm_check_{idx}", False)
        ]

        if unchecked_items:
            st.error("모든 체크리스트 항목을 확인해야 TBM을 완료할 수 있습니다.")
            st.write("미확인 항목:")
            for item in unchecked_items:
                st.write(f"- {item}")
            return

        if not signature_drawn:
            st.error("서명을 입력해야 TBM을 완료할 수 있습니다.")
            return

        # 확정 저장된 서명 PNG가 있으면 그것을 우선 사용한다(제출 버튼 클릭으로 인한
        # 리런에서 canvas_result 값이 일시적으로 비어 보이는 경우를 대비).
        signature_png_bytes = signature_confirmed_png
        if signature_png_bytes is None:
            try:
                signature_buf = BytesIO()
                Image.fromarray(
                    np.asarray(canvas_result.image_data).astype("uint8"), mode="RGBA"
                ).save(signature_buf, format="PNG")
                signature_png_bytes = signature_buf.getvalue()
            except Exception:
                st.error("서명 이미지를 처리하지 못했습니다. 서명을 다시 입력한 후 제출해 주세요.")
                return
        st.session_state.signature_png_bytes = signature_png_bytes

        st.session_state.checklist_data = {
            "저장시간": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "작업자명": worker_name,
            "작업명": work_name,
            "작업장소": work_location,
            "작업유형": work_type,
            "전체항목수": len(checklist_items),
            "확인항목수": checked_count,
            "미확인항목": [],
            "특이사항": remark,
            "TBM완료여부": "완료"
        }

        task_id = work_data.get("task_id", "")
        team_id = st.session_state.get("team_id", "")

        # 서명과 TBM 1차 정보(체크리스트 결과)를 작업일지 제출을 기다리지 않고 바로 반영한다.
        # work_logs는 (task_id, worker_name) 기준으로 upsert되어, 이후 작업일지 제출 시
        # 같은 행에 종료 미팅 내용만 추가되고 새 행이 생기지 않는다.
        try:
            upload_signature(task_id, team_id, worker_name, signature_png_bytes)

            upsert_work_log({
                "team_id": team_id,
                "team_name": st.session_state.get("team_name", ""),
                "task_id": task_id,
                "worker_name": worker_name,
                "work_name": work_name,
                "work_location": work_location,
                "work_time": work_time,
                "work_type": work_type,
                "material_name": chemical,
                "risk_score": float(score) if score is not None else None,
                "risk_level": level,
                "tbm_status": "완료",
                "daily_safety_check_result": remark,
                "is_tbm_leader": work_data.get("TBM리더여부", False),
                "main_hazard_1": work_data.get("main_hazard_1", ""),
                "main_hazard_2": work_data.get("main_hazard_2", ""),
                "main_hazard_3": work_data.get("main_hazard_3", ""),
                "safety_measure_1": work_data.get("safety_measure_1", ""),
                "safety_measure_2": work_data.get("safety_measure_2", ""),
                "safety_measure_3": work_data.get("safety_measure_3", ""),
                "leader_department": work_data.get("리더소속", ""),
                "leader_position": work_data.get("리더직책", ""),
                "leader_name": work_data.get("리더성명", ""),
                "tbm_start_time": work_data.get("TBM시작시간", ""),
                "submit_status": "TBM완료",
            }, task_id, worker_name)
        except Exception as e:
            st.error("서명/TBM 정보 저장 중 오류가 발생했습니다.")
            st.write(str(e))

        # 이 작업(task)의 TBM 중 가장 먼저 제출된 건이면 work_tasks.first_tbm_submitted_at에 기록한다.
        # 작업일지 제출을 기다리지 않고 체크리스트 제출("TBM 완료 및 저장") 시점에 바로 기록해야
        # 그 직후 회의록(Word)을 생성해도 TBM 일시가 정상적으로 표시된다.
        # IS NULL 조건으로 갱신하므로, 이미 값이 기록되어 있으면(=다른 작업자가 먼저 제출) 덮어쓰지 않는다.
        checklist_submit_time_str = (client_dt or datetime.now()).strftime("%Y-%m-%d %H:%M:%S")
        if task_id:
            try:
                supabase.table("work_tasks").update(
                    {"first_tbm_submitted_at": checklist_submit_time_str}
                ).eq("id", task_id).is_(
                    "first_tbm_submitted_at", "null"
                ).execute()
            except Exception as e:
                print("first_tbm_submitted_at 업데이트 오류:", e)

        # 작업모드 선택 화면으로 돌아갔을 때 방금 사용한 작업자명/작업이 그대로
        # 선택돼 있도록, 로그인 화면 위젯의 session_state 값을 미리 채워 둔다.
        # (텍스트 입력 위젯은 이 페이지를 벗어났다 오면 값이 비워지므로, 자동으로
        # 유지되지 않는다 — 위젯 재생성 전에 값을 직접 지정해야 한다.)
        st.session_state.login_mode_radio = "작업자"
        st.session_state.login_worker_name = worker_name
        st.session_state.login_selected_task = f"{work_name} / {work_data.get('예정시간', '-')}"
        if work_data.get("TBM리더여부", False):
            st.session_state.is_tbm_leader = True
            st.session_state.leader_department = work_data.get("리더소속", "")
            st.session_state.leader_position = work_data.get("리더직책", "")

        st.session_state.show_checklist_success_modal = True
        st.rerun()

    show_bottom_nav()

def _close_checklist_success_modal():
    st.session_state.show_checklist_success_modal = False


@st.dialog("TBM 완료", on_dismiss=_close_checklist_success_modal)
def show_checklist_success_popup():
    st.write("TBM이 완료되었습니다.")

    if st.button("확인", use_container_width=True):
        st.session_state.show_checklist_success_modal = False
        st.session_state.page = "login"
        st.query_params.clear()
        st.rerun()


def _close_journal_success_modal():
    st.session_state.show_journal_success_modal = False


@st.dialog("제출 완료", on_dismiss=_close_journal_success_modal)
def show_journal_success_popup():
    st.write("작업일지가 제출되었습니다.")

    if st.button("확인", use_container_width=True):
        st.session_state.show_journal_success_modal = False
        st.session_state.page = "login"
        st.query_params.clear()
        st.rerun()


def show_journal():

    work_data = st.session_state.get("work_data", {})

    if not (work_data.get("task_id") and work_data.get("작업자명")):
        _redirect_with_message("login", "먼저 작업자명과 오늘 작업을 선택해 주세요.")
        return

    if st.session_state.get("show_journal_success_modal"):
        show_journal_success_popup()

    render_topbar("journal", "Safety TBM", "journal-app-title", back_page="login", help_key="journal")
    _flash_pending_message()

    # 작업자 기기(브라우저)의 로컬 시각. "제출하기" 클릭 시 TBM 제출 시각으로 기록한다.
    client_dt = get_client_datetime()

    # =========================
    # 데이터 불러오기
    # =========================
    result = st.session_state.get("result", {})
    checklist_data = st.session_state.get("checklist_data", {})

    worker_name = work_data.get("작업자명", "미입력")
    work_name = work_data.get("작업명", "미입력")
    work_location = work_data.get("작업장소", "미입력")
    work_type = work_data.get("작업유형", "미입력")
    chemical = work_data.get("취급물질", "미입력")
    work_time = work_data.get("작업시간") or (client_dt or datetime.now()).strftime("%Y-%m-%d %H:%M")

    task_id = work_data.get("task_id", "")

    # 체크리스트 완료("TBM 완료 및 저장") 시점에 이미 저장된 work_logs 행을 DB에서
    # 직접 조회해, 화면 표시와 재제출 값의 기준으로 삼는다. st.session_state.checklist_data는
    # 작업자 전체가 공유하는 세션 값이라, "작업로그 작성하기"로 이 화면에 들어오기 전
    # 다른 작업자/작업으로 바꿨다면 최신 값이 아닐 수 있으므로 신뢰하지 않는다.
    existing_log = None
    if task_id and worker_name and worker_name != "미입력":
        try:
            existing_result = (
                supabase.table("work_logs")
                .select("*")
                .eq("task_id", task_id)
                .eq("worker_name", worker_name)
                .limit(1)
                .execute()
            )
            if existing_result.data:
                existing_log = existing_result.data[0]
        except Exception as e:
            print("기존 work_logs 조회 오류:", e)

    # TBM 체크리스트를 아직 완료하지 않은 작업/작업자에 대해 작업일지 화면으로
    # 곧바로 건너뛰려는 시도(네비게이션, URL 직접 진입 등)는 막고 체크리스트로 되돌린다.
    # 단, 방금 제출을 마친 직후(성공 팝업이 뜬 상태)는 이미 완료된 상태이므로 막지 않는다.
    if existing_log is None and not st.session_state.get("show_journal_success_modal"):
        _redirect_with_message("checklist", "TBM 체크리스트를 먼저 완료해야 작업일지를 작성할 수 있습니다.")
        return

    if existing_log:
        score = existing_log.get("risk_score")
        level = existing_log.get("risk_level") or "미산정"
        checklist_remark = existing_log.get("daily_safety_check_result") or ""
    else:
        score = result.get("score", None)
        level = result.get("level", "미산정")
        checklist_remark = checklist_data.get("특이사항", "")

    today_text = (client_dt or datetime.now()).strftime("%Y년 %m월 %d일")
    score_text = f"{float(score):.0f}점" if score is not None else "미산정"

    # =========================
    # 제목
    # =========================
    st.markdown(f"""
<div class="journal-title">작업일지 작성</div>
<div class="journal-subtitle">
    {today_text} | TBM 완료 후 작업 내용을 기록합니다.
</div>
""", unsafe_allow_html=True)

    # =========================
    # TBM 요약 정보
    # =========================
    checklist_remark_text = checklist_remark if checklist_remark else "체크리스트 단계에서 입력된 특이사항이 없습니다."

    summary_html = f"""
<div class="journal-card"><div class="journal-card-header"><div class="journal-card-title">TBM 요약 정보</div><div style="font-size:22px; color:#2170e4;">ℹ️</div></div><div class="journal-summary-row"><div class="journal-summary-label">작업자</div><div class="journal-summary-value">{worker_name}</div></div><div class="journal-summary-row"><div class="journal-summary-label">작업명</div><div class="journal-summary-value">{work_name}</div></div><div class="journal-summary-row"><div class="journal-summary-label">작업장소</div><div class="journal-summary-value">{work_location}</div></div><div class="journal-summary-row"><div class="journal-summary-label">작업유형</div><div class="journal-summary-value">{work_type}</div></div><div class="journal-summary-row"><div class="journal-summary-label">취급물질</div><div class="journal-summary-value">{chemical}</div></div><div class="journal-summary-row"><div class="journal-summary-label">위험도</div><div class="journal-summary-value"><span class="journal-risk-badge">{score_text} / {level}</span></div></div><div class="journal-message-box"><b>TBM 체크리스트 특이사항</b><br>{checklist_remark_text}</div></div>
"""

    st.markdown(summary_html, unsafe_allow_html=True)


    # =========================
    # 특이사항 기록 내역
    # =========================
    st.markdown("""
<div class="journal-card">
    <div class="journal-card-header">
        <div class="journal-card-title">작업 전 일일 안전점검 시행 결과</div>
    </div>
    <div class="journal-small-label">작업 후 종료 미팅</div>
</div>
""", unsafe_allow_html=True)

    journal_note = st.text_area(
        "작업 후 종료 미팅",
        placeholder="작업 후 중점위험요인 확인내역을 작성하세요.",
        height=160,
        key="journal_note_input",
        label_visibility="collapsed"
    )

    st.markdown("""
<div class="journal-message-box">
    ✅ 종료 미팅 내역은 안전관리자 작업보드에 기록될 예정입니다.
</div>
""", unsafe_allow_html=True)


    # =========================
    # 제출 카드
    # =========================
    st.markdown("""
<div class="journal-submit-card">
    <div class="journal-submit-title">일지 작성을 완료하시겠습니까?</div>
    <div class="journal-submit-desc">
        작성된 내용은 추후 안전관리자 작업로그 화면에 반영되도록 연결할 수 있습니다.
    </div>
</div>
""", unsafe_allow_html=True)

    if st.button("📤 제출하기", use_container_width=True):

        # 작업자 기기의 로컬 시각(get_client_datetime)을 우선 사용하고,
        # 아직 브라우저 값이 도착하지 않았으면 서버 시각으로 대체한다.
        submit_dt = client_dt or datetime.now()
        submit_time_str = submit_dt.strftime("%Y-%m-%d %H:%M:%S")

        # 체크리스트 단계에서 이미 저장된 값(existing_log)을 우선 신뢰하고,
        # 없는 경우에만 세션 값으로 대체한다.
        base_fields = existing_log or {}
        checklist_is_tbm_leader = base_fields.get("is_tbm_leader", st.session_state.work_data.get("TBM리더여부", False))

        if checklist_is_tbm_leader:
            st.session_state.work_data["TBM종료시간"] = submit_time_str

        # first_tbm_submitted_at은 이제 체크리스트 제출("TBM 완료 및 저장") 시점에
        # show_checklist()에서 기록한다 — 작업일지 제출과 독립적으로 동작해야 하므로
        # 여기서는 별도로 갱신하지 않는다.

        log_data = {
            "팀ID": st.session_state.get("team_id", ""),
            "팀명": st.session_state.get("team_name", ""),
            "작업자": worker_name,
            "작업명": work_name,
            "작업장소": work_location,
            "작업유형": work_type,
            "취급물질": chemical,
            "작업시간": work_time,
            "위험도점수": score_text,
            "위험등급": level,
            "TBM완료여부": "완료",

            "작업전일일안전점검결과": checklist_remark,

            "작업후종료미팅": journal_note,

            "TBM리더여부": st.session_state.work_data.get("TBM리더여부", False),
            "리더소속": st.session_state.work_data.get("리더소속", ""),
            "리더직책": st.session_state.work_data.get("리더직책", ""),
            "리더성명": st.session_state.work_data.get("리더성명", ""),

            "task_id": st.session_state.work_data.get("task_id", ""),

            "제출시간": submit_time_str,
            "제출상태": "제출완료"
        }

        st.session_state.journal_data = log_data
        st.session_state.work_logs.append(log_data)

        log_df = pd.DataFrame(st.session_state.work_logs)
        log_df.to_csv(
            st.session_state.work_log_csv_path,
            index=False,
            encoding="utf-8-sig"
        )

        try:
            # 체크리스트 완료("TBM 완료 및 저장") 시점에 이미 만들어 둔 work_logs 행에
            # 작업일지 정보(종료 미팅, 종료 시간, 최종 제출 상태)를 이어서 채운다.
            # (task_id, worker_name) 기준 upsert이므로 행이 2건으로 늘지 않는다.
            upsert_work_log({
                "team_id": st.session_state.get("team_id", ""),
                "team_name": st.session_state.get("team_name", ""),
                "task_id": st.session_state.work_data.get("task_id", ""),
                "worker_name": worker_name,
                "work_name": work_name,
                "work_location": work_location,
                "work_time": work_time,
                "work_type": st.session_state.work_data.get("작업유형", ""),
                "material_name": chemical,
                "risk_score": float(score) if score is not None else None,
                "risk_level": level,
                "tbm_status": "완료",
                "daily_safety_check_result": checklist_remark,
                "closing_meeting_result": journal_note,
                "is_tbm_leader": checklist_is_tbm_leader,
                "main_hazard_1": base_fields.get("main_hazard_1") or st.session_state.work_data.get("main_hazard_1", ""),
                "main_hazard_2": base_fields.get("main_hazard_2") or st.session_state.work_data.get("main_hazard_2", ""),
                "main_hazard_3": base_fields.get("main_hazard_3") or st.session_state.work_data.get("main_hazard_3", ""),
                "safety_measure_1": base_fields.get("safety_measure_1") or st.session_state.work_data.get("safety_measure_1", ""),
                "safety_measure_2": base_fields.get("safety_measure_2") or st.session_state.work_data.get("safety_measure_2", ""),
                "safety_measure_3": base_fields.get("safety_measure_3") or st.session_state.work_data.get("safety_measure_3", ""),
                "leader_department": base_fields.get("leader_department") or st.session_state.work_data.get("리더소속", ""),
                "leader_position": base_fields.get("leader_position") or st.session_state.work_data.get("리더직책", ""),
                "leader_name": base_fields.get("leader_name") or st.session_state.work_data.get("리더성명", ""),
                "tbm_start_time": base_fields.get("tbm_start_time") or st.session_state.work_data.get("TBM시작시간", ""),
                "tbm_end_time": st.session_state.work_data.get("TBM종료시간", ""),
                "submit_status": "제출완료",
            }, task_id, worker_name)

        except Exception as e:

            st.error("작업로그 DB 저장 실패")
            st.write(str(e))

        st.session_state.journal_submitted = True

        st.session_state.show_journal_success_modal = True
        st.rerun()

    show_bottom_nav()

def show_task_create():

    if not st.session_state.get("team_id"):
        st.warning("팀 접속 정보가 없습니다. 작업팀 접속 화면에서 다시 접속해 주세요.")

        if st.button("작업팀 접속 화면으로 이동", use_container_width=True):
            st.session_state.page = "team_access"
            st.query_params.clear()
            st.rerun()

        return

    if st.session_state.get("mode") != "안전관리자":
        st.warning("안전관리자 인증이 필요합니다. 작업모드 선택 화면에서 안전관리자로 접속해 주세요.")

        if st.button("작업모드 선택 화면으로 이동", use_container_width=True):
            st.session_state.page = "login"
            st.rerun()

        return

    render_topbar("task_create", "오늘 작업 입력", "manager-title", back_page="manager", help_key="manager")

    # 안전관리자 기기(브라우저)의 현재 접속 시각. 위험도 계산용 시간(작업 시작 시간)에는
    # 쓰지 않고, 유사사고사례를 월(月) 기준으로 찾을 때만 사용한다.
    client_dt = get_client_datetime()

    workers = get_team_workers(st.session_state.get("team_id", ""))

    work_type_options = [
        "정기작업",
        "비작업(순찰·경비)",
        "유지보수",
        "화기작업",
        "시운전·정지",
        "세척작업"
    ]

    work_type_map = {
        "정기작업": "ROUTINE",
        "비작업(순찰·경비)": "IDLE",
        "유지보수": "MAINTENANCE",
        "화기작업": "HOT_WORK",
        "시운전·정지": "STARTUP_SHUTDOWN",
        "세척작업": "CLEANING"
    }

    # =========================
    # 저장된 작업 템플릿 불러오기 / 삭제
    # =========================
    templates = get_work_templates(st.session_state.get("team_id", ""))

    if templates:
        st.markdown('<div class="manager-section-title">저장된 작업 불러오기</div>', unsafe_allow_html=True)

        template_options = {t["template_name"]: t for t in templates}

        tc_selected_template_name = st.selectbox(
            "불러올 템플릿 선택",
            list(template_options.keys()),
            key="tc_template_select"
        )

        with st.container(key="inline_pair_tc_template_btns"):
            tc_load_col, tc_delete_col = st.columns(2)

            with tc_load_col:
                tc_load_clicked = st.button("📥 불러오기", key="tc_load_template_btn", use_container_width=True)

            with tc_delete_col:
                tc_delete_clicked = st.button("🗑️ 삭제", key="tc_delete_template_btn", use_container_width=True)

        if tc_load_clicked:
            _template = template_options[tc_selected_template_name]

            st.session_state["tc_work_name"] = _template.get("work_name") or ""
            st.session_state["tc_work_content"] = _template.get("work_content") or ""
            st.session_state["tc_work_location"] = _template.get("work_location") or ""
            st.session_state["tc_tbm_place"] = _template.get("tbm_place") or ""

            if _template.get("work_type_display") in work_type_options:
                st.session_state["tc_work_type"] = _template.get("work_type_display")

            # streamlit_searchbox는 key가 이미 session_state에 있으면 default를 무시하므로,
            # key를 지워서 다음 렌더링에서 default_searchterm으로 다시 초기화되게 한다.
            st.session_state.pop("tc_chemical_searchbox", None)
            st.session_state["tc_chem_pending_default"] = _template.get("material_name") or ""

            st.success(f"'{tc_selected_template_name}' 템플릿을 불러왔습니다. 작업 시간을 입력해 주세요.")
            st.rerun()

        if tc_delete_clicked:
            delete_work_template(template_options[tc_selected_template_name]["id"])
            st.success(f"'{tc_selected_template_name}' 템플릿을 삭제했습니다.")
            st.rerun()

    tc_register_tomorrow = st.checkbox(
        "내일 작업을 미리 등록합니다",
        key="tc_register_tomorrow"
    )

    if tc_register_tomorrow:
        st.caption("📅 이 작업은 '내일 예정된 작업'으로 등록되며, 날짜가 바뀌면 자동으로 '오늘 작업'이 됩니다.")

    st.markdown('<div class="manager-section-title">작업 기본정보</div>', unsafe_allow_html=True)

    tc_work_name = st.text_input(
        "작업명",
        placeholder="예: 배관 플랜지 교체 작업",
        key="tc_work_name"
    )

    tc_work_content = st.text_area(
        "작업내용",
        placeholder="예: 노후 배관 플랜지 분리, 가스켓 교체 및 체결 작업",
        height=100,
        key="tc_work_content"
    )

    tc_work_location = st.text_input(
        "작업장소",
        placeholder="예: 2공장 반응기실",
        key="tc_work_location"
    )

    tc_tbm_place = st.text_input(
        "TBM 장소",
        placeholder="예: 현장 작업구역 앞",
        key="tc_tbm_place"
    )

    tc_work_type_display = st.selectbox(
        "작업 유형",
        work_type_options,
        key="tc_work_type"
    )

    st.markdown('<div class="field-label">작업 시작 시간</div>', unsafe_allow_html=True)

    _now_local = datetime.now()

    # number_input에 min_value/max_value를 걸면, 범위를 벗어난 값을 키보드로 직접 타이핑해
    # Enter를 눌렀을 때 위젯이 그 값을 서버에 아예 반영하지 않고 조용히 무시해버려서(직전의
    # 유효했던 값을 그대로 유지) 화면에는 사용자가 입력한 범위 밖 숫자가 그대로 남아있는데
    # 정작 서버는 다른 값을 들고 있는 상태가 되어버린다. 그래서 위젯 자체에는 범위 제한을
    # 걸지 않고, 위젯을 만들기 "전"에 직전 rerun에서 커밋된 session_state 값을 미리
    # 0~23 / 0~59로 보정해 둔다(위젯이 이미 만들어진 뒤에는 같은 key의 session_state를
    # 다시 대입할 수 없어 이 순서가 중요하다).
    _time_out_of_range = False

    if "tc_start_hour_input" in st.session_state:
        _raw_hour = st.session_state["tc_start_hour_input"]
        if not (0 <= _raw_hour <= 23):
            st.session_state["tc_start_hour_input"] = max(0, min(23, _raw_hour))
            _time_out_of_range = True

    if "tc_start_minute_input" in st.session_state:
        _raw_minute = st.session_state["tc_start_minute_input"]
        if not (0 <= _raw_minute <= 59):
            st.session_state["tc_start_minute_input"] = max(0, min(59, _raw_minute))
            _time_out_of_range = True

    if _time_out_of_range:
        st.warning("작업 시작 시간은 시 0~23, 분 0~59 범위로 입력해 주세요. 값을 범위 안으로 보정했습니다.")

    with st.container(key="tc_time_input_row"):
        tc_hour_col, tc_minute_col = st.columns(2)

        with tc_hour_col:
            tc_start_hour = st.number_input(
                "시",
                step=1,
                value=_now_local.hour,
                key="tc_start_hour_input",
            )

        with tc_minute_col:
            tc_start_minute = st.number_input(
                "분",
                step=1,
                value=_now_local.minute,
                key="tc_start_minute_input",
            )

    # 모바일에서 이 입력창을 탭했을 때 문자 키패드가 아니라 숫자 키패드가 바로 뜨도록,
    # number_input이 만든 실제 <input>에 inputmode="numeric"을 부여한다. Streamlit에는
    # 이 속성을 직접 지정하는 옵션이 없어, 휠 피커 시절과 같은 방식(streamlit_js_eval로
    # parent 문서를 조작)을 재사용한다. 이 위젯은 값이 바뀔 때마다 React가 리렌더링되며
    # inputmode 속성을 매번 지워버리므로, 한 번 스크립트가 실행된 뒤에도 계속 유지되도록
    # MutationObserver로 속성이 지워지는 즉시 다시 채워 넣는다(스크롤 리스너와 동일하게
    # 최초 1회만 리스너를 붙이고, 이후에는 리스너 자체가 계속 동작한다). JS 표현식 문자열은
    # 매 rerun 동일하게 유지해 재평가·재전송이 반복되지 않도록 한다.
    streamlit_js_eval(
        js_expressions="""
        (function() {
            function initFor(k) {
                var el = parent.document.querySelector('.st-key-' + k + ' input');
                if (!el) {
                    setTimeout(function() { initFor(k); }, 200);
                    return;
                }
                if (el.dataset.numericKeypadInit) return;
                el.dataset.numericKeypadInit = '1';
                function apply() {
                    if (el.getAttribute('inputmode') !== 'numeric') el.setAttribute('inputmode', 'numeric');
                    if (el.getAttribute('pattern') !== '[0-9]*') el.setAttribute('pattern', '[0-9]*');
                }
                apply();
                new MutationObserver(apply).observe(el, { attributes: true, attributeFilter: ['inputmode', 'pattern'] });
            }
            ['tc_start_hour_input', 'tc_start_minute_input'].forEach(initFor);
            return 'ok';
        })()
        """,
        key="tc_time_numeric_keypad",
    )

    tc_start_time = dt_time(tc_start_hour, tc_start_minute)

    # 템플릿을 방금 불러온 경우, streamlit_searchbox는 key가 이미 session_state에
    # 있으면 default를 무시하므로 key를 지우고 default_searchterm으로 재초기화한다.
    tc_chem_pending_default = st.session_state.pop("tc_chem_pending_default", None)

    if tc_chem_pending_default is not None:
        tc_chemical = st_searchbox(
            search_chemical_candidates,
            key="tc_chemical_searchbox",
            placeholder="화학물질명을 입력하세요. 예: 황산",
            clear_on_submit=False,
            default=tc_chem_pending_default,
            default_searchterm=tc_chem_pending_default,
            default_use_searchterm=True,
        )
    else:
        tc_chemical = st_searchbox(
            search_chemical_candidates,
            key="tc_chemical_searchbox",
            placeholder="화학물질명을 입력하세요. 예: 황산",
            clear_on_submit=False,
        )

    tc_risk_assessment_done = st.checkbox(
        "위험성평가 실시",
        value=True,
        key="tc_risk_assessment_done"
    )

    tc_assigned_workers = st.multiselect(
        "작업자 선택",
        workers,
        key="tc_assigned_workers"
    )

    # =========================
    # 작업을 템플릿으로 저장
    # =========================
    tc_save_as_template = st.checkbox("이 작업을 템플릿으로 저장", key="tc_save_as_template")

    if tc_save_as_template:
        tc_template_name = st.text_input(
            "템플릿 이름",
            placeholder="예: 정기 배관 점검",
            key="tc_template_name"
        )

        if st.button("💾 템플릿 저장", key="tc_save_template_btn", use_container_width=True):

            if not tc_template_name.strip():
                st.warning("템플릿 이름을 입력해 주세요.")
            elif not tc_work_name.strip():
                st.warning("작업명을 입력해 주세요.")
            else:
                create_work_template(
                    team_id=st.session_state.get("team_id", ""),
                    team_name=st.session_state.get("team_name", ""),
                    template_name=tc_template_name.strip(),
                    work_name=tc_work_name.strip(),
                    work_content=tc_work_content.strip(),
                    work_location=tc_work_location.strip(),
                    tbm_place=tc_tbm_place.strip(),
                    work_type_code=work_type_map[tc_work_type_display],
                    work_type_display=tc_work_type_display,
                    material_name=str(tc_chemical).strip() if tc_chemical else ""
                )
                st.success(f"'{tc_template_name.strip()}' 템플릿이 저장되었습니다.")
                st.session_state.pop("tc_template_name", None)
                # 체크박스는 이미 이 실행에서 인스턴스화됐으므로 값 대입이 아니라
                # 키를 지워서 다음 렌더링에서 기본값(False)으로 재초기화되게 한다.
                st.session_state.pop("tc_save_as_template", None)
                st.rerun()

    if st.button("⚠️ 위험도 분석하기", key="tc_analyze_btn", use_container_width=True):

        if not tc_work_name.strip():
            st.warning("작업명을 입력해 주세요.")
            return

        if not tc_work_content.strip():
            st.warning("작업내용을 입력해 주세요.")
            return

        if not tc_work_location.strip():
            st.warning("작업장소를 입력해 주세요.")
            return

        if not tc_chemical or not str(tc_chemical).strip():
            st.warning("취급물질을 입력하고 검색 결과에서 물질을 선택해 주세요.")
            return

        chemical_resolved = resolve_chemical_alias(str(tc_chemical).strip())
        work_type_code = work_type_map[tc_work_type_display]

        # 자동 시간 감지(get_client_datetime) 대신, 안전관리자가 직접 입력한
        # 작업 시작 시간을 위험도 계산에 그대로 사용한다.
        start_dt = datetime.combine(datetime.now().date(), tc_start_time)
        time_slot, _ = get_current_time_slot(start_dt)

        try:
            result, err = run_risk_scoring(chemical_resolved, work_type_code, time_slot)

            if err:
                st.error(err)
                return

            # get_risk_and_measure_messages / find_similar_accident가 작업유형으로
            # 매칭하므로 결과 dict에 넣어준다 (이게 빠져 있어서 유사사고사례가
            # 항상 "없음"으로 나왔던 원인 중 하나였다).
            result["작업유형"] = tc_work_type_display
            # get_archive_reference_cases가 작업명 키워드로 아카이브를 검색하고,
            # generate_ai_text가 [이번 작업 정보]에 원문 그대로 넣어주므로 채워준다.
            result["작업명"] = tc_work_name.strip()
            result["작업내용"] = tc_work_content.strip()

            # 유사사고사례의 "월" 기준은 작업 시작시간이 아니라 안전관리자
            # 기기의 현재 접속 시각으로 판정한다.
            current_month = client_dt.month if client_dt else datetime.now().month

            with st.spinner("AI가 TBM 내용을 생성 중입니다..."):
                ai_result = run_ai_hazard_analysis(result, current_month=current_month)

            # 새로 분석하면 이전 분석에서 골라뒀던 중점위험요인/대책 선택은 초기화한다.
            st.session_state.pop("tc_selected_hazard_idx", None)
            st.session_state.pop("tc_selected_measure_idx", None)

            st.session_state.task_analysis = {
                "work_name": tc_work_name.strip(),
                "work_content": tc_work_content.strip(),
                "work_location": tc_work_location.strip(),
                "tbm_place": tc_tbm_place.strip(),
                "risk_assessment_done": tc_risk_assessment_done,
                "work_type_code": work_type_code,
                "work_type_display": tc_work_type_display,
                "start_time": start_dt.strftime("%Y-%m-%d %H:%M:%S"),
                "scheduled_time": tc_start_time.strftime("%H:%M"),
                "time_slot": time_slot,
                "material_name": chemical_resolved,
                "chem_id": result.get("chem_id", ""),
                "chem_name": result.get("chem_name", ""),
                "risk_score": result["score"],
                "risk_level": result["level"],
                "main_hazard_1": ai_result["main_hazard_1"],
                "main_hazard_2": ai_result["main_hazard_2"],
                "main_hazard_3": ai_result["main_hazard_3"],
                "safety_measure_1": ai_result["safety_measure_1"],
                "safety_measure_2": ai_result["safety_measure_2"],
                "safety_measure_3": ai_result["safety_measure_3"],
                "similar_accident_text": ai_result["similar_accident_ai_text"] if ai_result["similar_accident"] is not None else "",
            }

        except Exception as e:
            st.error("위험도 분석 중 오류가 발생했습니다.")
            st.exception(e)

    # =========================
    # 분석 결과 + 중점위험요인 선택
    # =========================
    analysis = st.session_state.get("task_analysis")

    if analysis:
        st.markdown('<div class="manager-section-title">위험도 분석 결과</div>', unsafe_allow_html=True)

        render_risk_summary_card(float(analysis["risk_score"]))

        risk_items = [
            h for h in (analysis["main_hazard_1"], analysis["main_hazard_2"], analysis["main_hazard_3"])
            if h
        ]
        measure_items = [
            m for m in (analysis["safety_measure_1"], analysis["safety_measure_2"], analysis["safety_measure_3"])
            if m
        ]

        if risk_items:
            render_selectable_hazard_cards(risk_items, "tc_selected_hazard_idx")
        else:
            render_ai_hazard_cards(risk_items)

        if measure_items:
            render_selectable_measure_cards(measure_items, "tc_selected_measure_idx")
        else:
            render_safety_measure_cards(measure_items)

        render_similar_accident_card(analysis["similar_accident_text"])

        if analysis.get("chem_name"):
            st.markdown(
                f"""
<div class="result-info-caption">
    작업물질: {analysis['chem_name']} / CHEMID: {analysis['chem_id']}
</div>
""",
                unsafe_allow_html=True
            )

        hazard_idx = st.session_state.get("tc_selected_hazard_idx")
        measure_idx = st.session_state.get("tc_selected_measure_idx")

        selected_hazard = risk_items[hazard_idx] if hazard_idx is not None and hazard_idx < len(risk_items) else ""
        selected_measure = measure_items[measure_idx] if measure_idx is not None and measure_idx < len(measure_items) else ""

        create_btn_label = "➕ 내일 작업 등록" if tc_register_tomorrow else "➕ 오늘 작업 등록"

        if st.button(create_btn_label, key="create_work_task_btn", use_container_width=True):

            if not tc_assigned_workers:
                st.warning("작업자를 1명 이상 선택해 주세요.")
                return

            if risk_items and not selected_hazard:
                st.warning("중점위험요인을 선택해 주세요.")
                return

            if measure_items and not selected_measure:
                st.warning("중점위험요인의 대책을 선택해 주세요.")
                return

            try:
                target_dt = (client_dt or datetime.now())
                if tc_register_tomorrow:
                    target_dt = target_dt + timedelta(days=1)

                saved_task = create_work_task(
                    team_id=st.session_state.get("team_id", ""),
                    team_name=st.session_state.get("team_name", ""),
                    work_name=analysis["work_name"],
                    work_content=analysis["work_content"],
                    work_location=analysis["work_location"],
                    work_date=target_dt.strftime("%Y-%m-%d"),
                    scheduled_time=analysis["scheduled_time"],
                    tbm_place=analysis["tbm_place"],
                    assigned_workers=tc_assigned_workers,
                    risk_assessment_done=analysis["risk_assessment_done"],
                    work_type_code=analysis["work_type_code"],
                    work_type_display=analysis["work_type_display"],
                    start_time=analysis["start_time"],
                    time_slot=analysis["time_slot"],
                    material_name=analysis["material_name"],
                    chem_id=analysis["chem_id"],
                    chem_name=analysis["chem_name"],
                    risk_score=analysis["risk_score"],
                    risk_level=analysis["risk_level"],
                    main_hazard_1=analysis["main_hazard_1"],
                    main_hazard_2=analysis["main_hazard_2"],
                    main_hazard_3=analysis["main_hazard_3"],
                    safety_measure_1=analysis["safety_measure_1"],
                    safety_measure_2=analysis["safety_measure_2"],
                    safety_measure_3=analysis["safety_measure_3"],
                    selected_hazard=selected_hazard,
                    selected_measure=selected_measure,
                    similar_accident_text=analysis["similar_accident_text"]
                )

                if saved_task:
                    st.session_state.pop("task_analysis", None)
                    st.session_state.pop("tc_selected_hazard_idx", None)
                    st.session_state.pop("tc_selected_measure_idx", None)
                    if tc_register_tomorrow:
                        st.success("내일 예정 작업으로 등록되었습니다.")
                    else:
                        st.success("오늘 작업이 등록되었습니다.")
                    st.session_state.page = "manager"
                    st.rerun()
                else:
                    st.error("작업 등록 결과가 비어 있습니다.")

            except Exception as e:
                st.error("작업 등록 중 오류가 발생했습니다.")
                st.write(str(e))

def show_manager_dashboard():

    if not st.session_state.get("team_id"):
        st.warning("팀 접속 정보가 없습니다. 작업팀 접속 화면에서 다시 접속해 주세요.")

        if st.button("작업팀 접속 화면으로 이동", use_container_width=True):
            st.session_state.page = "team_access"
            st.query_params.clear()
            st.rerun()

        return

    # team_id만으로는 안전관리자 비밀번호 인증 여부를 알 수 없다 — 작업자가 URL 등으로
    # 이 화면에 직접 진입해 관리자 대시보드를 보는 것을 막는다.
    if st.session_state.get("mode") != "안전관리자":
        st.warning("안전관리자 인증이 필요합니다. 작업모드 선택 화면에서 안전관리자로 접속해 주세요.")

        if st.button("작업모드 선택 화면으로 이동", use_container_width=True):
            st.session_state.page = "login"
            st.rerun()

        return

    if not st.session_state.get("_signature_cleanup_done"):
        cleanup_expired_signatures(st.session_state.get("team_id"))
        st.session_state["_signature_cleanup_done"] = True

    # 안전관리자 기기의 로컬 날짜. "오늘 작업" 판정에 서버(UTC) 날짜를 쓰면
    # 한국 시간 00시~09시 사이에는 하루 어긋난 작업 목록이 나오므로 이걸 기준으로 삼는다.
    manager_client_dt = get_client_datetime()

    st.markdown("""
<style>
.manager-title {
    font-size: 24px;
    font-weight: 900;
    color: white;
}

.manager-summary-grid {
    display: grid;
    grid-template-columns: repeat(3, 1fr);
    gap: 10px;
    margin-bottom: 24px;
}

.manager-summary-card {
    background: white;
    border: 1px solid #d8dee9;
    border-radius: 16px;
    padding: 14px 10px;
    text-align: center;
    box-shadow: 0 4px 12px rgba(15,23,42,0.07);
}

.manager-summary-label {
    font-size: 12px;
    font-weight: 900;
    color: #45474c;
    margin-bottom: 8px;
}

.manager-summary-value {
    font-size: 28px;
    font-weight: 950;
    color: #091426;
}

.manager-summary-red {
    color: #ef4444;
}

.manager-summary-green {
    color: #16a34a;
}

.manager-section-title {
    font-size: 22px;
    font-weight: 900;
    color: #091426;
    margin: 24px 0 12px 0;
}

.log-card {
    background: white;
    border: 1px solid #d8dee9;
    border-radius: 16px;
    padding: 14px;
    margin-bottom: 10px;
    box-shadow: 0 3px 10px rgba(15,23,42,0.06);
}

.log-top-row {
    display: flex;
    align-items: center;
    justify-content: space-between;
    margin-bottom: 8px;
}

.log-work-name {
    font-size: 16px;
    font-weight: 900;
    color: #091426;
}

.log-meta {
    font-size: 13px;
    color: #45474c;
    line-height: 1.5;
}

.status-badge {
    display: inline-block;
    padding: 6px 10px;
    border-radius: 999px;
    font-size: 12px;
    font-weight: 900;
}

.status-done {
    background: #dcfce7;
    color: #15803d;
}

.status-progress {
    background: #ffedd5;
    color: #ea580c;
}

.tbm-history-row {
    padding: 0;
    line-height: 1.3;
}

.tbm-history-date {
    font-size: 13px;
    color: #45474c;
    white-space: nowrap;
}

.tbm-history-name {
    font-size: 15px;
    font-weight: 700;
    color: #091426;
}

div[data-testid="stMarkdownContainer"] hr.tbm-history-divider {
    border: none;
    border-top: 1px solid #e5e7eb;
    margin: 0 !important;
}

.st-key-tbm_history_list button[data-testid="stBaseButton-secondary"] {
    padding: 0.1rem 0.5rem;
    height: auto;
    min-height: 0;
    line-height: 1.3;
}

.st-key-tbm_history_list div[data-testid="stVerticalBlock"] {
    gap: 0.1rem;
}

.st-key-tbm_history_list div[data-testid="stElementContainer"] {
    margin: 0 !important;
}
</style>
""", unsafe_allow_html=True)

    render_topbar("manager", "안전관리자 대시보드", "manager-title", help_key="manager")

    if st.button("📝 오늘 작업 입력", key="goto_task_create_btn", use_container_width=True):
        st.session_state.page = "task_create"
        st.rerun()

    st.markdown('<div class="manager-section-title">오늘 등록된 작업</div>', unsafe_allow_html=True)

    try:
        today_tasks = get_today_work_tasks(
            st.session_state.get("team_id", ""),
            client_date=(manager_client_dt or datetime.now()).strftime("%Y-%m-%d")
        )
    except Exception as e:
        st.error("오늘 작업 목록을 불러오지 못했습니다.")
        st.write(str(e))
        today_tasks = []

    if not today_tasks:
        st.caption("아직 등록된 작업이 없습니다.")
    else:
        for task in today_tasks:
            if st.button(
                f"📋 {task.get('work_name', '-')}",
                key=f"task_detail_btn_{task.get('id')}",
                use_container_width=True
            ):
                st.session_state.selected_task_id = task.get("id")
                st.session_state.page = "task_detail"
                st.query_params.clear()
                st.rerun()

    st.markdown('<div class="manager-section-title">내일 예정된 작업</div>', unsafe_allow_html=True)

    try:
        tomorrow_date = ((manager_client_dt or datetime.now()) + timedelta(days=1)).strftime("%Y-%m-%d")
        tomorrow_tasks = get_work_tasks_by_date(st.session_state.get("team_id", ""), tomorrow_date)
    except Exception as e:
        st.error("내일 예정된 작업 목록을 불러오지 못했습니다.")
        st.write(str(e))
        tomorrow_tasks = []

    if not tomorrow_tasks:
        st.caption("내일 예정된 작업이 없습니다.")
    else:
        for task in tomorrow_tasks:

            task_workers = ", ".join(task.get("assigned_workers") or [])

            tomorrow_task_card_html = f"""
    <div class="log-card">
    <div class="log-top-row">
    <div class="log-work-name">{task.get("work_name", "-")}</div>
    <span class="status-badge status-progress">
    📅 내일 예정 · {task.get("scheduled_time", "-")}
    </span>
    </div>

    <div class="log-meta">
    작업내용: {task.get("work_content", "-")}<br>
    작업장소: {task.get("work_location", "-")}<br>
    TBM 장소: {task.get("tbm_place", "-")}<br>
    작업자: {task_workers}
    </div>
    </div>
    """

            st.markdown(
                tomorrow_task_card_html,
                unsafe_allow_html=True
            )

    st.markdown('<div class="manager-section-title">작업자 관리</div>', unsafe_allow_html=True)

    new_worker_name = st.text_input(
        "작업자 추가",
        placeholder="추가할 작업자 이름 입력",
        key="manager_add_worker_input"
    )

    if st.button("➕ 작업자 추가", key="manager_add_worker_btn", use_container_width=True):
        if not new_worker_name.strip():
            st.warning("추가할 작업자 이름을 입력해 주세요.")
        else:
            success, message = add_worker_to_team(
                st.session_state.get("team_id", ""),
                new_worker_name.strip()
            )

            if success:
                st.success(message)
                st.rerun()
            else:
                st.warning(message)
    try:
        team_id = st.session_state.get("team_id", "")

        if not team_id:
            workers = []
        else:
            team_result = (
                supabase.table("teams")
                .select("workers")
                .eq("id", team_id)
                .execute()
            )

            workers = (
                team_result.data[0].get("workers", [])
                if team_result.data
                else []
            )

        if workers:
            st.write("등록된 작업자")
            for worker in workers:
                st.markdown(f"- {worker}")
        else:
            st.caption("등록된 작업자가 없습니다.")

    except Exception as e:
        st.error("작업자 목록을 불러오지 못했습니다.")
        st.write(str(e))

    try:
        team_id = st.session_state.get("team_id", "")

        if not team_id:
            db_logs = []

        else:
            result = (
                supabase.table("work_logs")
                .select("*")
                .eq("team_id", team_id)
                .order("created_at", desc=True)
                .execute()
            )

            db_logs = result.data if result.data else []

    except Exception as e:
        st.error("작업로그를 불러오는 중 오류가 발생했습니다.")
        st.write(str(e))
        db_logs = []

    logs = []

    for log in db_logs:

        risk_score = log.get("risk_score")

        if isinstance(risk_score, (int, float)):
            risk_score_text = f"{risk_score:.0f}점"
        else:
            risk_score_text = "-"

        logs.append({
            "작업명": log.get("work_name", "-"),
            "작업자": log.get("worker_name", "-"),
            "작업시간": log.get("work_time", "-"),
            "위험도점수": risk_score_text,
            "위험등급": log.get("risk_level", "-"),
            "TBM완료여부": log.get("tbm_status", "-"),
            "TBM특이사항": log.get("checklist_remark", "-"),
            "작업특이사항": log.get("journal_note", "-"),
            "제출상태": log.get("submit_status", "-")
        })

    log_df = pd.DataFrame(logs)

    if "작업명" in log_df.columns:
        total_work_count = log_df["작업명"].nunique()
    else:
        total_work_count = 0

    high_risk_count = log_df[
        log_df["위험등급"].astype(str).str.contains("위험경고|고위험|🔴", na=False)
    ]["작업명"].nunique() if not log_df.empty else 0

    working_count = len(
        log_df[log_df["제출상태"].astype(str) != "제출완료"]
    ) if "제출상태" in log_df.columns else 0

    st.markdown(f"""
<div class="manager-section-title">오늘의 작업 현황</div>
<div class="manager-summary-grid">
    <div class="manager-summary-card">
        <div class="manager-summary-label">우리 팀 작업 수</div>
        <div class="manager-summary-value">{total_work_count}<span style="font-size:14px;">건</span></div>
    </div>
    <div class="manager-summary-card">
        <div class="manager-summary-label">고위험 작업 수</div>
        <div class="manager-summary-value manager-summary-red">{high_risk_count}<span style="font-size:14px;">건</span></div>
    </div>
    <div class="manager-summary-card">
        <div class="manager-summary-label">현재 작업중 팀원</div>
        <div class="manager-summary-value manager-summary-green">{working_count}<span style="font-size:14px;">명</span></div>
    </div>
</div>
""", unsafe_allow_html=True)

    st.markdown('<div class="manager-section-title">TBM 이력</div>', unsafe_allow_html=True)

    try:
        past_tasks = get_recent_past_work_tasks(
            st.session_state.get("team_id", ""),
            today_str=(manager_client_dt or datetime.now()).strftime("%Y-%m-%d"),
            days=7
        )
    except Exception as e:
        st.error("TBM 이력을 불러오지 못했습니다.")
        st.write(str(e))
        past_tasks = []

    if not past_tasks:
        st.caption("최근 7일 이내 지난 작업이 없습니다.")
    else:
        with st.container(key="tbm_history_list"):
            st.markdown('<hr class="tbm-history-divider">', unsafe_allow_html=True)

            for task in past_tasks:

                col_date, col_name, col_btn = st.columns([2, 5, 1], vertical_alignment="center")

                with col_date:
                    st.markdown(
                        f'<div class="tbm-history-row tbm-history-date">{task.get("work_date", "-")}</div>',
                        unsafe_allow_html=True
                    )

                with col_name:
                    st.markdown(
                        f'<div class="tbm-history-row tbm-history-name">{task.get("work_name", "-")}</div>',
                        unsafe_allow_html=True
                    )

                with col_btn:
                    try:
                        past_log_result = (
                            supabase.table("work_logs")
                            .select("*")
                            .eq("task_id", task.get("id", ""))
                            .execute()
                        )

                        past_task_logs = past_log_result.data if past_log_result.data else []

                        past_signatures_by_worker = fetch_signatures_for_task(task.get("id", ""))
                        past_docx_file = generate_tbm_docx(task, past_task_logs, past_signatures_by_worker)

                        st.download_button(
                            label="📋",
                            data=past_docx_file,
                            file_name=f"TBM회의록_{task.get('work_name', '작업')}_{task.get('work_date', '')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key=f"download_past_tbm_docx_{task.get('id')}",
                            help="TBM 회의록 다운로드",
                            use_container_width=True
                        )
                    except Exception as e:
                        st.error("⚠️")
                        st.write(str(e))

                st.markdown('<hr class="tbm-history-divider">', unsafe_allow_html=True)

    show_bottom_nav()

def show_task_detail():

    if not st.session_state.get("team_id"):
        st.warning("팀 접속 정보가 없습니다. 작업팀 접속 화면에서 다시 접속해 주세요.")

        if st.button("작업팀 접속 화면으로 이동", use_container_width=True):
            st.session_state.page = "team_access"
            st.query_params.clear()
            st.rerun()

        return

    if st.session_state.get("mode") != "안전관리자":
        st.warning("안전관리자 인증이 필요합니다. 작업모드 선택 화면에서 안전관리자로 접속해 주세요.")

        if st.button("작업모드 선택 화면으로 이동", use_container_width=True):
            st.session_state.page = "login"
            st.rerun()

        return

    task_id = st.session_state.get("selected_task_id")

    if not task_id:
        _redirect_with_message("manager", "먼저 대시보드에서 작업을 선택해 주세요.")
        return

    try:
        task_result = (
            supabase.table("work_tasks")
            .select("*")
            .eq("id", task_id)
            .limit(1)
            .execute()
        )
        task = task_result.data[0] if task_result.data else None
    except Exception as e:
        st.error("작업 정보를 불러오지 못했습니다.")
        st.write(str(e))
        task = None

    if not task:
        _redirect_with_message("manager", "해당 작업을 찾을 수 없습니다.")
        return

    st.markdown("""
<style>
.task-detail-title {
    font-size: 22px;
    font-weight: 900;
    color: #091426;
}

.manager-section-title {
    font-size: 22px;
    font-weight: 900;
    color: #091426;
    margin: 24px 0 12px 0;
}

.log-card {
    background: white;
    border: 1px solid #d8dee9;
    border-radius: 16px;
    padding: 14px;
    margin-bottom: 10px;
    box-shadow: 0 3px 10px rgba(15,23,42,0.06);
}

.log-top-row {
    display: flex;
    align-items: center;
    justify-content: space-between;
    margin-bottom: 8px;
}

.log-work-name {
    font-size: 16px;
    font-weight: 900;
    color: #091426;
}

.log-meta {
    font-size: 13px;
    color: #45474c;
    line-height: 1.5;
}

.status-badge {
    display: inline-block;
    padding: 6px 10px;
    border-radius: 999px;
    font-size: 12px;
    font-weight: 900;
    margin-right: 6px;
}

.status-done {
    background: #dcfce7;
    color: #15803d;
}

.status-progress {
    background: #ffedd5;
    color: #ea580c;
}
</style>
""", unsafe_allow_html=True)

    render_topbar("task_detail", task.get("work_name", "작업 상세"), "task-detail-title", back_page="manager", help_key="manager")
    _flash_pending_message()

    task_info_html = f"""
<div class="log-card">
    <div class="log-meta">
    작업내용: {task.get("work_content", "-")}<br>
    작업장소: {task.get("work_location", "-")}<br>
    TBM 장소: {task.get("tbm_place", "-")}<br>
    예정시간: {task.get("scheduled_time", "-")}
    </div>
</div>
"""
    st.markdown(task_info_html, unsafe_allow_html=True)

    try:
        log_result = (
            supabase.table("work_logs")
            .select("*")
            .eq("task_id", task_id)
            .execute()
        )
        task_logs = log_result.data if log_result.data else []
    except Exception as e:
        st.error("작업로그를 불러오지 못했습니다.")
        st.write(str(e))
        task_logs = []

    # (task_id, worker_name) 기준으로 하나의 work_logs 행에 체크리스트→작업일지가
    # 이어서 채워지는 구조이므로, worker_name별 최신 행 하나만 있으면 된다.
    logs_by_worker = {log.get("worker_name"): log for log in task_logs}

    assigned_workers = task.get("assigned_workers") or []

    st.markdown('<div class="manager-section-title">작업자별 제출 현황</div>', unsafe_allow_html=True)

    if not assigned_workers:
        st.caption("배정된 작업자가 없습니다.")
    else:
        for worker_name in assigned_workers:
            worker_log = logs_by_worker.get(worker_name)
            submit_status = worker_log.get("submit_status") if worker_log else None

            # work_logs.submit_status는 체크리스트 제출 시 "TBM완료", 작업일지 제출 시
            # "제출완료"로 갱신된다. 행이 아예 없으면 TBM조차 아직 안 낸 것이다.
            tbm_done = submit_status in ("TBM완료", "제출완료")
            journal_done = submit_status == "제출완료"

            tbm_badge = (
                '<span class="status-badge status-done">TBM 완료</span>'
                if tbm_done else
                '<span class="status-badge status-progress">TBM 미완료</span>'
            )
            journal_badge = (
                '<span class="status-badge status-done">작업일지 완료</span>'
                if journal_done else
                '<span class="status-badge status-progress">작업일지 미완료</span>'
            )

            worker_card_html = f"""
<div class="log-card">
    <div class="log-top-row">
        <div class="log-work-name">{worker_name}</div>
    </div>
    <div>{tbm_badge}{journal_badge}</div>
</div>
"""
            st.markdown(worker_card_html, unsafe_allow_html=True)

    st.markdown('<div class="manager-section-title">TBM 회의록 / 작업로그</div>', unsafe_allow_html=True)

    try:
        signatures_by_worker = fetch_signatures_for_task(task_id)
        docx_file = generate_tbm_docx(task, task_logs, signatures_by_worker)

        downloaded = st.download_button(
            label="📄 TBM 회의록 출력하기",
            data=docx_file,
            file_name=f"TBM회의록_{task.get('work_name', '작업')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key=f"download_tbm_docx_detail_{task_id}",
            use_container_width=True
        )

        if downloaded and not task.get("tbm_docx_generated_at"):
            try:
                supabase.table("work_tasks").update({
                    "tbm_docx_generated_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                }).eq("id", task_id).execute()
            except Exception as e:
                print("tbm_docx_generated_at 업데이트 오류:", e)

    except Exception as e:
        st.error("TBM 회의록 생성 실패")
        st.write(str(e))

    csv_rows = []
    for log in task_logs:
        risk_score = log.get("risk_score")
        risk_score_text = f"{risk_score:.0f}점" if isinstance(risk_score, (int, float)) else "-"

        csv_rows.append({
            "작업명": log.get("work_name", "-"),
            "작업자": log.get("worker_name", "-"),
            "작업시간": log.get("work_time", "-"),
            "위험도점수": risk_score_text,
            "위험등급": log.get("risk_level", "-"),
            "TBM완료여부": log.get("tbm_status", "-"),
            "TBM특이사항": log.get("daily_safety_check_result", "-"),
            "작업특이사항": log.get("closing_meeting_result", "-"),
            "제출상태": log.get("submit_status", "-")
        })

    csv_data = pd.DataFrame(csv_rows).to_csv(index=False, encoding="utf-8-sig")

    st.download_button(
        label="📥 이 작업의 작업로그 CSV 다운로드",
        data=csv_data,
        file_name=f"작업로그_{task.get('work_name', '작업')}.csv",
        mime="text/csv",
        key=f"download_task_csv_{task_id}",
        use_container_width=True
    )

    show_bottom_nav()

# ---------- OPENAI API키 ----------------
client = OpenAI(
    api_key=st.secrets["OPENAI_API_KEY"]
)
supabase = create_client(
    st.secrets["SUPABASE_URL"],
    st.secrets["SUPABASE_KEY"]
)

# =========================
# 2) 기본 설정
# =========================
SERVICE_KEY = st.secrets["DATA_API_KEY"]

strength_mapping = {
    0: 0,
    1: 1,
    2: 3,
    3: 6
}

hazard_features = [
    "금속부식성 물질",
    "급성 독성(흡입)",
    "인화성",
    "급성 수생환경 유해성",
    "산화성",
    "고압가스",
    "자연발화 및 과산화물",
]

col_alias = {
    "급성독성(흡입)": "급성 독성(흡입)",
    "급성 독성(흡입)": "급성 독성(흡입)",
    "금속부식성": "금속부식성 물질",
    "금속부식성 물질": "금속부식성 물질",
    "수생환경": "급성 수생환경 유해성",
    "급성 수생환경 유해성": "급성 수생환경 유해성",
    "자연발화과산화물": "자연발화 및 과산화물",
    "자연발화 및 과산화물": "자연발화 및 과산화물",
    "인화성": "인화성",
    "산화성": "산화성",
    "고압가스": "고압가스",
}

interaction_rules = [
    ("인화성", "work_type_HOT_WORK", "인화성_HOT_WORK"),
    ("인화성", "work_type_MAINTENANCE", "인화성_MAINTENANCE"),
    ("인화성", "work_type_CLEANING", "인화성_CLEANING"),
    ("인화성", "work_type_STARTUP_SHUTDOWN", "인화성_STARTUP_SHUTDOWN"),
    ("고압가스", "work_type_MAINTENANCE", "고압가스_MAINTENANCE"),
    ("고압가스", "work_type_STARTUP_SHUTDOWN", "고압가스_STARTUP_SHUTDOWN"),
    ("산화성", "work_type_HOT_WORK", "산화성_HOT_WORK"),
    ("산화성", "work_type_CLEANING", "산화성_CLEANING"),
    ("산화성", "work_type_MAINTENANCE", "산화성_MAINTENANCE"),
    ("급성 독성(흡입)", "work_type_MAINTENANCE", "급성독성흡입_MAINTENANCE"),
    ("급성 독성(흡입)", "work_type_CLEANING", "급성독성흡입_CLEANING"),
    ("급성 독성(흡입)", "work_type_STARTUP_SHUTDOWN", "급성독성흡입_STARTUP_SHUTDOWN"),
    ("금속부식성 물질", "work_type_MAINTENANCE", "금속부식성_MAINTENANCE"),
    ("금속부식성 물질", "work_type_CLEANING", "금속부식성_CLEANING"),
    ("자연발화 및 과산화물", "work_type_MAINTENANCE", "자연발화과산화물_MAINTENANCE"),
    ("자연발화 및 과산화물", "work_type_CLEANING", "자연발화과산화물_CLEANING"),
    ("자연발화 및 과산화물", "work_type_STARTUP_SHUTDOWN", "자연발화과산화물_STARTUP_SHUTDOWN"),
    ("급성 수생환경 유해성", "work_type_CLEANING", "수생환경_CLEANING"),
    ("급성 수생환경 유해성", "work_type_MAINTENANCE", "수생환경_MAINTENANCE"),
]

# =========================
# 3) API 함수
# =========================
def test_msds_api_by_cas(cas_no, service_key):
    url = "https://apis.data.go.kr/B552468/msdschem/getChemList"
    params = {
        "serviceKey": service_key,
        "searchWrd": cas_no,
        "searchCnd": "1",
        "numOfRows": "10",
        "pageNo": "1"
    }
    response = requests.get(url, params=params, timeout=10)
    return response.status_code, response.text


def resolve_chemical_alias(name):
    if not name:
        return name

    key = name.strip()

    if key in CHEMICAL_ALIASES:
        return CHEMICAL_ALIASES[key]

    key_cf = key.casefold()
    for alias, official in CHEMICAL_ALIASES.items():
        if alias.casefold() == key_cf:
            return official

    return key


def search_chemical_names_api(search_term, service_key, num_of_rows=10):
    if not search_term or not search_term.strip():
        return []

    url = "https://apis.data.go.kr/B552468/msdschem/getChemList"
    params = {
        "serviceKey": service_key,
        "searchWrd": search_term.strip(),
        "searchCnd": "0",  # 국문명 검색
        "numOfRows": str(num_of_rows),
        "pageNo": "1"
    }

    try:
        response = requests.get(url, params=params, timeout=10)

        if response.status_code != 200:
            return []

        root = ET.fromstring(response.text)

        if root.findtext(".//resultCode") != "00":
            return []

        names = []
        for item in root.findall(".//item"):
            chem_name_kor = item.findtext("chemNameKor")
            if chem_name_kor and chem_name_kor not in names:
                names.append(chem_name_kor)

        return names

    except Exception:
        return []


def search_chemical_candidates(searchterm):
    if not searchterm or not searchterm.strip():
        return []

    term = searchterm.strip()
    term_cf = term.casefold()
    candidates = []

    # 1) 관용명 부분일치 → 정식명 후보 추가 (보조 기능)
    for alias, official in CHEMICAL_ALIASES.items():
        if term_cf in alias.casefold() and official not in candidates:
            candidates.append(official)

    # 2) 관용명이 매칭되면 정식명으로, 아니면 입력값 그대로 MSDS API 검색
    api_term = resolve_chemical_alias(term)
    for name in search_chemical_names_api(api_term, SERVICE_KEY):
        if name not in candidates:
            candidates.append(name)

    # 3) 관용명 변환이 있었던 경우, 원본 입력어로도 검색해 후보를 보강
    if api_term != term:
        for name in search_chemical_names_api(term, SERVICE_KEY):
            if name not in candidates:
                candidates.append(name)

    return candidates[:15]


def get_chemid_by_name(chem_name, service_key):
    url = "https://apis.data.go.kr/B552468/msdschem/getChemList"

    params = {
        "serviceKey": service_key,
        "searchWrd": chem_name,
        "searchCnd": "0",  # 🔥 국문명 검색
        "numOfRows": "10",
        "pageNo": "1"
    }

    response = requests.get(url, params=params, timeout=10)

    if response.status_code != 200:
        return None, None, f"API 호출 실패: {response.status_code}"

    root = ET.fromstring(response.text)

    if root.findtext(".//resultCode") != "00":
        return None, None, root.findtext(".//resultMsg")

    items = root.findall(".//item")

    if not items:
        return None, None, "검색 결과 없음"

    # 🔥 첫 번째 결과 사용
    chem_id = items[0].findtext("chemId")
    chem_name_kor = items[0].findtext("chemNameKor")

    return chem_id, chem_name_kor, None

    return None, None, "물질명과 정확히 일치하는 물질이 없습니다."


def get_hazard_by_chemid(chem_id, service_key):
    url = "https://apis.data.go.kr/B552468/msdschem/getChemDetail02"
    params = {
        "serviceKey": service_key,
        "chemId": chem_id
    }
    response = requests.get(url, params=params, timeout=10)
    return response.status_code, response.text


def extract_classification_text(detail_xml):
    root = ET.fromstring(detail_xml)

    for item in root.findall(".//item"):
        if item.findtext("msdsItemNameKor") == "유해성·위험성 분류":
            return item.findtext("itemDetail")

    return ""


def parse_api_classification(classification_text):
    rows = []

    if not classification_text:
        return rows

    for item in classification_text.split("|"):
        item = item.strip()

        if ":" not in item:
            continue

        hazard, category = item.split(":", 1)

        rows.append({
            "API_위해성": hazard.strip(),
            "세부분류": category.strip()
        })

    return rows


def normalize_text(x):
    return (
        str(x)
        .replace(" ", "")
        .replace("\u3000", "")
        .replace("\n", "")
        .replace("\r", "")
        .replace("\t", "")
        .strip()
    )


def map_hazard_scores_by_excel(classification_text, mapping_df):
    hazard_scores = {}
    parsed_rows = parse_api_classification(classification_text)

    temp = mapping_df.copy()
    temp["_api_hazard_norm"] = temp["API_위해성"].apply(normalize_text)
    temp["_category_norm"] = temp["세부분류"].apply(normalize_text)

    for row in parsed_rows:
        api_hazard = normalize_text(row["API_위해성"])
        api_category = normalize_text(row["세부분류"])

        matched = temp[
            (temp["_api_hazard_norm"] == api_hazard) &
            (temp["_category_norm"] == api_category)
        ]

        if matched.empty:
            continue

        for _, m in matched.iterrows():
            model_group = str(m["물질군"]).strip()
            raw_score = pd.to_numeric(m["점수"], errors="coerce")

            if pd.isna(raw_score):
                continue

            if model_group not in hazard_scores:
                hazard_scores[model_group] = 0

            hazard_scores[model_group] = max(hazard_scores[model_group], float(raw_score))

    return hazard_scores

# =========================
# 4) 점수 보정 함수
# =========================
def convert_strength_score(raw_score):
    raw_score = pd.to_numeric(raw_score, errors="coerce")

    if pd.isna(raw_score):
        return 0

    raw_score = int(raw_score)
    return strength_mapping.get(raw_score, 0)


def sigmoid_rescale(score, k=0.08, center=30):
    return 100 / (1 + np.exp(-k * (score - center)))


def score_to_level(score):
    if score <= 40:
        return "🟢 안전유의"
    elif score <= 70:
        return "🟡 작업주의"
    else:
        return "🔴 위험경고"


def calculate_final_score(input_df, pred_prob, work_type, time_slot, chem_info_missing=0):
    base_prob = pred_prob

    damage_boost = 0

    if "time_slot_야간" in input_df.columns and input_df.loc[0, "time_slot_야간"] == 1:
        damage_boost += 0.10

    if "인화성" in input_df.columns and input_df.loc[0, "인화성"] > 0:
        damage_boost += 0.05

    uncertainty_boost = 0

    if chem_info_missing == 1:
        work_uncertainty_map = {
            "HOT_WORK": 0.15,
            "MAINTENANCE": 0.12,
            "CLEANING": 0.10,
            "STARTUP_SHUTDOWN": 0.10,
            "ROUTINE": 0.05,
            "IDLE": 0.00
        }

        time_uncertainty_map = {
            "야간": 0.05,
            "저녁": 0.03,
            "주간": 0.00,
            "미상": 0.00,
            "unknown": 0.00
        }

        uncertainty_boost += work_uncertainty_map.get(work_type, 0.05)
        uncertainty_boost += time_uncertainty_map.get(time_slot, 0.03)

    if time_slot in ["미상", "unknown"] and work_type in ["HOT_WORK", "MAINTENANCE"]:
        time_unknown_boost = 0.12
    elif time_slot in ["미상", "unknown"]:
        time_unknown_boost = 0.05
    else:
        time_unknown_boost = 0

    combined_prob = base_prob * (
        1
        + damage_boost
        + uncertainty_boost
        + time_unknown_boost
    )

    combined_prob = np.clip(combined_prob, 0, 1)

    raw_score = combined_prob * 100
    final_score = sigmoid_rescale(raw_score, k=0.08, center=30)
    final_score = float(np.clip(final_score, 0, 100))

    return final_score, score_to_level(final_score), {
        "base_prob": base_prob,
        "damage_boost": damage_boost,
        "uncertainty_boost": uncertainty_boost,
        "time_unknown_boost": time_unknown_boost,
        "combined_prob": combined_prob,
        "raw_score": raw_score
    }
def get_client_datetime():
    """
    서버 시간 대신 접속 기기(모바일 등)의 로컬 시간을 가져온다.
    서버가 배포된 곳과 접속 기기의 시간대/시계가 다르면 datetime.now()(서버 시간)는
    실제 기기 시간과 어긋나므로, 브라우저의 Date 객체 값을 그대로 받아온다.
    값이 아직 브라우저에서 도착하지 않았으면 None을 반환한다(호출부에서 서버 시간으로 대체).
    """
    js_now = streamlit_js_eval(
        js_expressions="""
            (function() {
                var d = new Date();
                function pad(n) { return n < 10 ? '0' + n : '' + n; }
                return d.getFullYear() + '-' + pad(d.getMonth() + 1) + '-' + pad(d.getDate())
                    + ' ' + pad(d.getHours()) + ':' + pad(d.getMinutes()) + ':' + pad(d.getSeconds());
            })()
        """,
        key="client_now",
    )
    if not js_now:
        return None
    try:
        return datetime.strptime(js_now, "%Y-%m-%d %H:%M:%S")
    except (TypeError, ValueError):
        return None

def get_current_time_slot(now=None):
    if now is None:
        now = datetime.now()
    current_hour = now.hour

    if 7 <= current_hour < 15:
        return "주간", now
    elif 15 <= current_hour < 23:
        return "저녁", now
    else:
        return "야간", now

# =========================
# 5) 모델 입력 생성
# =========================
def make_input_data(work_type, time_slot):
    input_df = pd.DataFrame(columns=model_columns)
    input_df.loc[0] = 0

    work_col = f"work_type_{work_type}"
    if work_col in input_df.columns:
        input_df.loc[0, work_col] = 1

    if work_col == "work_type_IDLE" and work_col in input_df.columns:
        input_df.loc[0, work_col] = 0.5

    time_col = f"time_slot_{time_slot}"
    if time_col in input_df.columns:
        input_df.loc[0, time_col] = 1

    hazard_scores = st.session_state.get("hazard_scores", {})

    for raw_col, raw_score in hazard_scores.items():
        model_col = col_alias.get(raw_col, raw_col)

        if model_col in input_df.columns:
            model_score = convert_strength_score(raw_score)
            input_df.loc[0, model_col] = model_score

    for hazard_col, work_col, new_col in interaction_rules:
        if hazard_col in input_df.columns and work_col in input_df.columns and new_col in input_df.columns:
            input_df.loc[0, new_col] = input_df.loc[0, hazard_col] * input_df.loc[0, work_col]

    input_df = input_df.apply(pd.to_numeric, errors="coerce").fillna(0)

    return input_df

# =========================
# 6) 화면
# =========================

if "hazard_scores" not in st.session_state:
    st.session_state.hazard_scores = {}

if "chem_id" not in st.session_state:
    st.session_state.chem_id = ""

if "chem_name" not in st.session_state:
    st.session_state.chem_name = ""

if "classification_text" not in st.session_state:
    st.session_state.classification_text = ""

if "result" not in st.session_state:
    st.session_state.result = {}

if "checklist_data" not in st.session_state:
    st.session_state.checklist_data = {}

if "checklist_remark" not in st.session_state:
    st.session_state.checklist_remark = ""

if "journal_data" not in st.session_state:
    st.session_state.journal_data = {}

if "journal_submitted" not in st.session_state:
    st.session_state.journal_submitted = False

if "work_logs" not in st.session_state:
    st.session_state.work_logs = []

if "work_log_csv_path" not in st.session_state:
    st.session_state.work_log_csv_path = "safety_tbm_work_logs.csv"



# =========================
# 페이지 실행부
# =========================
if st.session_state.page == "team_access":
    show_team_access()

elif st.session_state.page == "create_team":
    show_create_team()

elif st.session_state.page == "login":
    show_login()

elif st.session_state.page == "input":
    show_work_input()

elif st.session_state.page == "result":
    show_risk_result()

elif st.session_state.page == "task_info":
    show_task_info()

elif st.session_state.page == "checklist":
    show_checklist()

elif st.session_state.page == "journal":
    show_journal()

elif st.session_state.page == "manager":
    show_manager_dashboard()

elif st.session_state.page == "task_create":
    show_task_create()

elif st.session_state.page == "task_detail":
    show_task_detail()

show_active_help_popup()